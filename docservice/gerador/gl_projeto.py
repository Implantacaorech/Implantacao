# -*- coding: utf-8 -*-
"""Geração fiel — parte do PROJETO de Implantação (detalhamento + tabelas + respostas)."""
import db
import preencher_layout as PL
from gl_comum import (_conteudo, _num, _por_extenso, _hoje, _norm, _eh_marcador,
                      _inserir_textos_depois, _PROJ_GRUPOS)


def _repl_projeto(p):
    val = _conteudo(p, "projeto")
    cli = (p.get("cliente") or "").strip()
    repl, paras = [], []
    if cli:
        # ocorrências no corpo (literais ASCII) + linha-cabeçalho via prefixo
        repl += [("<Nome do Cliente >", cli), ("<Nome do Cliente>", cli)]
        paras.append(("Nome do Cliente:", "Nome do Cliente: %s" % cli))
    if val("cnpj", "cnpj"):
        paras.append(("CNPJ", "CNPJ: %s" % val("cnpj", "cnpj")))
    if val("objetivos", "observacoes"):
        repl.append(("<(preencher)>", val("objetivos", "observacoes")))
    # Equipes (telas de edição estruturada)
    if val("gerente_contas", "gci"):
        paras.append(("Gerente de Contas do Projeto", "Gerente de Contas do Projeto: %s" % val("gerente_contas", "gci")))
    if val("redator"):
        paras.append(("Redator do Projeto", "Redator do Projeto: %s" % val("redator")))
    if val("consultor", "consultor"):
        paras.append(("Consultor/Implantador", "Consultor/Implantador: %s" % val("consultor", "consultor")))
    # Lado CLIENTE das Equipes de Trabalho — o layout traz a linha "Encarregado pelo Projeto:"
    # em paralelo às três da Rech, mas ela nunca era preenchida: a tela de edição gravava
    # `encarregado` em DocConteudo e a geração ignorava, então o campo saía em branco em todo
    # Projeto (defeito herdado do webapp/gl_projeto.py do Flask).
    if val("encarregado", "contato_nome"):
        paras.append(("Encarregado pelo Projeto",
                      "Encarregado pelo Projeto: %s" % val("encarregado", "contato_nome")))
    hb, hc = _num(p.get("horas_bonificadas")), _num(p.get("horas_cobradas"))
    if hb:
        repl.append(("<XX horas bonificadas>", "%s horas bonificadas" % hb))
    if hc:
        repl.append(("<XX horas cobradas>", "%s horas cobradas" % hc))
    paras.append(("Novo Hamburgo", "Novo Hamburgo, %s." % _por_extenso(_hoje())))
    return repl, paras


def _linhas(texto):
    """Quebra um textarea da tela de edição em linhas não vazias."""
    return [l.strip() for l in (texto or "").splitlines() if l.strip()]


def _preencher_escopo_projeto(doc, projeto):
    """Escopo → 'Empresas contempladas' e 'Conversões': escreve os textos que a tela do
    Projeto grava em DocConteudo (`empresas`, `conversoes`).

    Os dois campos existem na tela desde que ela foi criada, mas nenhum gerador os lia — o
    marcador '(preencher)' do layout acabava apagado por `remover_marcadores_docx` e o texto
    digitado se perdia. Cada valor entra como parágrafo(s) logo abaixo do rótulo do layout,
    preservando o estilo do título (mesma mecânica dos blocos do Levantamento)."""
    cont = db.doc_conteudo(projeto.get("id"), "projeto") if projeto.get("id") else {}
    if not cont:
        return 0
    alvos = [
        ("estão contempladas no referido projeto as seguintes empresas", "empresas"),
        ("conversões", "conversoes"),
    ]
    n = 0
    for rotulo, campo in alvos:
        itens = _linhas(cont.get(campo))
        if not itens:
            continue
        for p in doc.paragraphs:
            if _norm(p.text).startswith(rotulo):
                _inserir_textos_depois(p, itens)
                n += 1
                break
    return n


def _emitir_lista_detalhamento(anchor_p, itens):
    """O modelo traz UM parágrafo-marcador (com bullet) para o detalhamento. Em vez de emendar
    tudo numa linha, emite UMA linha por item, clonando o parágrafo para preservar a formatação
    (bullet/indentação) do modelo."""
    import copy
    from docx.text.paragraph import Paragraph
    if not itens:
        return
    PL._aplica_no_paragrafo(anchor_p, itens[0])         # 1º item no próprio marcador
    el = anchor_p._p
    for txt in itens[1:]:
        novo = copy.deepcopy(anchor_p._p)               # clona pPr (numeração/indent) + run
        el.addnext(novo)
        PL._aplica_no_paragrafo(Paragraph(novo, anchor_p._parent), txt)
        el = novo


def _preencher_detalhamento_projeto(doc, projeto, guia=False):
    """No 'Detalhamento das Rotinas' do Projeto: mantém SÓ as áreas dos módulos
    contratados (remove as demais e os grupos que ficarem vazios) e preenche cada
    bloco com os dados do Levantamento — 'Módulos Previstos' = módulos contratados
    da área; 'Detalhamento' = respostas do Levantamento da área. A tela do Projeto
    (DocConteudo det_<area>_*) sobrepõe quando preenchida. Com `guia=True` (modelo
    para preenchimento manual), sem respostas usa as PERGUNTAS do Índice como guia."""
    import doc_edit, re as _re
    pid = projeto.get("id")
    sigs = {m.strip().upper() for m in _re.split(r"[,;\n]+", projeto.get("modulos", "") or "") if m.strip()}
    nomes = {m["sigla"].upper(): m["nome"] for m in db.indice_modulos()}
    cont = db.doc_conteudo(pid, "projeto") if pid else {}

    # área (nome normalizado) -> (k, siglas contratadas) e sigla -> área k
    area_info, sigla_area, inter_by_k = {}, {}, {}
    for (k, nome, ss) in doc_edit._PROJ_AREAS:
        inter = [s for s in ss if s in sigs]
        area_info[_norm(nome)] = (k, inter)
        inter_by_k[k] = inter
        for s in ss:
            sigla_area[s] = k
    contratadas = {a for a, (k, inter) in area_info.items() if inter}

    # respostas do Levantamento por área (k) — alimentam o 'Detalhamento'
    resp_por_k = {}
    for r in (db.levantamento_respostas(pid) if pid else []):
        k = sigla_area.get((r.get("modulo_sigla") or "").upper())
        resp = (r.get("resposta") or "").strip()
        if k and resp:
            resp_por_k.setdefault(k, []).append("%s: %s" % ((r.get("topico") or "").strip(), resp))

    def _valor(k, campo):
        v = (cont.get("det_%s_%s" % (k, campo)) or "").strip()
        if v:
            return v
        if campo == "modulos":
            return ", ".join(("%s — %s" % (s, nomes[s])) if nomes.get(s) else s for s in inter_by_k.get(k, []))
        if campo == "detalhamento":
            base = resp_por_k.get(k)
            if base:
                return base                # lista -> 1 linha por tópico (respeita o modelo)
            if guia:                       # modelo manual: usa as perguntas do Índice como guia
                qs = []
                for sig in inter_by_k.get(k, []):
                    tops, _ = db.indice_listar(modulo=sig)
                    qs += [(t.get("topico") or "").strip() for t in tops if (t.get("topico") or "").strip()]
                return qs
        return ""

    rotulos = [("módulos previsto", "modulos"), ("detalhamento das rotinas", "detalhamento"),
               ("particularidade", "particularidade"), ("não está previsto", "naoprevisto")]

    paras = doc.paragraphs
    ini = next((i for i, p in enumerate(paras) if _norm(p.text) == "detalhamento das rotinas"), None)
    if ini is None:
        return 0
    fim = next((i for i in range(ini + 1, len(paras)) if _norm(paras[i].text).startswith("responsabilidades")), len(paras))

    # segmenta o trecho em grupos e áreas (com seus parágrafos de conteúdo)
    segs, cur = [], None
    for i in range(ini + 1, fim):
        p = paras[i]
        tl = _norm(p.text)
        if tl in _PROJ_GRUPOS:
            cur = {"kind": "group", "name": tl, "head": p}
            segs.append(cur)
        elif tl in area_info:
            cur = {"kind": "area", "name": tl, "head": p, "paras": []}
            segs.append(cur)
        elif cur and cur["kind"] == "area":
            cur["paras"].append(p)

    remover, n = [], 0
    for seg in segs:
        if seg["kind"] == "group":
            if not (_PROJ_GRUPOS[seg["name"]] & contratadas):
                remover.append(seg["head"])
            continue
        if seg["name"] not in contratadas:          # área não contratada -> remove o bloco inteiro
            remover.append(seg["head"])
            remover.extend(seg["paras"])
            continue
        k, campo = area_info[seg["name"]][0], None    # área contratada -> preenche os <XX>
        for p in seg["paras"]:
            tl = _norm(p.text)
            lab = next((c for (kw, c) in rotulos if kw in tl), None) if tl else None
            if lab:
                campo = lab
            elif campo and _eh_marcador(p.text):
                val = _valor(k, campo)
                if isinstance(val, list):
                    if val:                                  # detalhamento -> uma linha por item
                        _emitir_lista_detalhamento(p, val)
                        n += 1
                elif val:
                    PL._aplica_no_paragrafo(p, val)
                    n += 1
                campo = None

    for p in remover:
        el = p._p
        if el.getparent() is not None:
            el.getparent().remove(el)
    return n


def _preencher_projeto_tabelas(doc, projeto):
    """Preenche no Projeto a Tabela de Usuários e o Cronograma Macro a partir do DocConteudo."""
    cont = db.doc_conteudo(projeto.get("id"), "projeto") if projeto.get("id") else {}
    if not cont:
        return
    crono = [("levantamento de requisitos", "crono_levantamento"),
             ("elaboração do cronograma", "crono_cronograma"), ("elaboracao do cronograma", "crono_cronograma"),
             ("parametriz", "crono_parametrizacao"), ("treinamento", "crono_treinamento"),
             ("simula", "crono_simulacao"), ("início do uso", "crono_inicio"), ("inicio do uso", "crono_inicio"),
             ("finaliza", "crono_finalizacao")]
    for t in doc.tables:
        hdr = [(c.text or "").strip().lower() for c in t.rows[0].cells] if t.rows else []
        if not hdr:
            continue
        # Cronograma Macro: Fase | Etapa | Período previsto
        if hdr[0] == "fase" and any("previsto" in h for h in hdr):
            col_per = next((j for j, h in enumerate(hdr) if "previsto" in h), len(hdr) - 1)
            col_et = next((j for j, h in enumerate(hdr) if h == "etapa"), 1)
            for row in t.rows[1:]:
                et = (row.cells[col_et].text or "").strip().lower()
                key = next((k for (kw, k) in crono if kw in et), None)
                if key and cont.get(key):
                    row.cells[col_per].text = cont[key]
        # Tabela de Usuários: Nome | E-mail | Área | Assina
        if hdr[0] == "nome" and any("assina" in h for h in hdr):
            base = t.rows[1:]
            usuarios = []
            # 5 linhas: mesma quantidade dos Usuários-chave do Levantamento, de onde a
            # etapa 10 herda a tabela. `t.add_row()` cobre o que passar do modelo.
            for i in range(5):
                nome = (cont.get("usu_%d_nome" % i) or "").strip()
                if nome:
                    usuarios.append([nome, cont.get("usu_%d_email" % i, ""),
                                     cont.get("usu_%d_area" % i, ""), cont.get("usu_%d_assina" % i, "")])
            for idx, u in enumerate(usuarios):
                row = base[idx] if idx < len(base) else t.add_row()
                for j, v in enumerate(u):
                    if j < len(row.cells):
                        row.cells[j].text = v
