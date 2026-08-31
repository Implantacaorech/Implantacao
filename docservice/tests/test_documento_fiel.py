# -*- coding: utf-8 -*-
"""Testes do endpoint /gerar/documento-fiel (Levantamento, Projeto, Cronograma, Termo) —
espelha webapp/test_painel.py:test_gerar_layout_* usando os templates reais de
tools/templates/layouts/ (os mesmos que o Cadastro de Modelos semeia em produção)."""
import base64
import io
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

import openpyxl
from docx import Document
from fastapi.testclient import TestClient

from main import app

client = TestClient(app)

_LAYOUTS = Path(__file__).resolve().parents[2] / "tools" / "templates" / "layouts"


def _template_base64(nome):
    return base64.b64encode((_LAYOUTS / nome).read_bytes()).decode("ascii")


def _texto_completo(doc):
    return "\n".join(p.text for p in doc.paragraphs)


def test_gera_termo_preenche_cliente_numero_projeto_e_grade():
    payload = {
        "slug": "termo",
        "modeloBase64": _template_base64("termo.docx"),
        "projeto": {
            "id": 1,
            "cliente": "Cliente Teste LTDA",
            "numeroProjeto": "PRJ-42",
            "dataEncerramento": "2026-08-20",
            "modulos": "FAT, EST",
        },
    }
    r = client.post("/gerar/documento-fiel", json=payload)
    assert r.status_code == 200
    assert r.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    doc = Document(io.BytesIO(r.content))
    texto = _texto_completo(doc)
    assert "Cliente Teste LTDA" in texto
    assert "PRJ-42" in texto
    assert "<Razão Social Longa>" not in texto
    assert "<Número do projeto quando se aplicar>" not in texto

    grade = doc.tables[0]
    linhas = [[c.text.strip() for c in row.cells] for row in grade.rows[1:]]
    assert linhas[0][0] == "FAT"
    assert linhas[0][2] == "Implantado"
    assert linhas[0][3] == "Sim"
    assert linhas[1][0] == "EST"


def test_gera_levantamento_preenche_cliente_e_data():
    payload = {
        "slug": "levantamento",
        "modeloBase64": _template_base64("levantamento.docx"),
        "projeto": {
            "id": 1,
            "cliente": "Cliente Teste LTDA",
            "dataLevantamento": "2026-07-10",
            "gci": "Ana",
            "consultor": "Beto",
            "modulos": "FAT",
        },
        "indiceModulos": [{"sigla": "FAT", "nome": "Faturamento"}],
        "indiceTopicos": [
            {"moduloSigla": "FAT", "topico": "Emissão de nota fiscal", "adicional": ""}
        ],
    }
    r = client.post("/gerar/documento-fiel", json=payload)
    assert r.status_code == 200
    doc = Document(io.BytesIO(r.content))
    texto = _texto_completo(doc)
    assert "Cliente Teste LTDA" in texto
    assert "10/07/2026" in texto
    assert "Ana / Beto" in texto
    assert "< Nome Cliente >" not in texto


def test_gera_projeto_modo_modelo_preenche_cliente():
    payload = {
        "slug": "projeto",
        "modo": "modelo",
        "modeloBase64": _template_base64("projeto.docx"),
        "projeto": {
            "id": 1,
            "cliente": "Cliente Teste LTDA",
            "cnpj": "00.000.000/0001-00",
            "modulos": "FAT",
        },
    }
    r = client.post("/gerar/documento-fiel", json=payload)
    assert r.status_code == 200
    doc = Document(io.BytesIO(r.content))
    texto = _texto_completo(doc)
    assert "Nome do Cliente: Cliente Teste LTDA" in texto
    assert "CNPJ: 00.000.000/0001-00" in texto


def test_gera_projeto_modo_auto_preenche_escopo_equipes_e_detalhamento():
    """modo=auto é o caminho do PASSO 10 (Criação do Projeto) — gerar o documento é o que
    conclui o passo. Cobre os campos da tela 'Projeto — edição estruturada' que a geração
    ignorava (`empresas`, `conversoes`, `encarregado`), o Detalhamento vindo das respostas
    do Levantamento e a remoção das áreas não contratadas."""
    payload = {
        "slug": "projeto",
        "modo": "auto",
        "modeloBase64": _template_base64("projeto.docx"),
        "projeto": {
            "id": 1,
            "cliente": "Metalurgica Teste Ltda",
            "cnpj": "12.345.678/0001-90",
            "gci": "GCI Teste",
            "consultor": "Consultor Teste",
            "modulos": "FAT",
            "horasCobradas": "120",
            "horasBonificadas": "20",
        },
        "docConteudo": {
            "objetivos": "Padronizar o processo comercial no SIGER.",
            "empresas": "Matriz - 12.345.678/0001-90\nFilial - 12.345.678/0002-70",
            "conversoes": "Converter clientes, produtos e titulos em aberto.",
            "redator": "Redator Teste",
            "encarregado": "Fulano da Silva",
        },
        "indiceModulos": [{"sigla": "FAT", "nome": "Faturamento"}],
        "indiceTopicos": [{"moduloSigla": "FAT", "topico": "Emissao de pedido", "adicional": ""}],
        "levantamentoRespostas": [
            {"moduloSigla": "FAT", "topico": "Emissao de pedido",
             "resposta": "Pedido digitado pelo representante."}
        ],
        "cronogramaItens": [],
    }
    r = client.post("/gerar/documento-fiel", json=payload)
    assert r.status_code == 200

    doc = Document(io.BytesIO(r.content))
    texto = _texto_completo(doc)

    # Cabeçalho e objetivos (já cobertos no modo=modelo, mantidos como âncora do modo=auto)
    assert "Nome do Cliente: Metalurgica Teste Ltda" in texto
    assert "CNPJ: 12.345.678/0001-90" in texto
    assert "Padronizar o processo comercial no SIGER." in texto

    # Escopo — campos que a geração perdia: a tela gravava e o .docx saía em branco.
    assert "Matriz - 12.345.678/0001-90" in texto
    assert "Filial - 12.345.678/0002-70" in texto
    assert "Converter clientes, produtos e titulos em aberto." in texto

    # Equipes — as três linhas da Rech e a do cliente.
    assert "Gerente de Contas do Projeto: GCI Teste" in texto
    assert "Redator do Projeto: Redator Teste" in texto
    assert "Consultor/Implantador: Consultor Teste" in texto
    assert "Encarregado pelo Projeto: Fulano da Silva" in texto

    # Detalhamento das Rotinas: só a área contratada, preenchida pelo Levantamento.
    assert "Vendas e Faturamento" in texto
    assert "Emissao de pedido: Pedido digitado pelo representante." in texto
    assert "Controle de Compras" not in texto     # área sem módulo contratado -> removida
    assert "Gestão da Produção" not in texto      # grupo que ficou vazio -> removido

    # Tempo estimado e fecho.
    assert "120 horas cobradas" in texto
    assert "20 horas bonificadas" in texto

    # Nenhum marcador do layout pode vazar para o documento entregue ao cliente.
    assert "<" not in texto or ">" not in texto
    assert "(preencher)" not in texto

    # Erro de digitacao do layout oficial, corrigido em 2026-08-20: era "Da de Inicio".
    macro = [t for t in doc.tables if t.rows and (t.rows[0].cells[0].text or "").strip().lower() == "fase"]
    assert macro, "tabela Cronograma Macro nao encontrada no layout"
    etapas = [(row.cells[1].text or "").strip() for row in macro[0].rows[1:]]
    assert "Data de Início do Uso oficial" in etapas
    assert "Da de Início do Uso oficial" not in etapas


def test_gera_projeto_preenche_o_bloco_cadastros_e_nao_deixa_ponto_orfao():
    """Bloco "Cadastros" do layout: os três campos passaram a existir na tela do Projeto e o
    gerador os escreve NO LUGAR do marcador. Antes o bloco saía vazio em todo Projeto e o
    ponto final grudado no marcador ("<Detalhar ...>.") ficava órfão numa linha sozinha."""
    payload = {
        "slug": "projeto",
        "modo": "auto",
        "modeloBase64": _template_base64("projeto.docx"),
        "projeto": {"id": 1, "cliente": "Cliente Cadastros", "modulos": "FAT"},
        "docConteudo": {
            "cad_clientes": "Codigo sequencial; compartilhar cadastro entre as duas empresas.",
            "cad_produtos": "Codificacao por familia.\nUnidade de medida vem do fornecedor.",
        },
        "indiceModulos": [{"sigla": "FAT", "nome": "Faturamento"}],
        "indiceTopicos": [],
        "levantamentoRespostas": [],
        "cronogramaItens": [],
    }
    r = client.post("/gerar/documento-fiel", json=payload)
    assert r.status_code == 200
    doc = Document(io.BytesIO(r.content))
    textos = [p.text for p in doc.paragraphs]
    texto = "\n".join(textos)

    assert "Codigo sequencial; compartilhar cadastro entre as duas empresas." in texto
    # Textarea com várias linhas vira uma linha por parágrafo, não tudo emendado.
    assert "Codificacao por familia." in textos
    assert "Unidade de medida vem do fornecedor." in textos

    # `cad_outros` ficou em branco: o marcador some E o ponto órfão não fica para trás.
    assert "Outros pontos gerais do projeto" in textos
    assert not [t for t in textos if t.strip() in (".", ":", "-")]


def test_gera_projeto_nao_repete_o_valor_de_cadastros_no_campo_seguinte():
    """O marcador substituído é o do PRÓPRIO rótulo. O 'CNPJ:' logo abaixo de 'Estão
    contempladas ... as seguintes empresas:' é outro campo e não pode ser sobrescrito."""
    payload = {
        "slug": "projeto",
        "modo": "auto",
        "modeloBase64": _template_base64("projeto.docx"),
        "projeto": {"id": 1, "cliente": "Cliente Escopo", "cnpj": "11.111.111/0001-11",
                    "modulos": "FAT"},
        "docConteudo": {"empresas": "Matriz e filial."},
        "indiceModulos": [{"sigla": "FAT", "nome": "Faturamento"}],
        "indiceTopicos": [],
        "levantamentoRespostas": [],
        "cronogramaItens": [],
    }
    r = client.post("/gerar/documento-fiel", json=payload)
    assert r.status_code == 200
    textos = [p.text for p in Document(io.BytesIO(r.content)).paragraphs]
    assert "Matriz e filial." in textos
    assert "CNPJ: 11.111.111/0001-11" in textos


def test_cronograma_macro_do_projeto_sai_em_dd_mm_aaaa():
    """A tela grava a data em ISO (o <input type="date"> nao produz outra coisa), mas o
    documento que vai ao cliente e brasileiro. Valor antigo, digitado a mao em dd/mm/aaaa,
    tem de continuar saindo igual."""
    payload = {
        "slug": "projeto",
        "modo": "auto",
        "modeloBase64": _template_base64("projeto.docx"),
        "projeto": {"id": 1, "cliente": "Cliente Datas", "modulos": "FAT"},
        "docConteudo": {
            "crono_levantamento": "2026-08-10",      # ISO, vindo do seletor de data
            "crono_treinamento": "15/09/2026",       # legado, digitado a mao
            "crono_inicio": "a combinar",            # nao e data: sai como esta
        },
        "indiceModulos": [{"sigla": "FAT", "nome": "Faturamento"}],
        "indiceTopicos": [],
        "levantamentoRespostas": [],
        "cronogramaItens": [],
    }
    r = client.post("/gerar/documento-fiel", json=payload)
    assert r.status_code == 200

    doc = Document(io.BytesIO(r.content))
    macro = [t for t in doc.tables
             if t.rows and (t.rows[0].cells[0].text or "").strip().lower() == "fase"][0]
    por_etapa = {(row.cells[1].text or "").strip(): (row.cells[2].text or "").strip()
                 for row in macro.rows[1:]}

    assert por_etapa["Levantamento de requisitos"] == "10/08/2026"
    assert por_etapa["Treinamento"] == "15/09/2026"
    assert por_etapa["Data de Início do Uso oficial"] == "a combinar"
    # Nenhuma data pode vazar em ISO para o documento assinado pelo cliente.
    assert "2026-08-10" not in "\n".join(por_etapa.values())


def _realces(doc):
    """Quantos realces (caneta-marcador do Word) restam no documento."""
    from docx.oxml.ns import qn
    return sum(1 for _ in doc.element.body.iter(qn("w:highlight")))


def _sombreamentos(doc):
    """Sombreamentos com cor de verdade (a identidade visual da Rech vive aqui)."""
    from docx.oxml.ns import qn
    achados = []
    for shd in doc.element.body.iter(qn("w:shd")):
        fill = (shd.get(qn("w:fill")) or "").upper()
        if fill and fill not in ("AUTO", "FFFFFF"):
            achados.append(fill)
    return achados


def test_documento_gerado_nao_leva_realce_do_layout():
    """Os layouts marcam em verde/amarelo os pontos a preencher. O texto escrito pela geração
    HERDA esse realce, e o documento que vai ao cliente saía riscado — na tela e na impressão
    (pedido do usuário em 2026-08-25). O realce sai; o sombreamento fica."""
    base = Document(_LAYOUTS / "projeto.docx")
    assert _realces(base) > 0, "o layout de teste precisa ter realce para o caso provar algo"

    payload = {
        "slug": "projeto",
        "modo": "auto",
        "modeloBase64": _template_base64("projeto.docx"),
        "projeto": {"id": 1, "cliente": "Cliente Sem Realce", "modulos": "FAT"},
        "docConteudo": {"objetivos": "Texto que cai em cima de um marcador realcado."},
        "indiceModulos": [{"sigla": "FAT", "nome": "Faturamento"}],
        "indiceTopicos": [],
        "levantamentoRespostas": [],
        "cronogramaItens": [],
    }
    r = client.post("/gerar/documento-fiel", json=payload)
    assert r.status_code == 200

    gerado = Document(io.BytesIO(r.content))
    assert _realces(gerado) == 0
    assert "Texto que cai em cima de um marcador realcado." in _texto_completo(gerado)


def test_a_limpeza_de_realce_nao_leva_junto_a_identidade_visual():
    """O Levantamento usa sombreamento como IDENTIDADE: a faixa azul 0047BA do titulo e as
    tarjas CCDBF2 dos cabecalhos de tabela. Realce e anotacao e sai; sombreamento e layout e
    fica — apagar os dois deixaria o documento sem a marca da Rech."""
    base = Document(_LAYOUTS / "levantamento.docx")
    cores_base = set(_sombreamentos(base))
    assert cores_base, "o Levantamento precisa ter sombreamento colorido para o caso valer"

    payload = {
        "slug": "levantamento",
        "modo": "auto",
        "modeloBase64": _template_base64("levantamento.docx"),
        "projeto": {"id": 1, "cliente": "Cliente Identidade", "modulos": "FAT"},
        "docConteudo": {},
        "indiceModulos": [{"sigla": "FAT", "nome": "Faturamento"}],
        "indiceTopicos": [],
        "levantamentoRespostas": [],
        "cronogramaItens": [],
    }
    r = client.post("/gerar/documento-fiel", json=payload)
    assert r.status_code == 200

    gerado = Document(io.BytesIO(r.content))
    assert _realces(gerado) == 0
    assert cores_base.issubset(set(_sombreamentos(gerado)))


def test_gera_cronograma_preenche_cliente_consultor_horas_e_linhas():
    payload = {
        "slug": "cronograma",
        "modeloBase64": _template_base64("cronograma.xlsx"),
        "projeto": {
            "id": 1,
            "cliente": "Cliente Teste LTDA",
            "consultor": "Ana Consultora",
            "horasCobradas": "40",
            "horasBonificadas": "8",
        },
        "cronogramaItens": [
            {
                "etapa": "Abertura",
                "topicos": "Parametrização",
                "horas": "4",
                "data": "2026-08-10",
                "modalidade": "Remoto",
                "status": "Previsto",
            },
            {
                "etapa": "Treinamento — Faturamento",
                "topicos": "Cadastros e rotinas",
                "horas": "8",
                "data": "2026-08-17",
                "modalidade": "Presencial",
                "status": "Previsto",
            },
        ],
    }
    r = client.post("/gerar/documento-fiel", json=payload)
    assert r.status_code == 200
    assert r.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
    )

    wb = openpyxl.load_workbook(io.BytesIO(r.content))
    ws = wb["Cronograma de visitas"]
    assert ws["C6"].value == "Cliente Teste LTDA"
    assert ws["C3"].value == "Ana Consultora"
    assert ws["J3"].value == "40"
    assert ws["L4"].value == "8"
    assert ws["B9"].value == "2026-08-10"
    assert ws["N9"].value == "Abertura — Parametrização"
    assert ws["B10"].value == "2026-08-17"
    assert ws["N10"].value == "Treinamento — Faturamento — Cadastros e rotinas"


def test_gera_cronograma_sem_itens_ainda_preenche_cabecalho():
    payload = {
        "slug": "cronograma",
        "modeloBase64": _template_base64("cronograma.xlsx"),
        "projeto": {"id": 1, "cliente": "Cliente Sem Itens", "consultor": "Beto"},
    }
    r = client.post("/gerar/documento-fiel", json=payload)
    assert r.status_code == 200
    wb = openpyxl.load_workbook(io.BytesIO(r.content))
    ws = wb["Cronograma de visitas"]
    assert ws["C6"].value == "Cliente Sem Itens"
    assert ws["C3"].value == "Beto"
    assert ws["B9"].value is None


def test_slug_invalido_devolve_422():
    r = client.post(
        "/gerar/documento-fiel",
        json={"slug": "checklist", "modeloBase64": base64.b64encode(b"x").decode()},
    )
    assert r.status_code == 422


def test_modelo_base64_invalido_devolve_422():
    r = client.post(
        "/gerar/documento-fiel",
        json={"slug": "termo", "modeloBase64": "***nao-e-base64***", "projeto": {"id": 1, "cliente": "X"}},
    )
    assert r.status_code == 422
