# -*- coding: utf-8 -*-
"""Testes do endpoint /preview — pré-visualização WYSIWYG de documentos gerados/anexados.
Espelha webapp/routes_fluxo.py:projeto_doc_ver (docview.to_pdf/to_html)."""
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

import openpyxl
from docx import Document
from fastapi.testclient import TestClient

from main import app

client = TestClient(app)


def _docx_temp(tmp_path, nome="doc.docx"):
    doc = Document()
    doc.add_heading("Título de teste", level=0)
    doc.add_paragraph("Parágrafo com acentuação: ção, ã, é, ü.")
    tabela = doc.add_table(rows=2, cols=2)
    tabela.rows[0].cells[0].text = "Coluna A"
    tabela.rows[0].cells[1].text = "Coluna B"
    tabela.rows[1].cells[0].text = "1"
    tabela.rows[1].cells[1].text = "2"
    caminho = tmp_path / nome
    doc.save(caminho)
    return caminho


def _xlsx_temp(tmp_path, nome="planilha.xlsx"):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Aba1"
    ws.append(["Módulo", "Descrição"])
    ws.append(["FAT", "Vendas e Faturamento"])
    caminho = tmp_path / nome
    wb.save(caminho)
    return caminho


def test_arquivo_inexistente_404():
    r = client.post("/preview", json={"caminho": "C:\\caminho\\que\\nao\\existe.docx"})
    assert r.status_code == 404


def test_xlsx_sempre_vira_html(tmp_path):
    caminho = _xlsx_temp(tmp_path)
    r = client.post("/preview", json={"caminho": str(caminho)})
    assert r.status_code == 200
    corpo = r.json()
    assert corpo["tipo"] == "html"
    assert "Aba1" in corpo["html"]
    assert "FAT" in corpo["html"]
    assert "Vendas e Faturamento" in corpo["html"]


def test_docx_vira_pdf_ou_html(tmp_path):
    caminho = _docx_temp(tmp_path)
    r = client.post("/preview", json={"caminho": str(caminho)})
    assert r.status_code == 200
    corpo = r.json()
    assert corpo["tipo"] in ("pdf", "html")
    if corpo["tipo"] == "pdf":
        assert len(corpo["conteudoBase64"]) > 100
    else:
        assert "Título de teste" in corpo["html"]
        assert "Coluna A" in corpo["html"]


def test_tipo_nao_suportado_vira_html_com_aviso(tmp_path):
    caminho = tmp_path / "arquivo.txt"
    caminho.write_text("conteúdo qualquer")
    r = client.post("/preview", json={"caminho": str(caminho)})
    assert r.status_code == 200
    corpo = r.json()
    assert corpo["tipo"] == "html"
    assert "indisponível" in corpo["html"]
