# -*- coding: utf-8 -*-
"""
verificar.py — health-check / smoke test de TODA a estrutura.

Roda o ambiente, os templates, os 12 geradores (com os dados de exemplo), o
conversor verbal, o importador e a estrutura de agentes/skills, e imprime um
relatório PASS/FALHA/PULAR. Serve como teste de regressão.

Uso (a partir da pasta tools/):
    python verificar.py
Código de saída != 0 se houver qualquer FALHA.
"""
import os
import re
import io
import sys
import glob
import contextlib
import importlib

HERE = os.path.dirname(os.path.abspath(__file__))
REPO = os.path.dirname(HERE)
if not getattr(sys, "frozen", False):
    sys.path.insert(0, HERE)
import _common as C   # noqa: E402
OUT = C.OUT

R = []  # (area, nome, status, detalhe)


def reg(area, nome, status, detalhe=""):
    R.append((area, nome, status, detalhe))


# 1) Ambiente -----------------------------------------------------------------
def check_ambiente():
    reg("Ambiente", "Python " + sys.version.split()[0], "OK")
    for mod in ("docx", "openpyxl", "yaml"):
        try:
            m = importlib.import_module(mod)
            reg("Ambiente", mod, "OK", getattr(m, "__version__", ""))
        except Exception as e:
            reg("Ambiente", mod, "FALHA", str(e)[:60])


# 2) Templates ----------------------------------------------------------------
def check_templates():
    for t in ("base_projeto_tokenizado", "base_termo", "base_levantamento"):
        p = os.path.join(HERE, "templates", t + ".docx")
        existe = os.path.exists(p)
        reg("Templates", t + ".docx", "OK" if existe else "PULAR",
            "presente" if existe else "ausente (letterhead local; sem ele cai no estilo limpo)")


# 3) Geradores ----------------------------------------------------------------
GERADORES = [
    ("gerar_kit_mudanca", "Kit_Gestao_Mudanca_", "xlsx"),
    ("gerar_roteiros_teste", "Roteiros_SIT_UAT_", "xlsx"),
    ("gerar_aceite_uat", "Termo_Aceite_UAT_", "docx"),
    ("gerar_reconciliacao_conversao", "Reconciliacao_Conversao_", "xlsx"),
    ("gerar_painel_hypercare", "Painel_Hypercare_", "xlsx"),
    ("gerar_log_fitgap", "Log_FitGap_", "xlsx"),
    ("gerar_painel_kpi", "Painel_KPIs_", "xlsx"),
    ("gerar_raid", "RAID_", "xlsx"),
    ("gerar_dossie_cliente", "Dossie_", "docx"),
    ("gerar_projeto_implantacao", "Projeto_Implantacao_", "docx"),
    ("gerar_termo_encerramento", "Termo_Encerramento_", "docx"),
    ("gerar_levantamento", "Levantamento_", "docx"),
    ("gerar_checklist_consultor", "CheckList_Consultor_", "xlsx"),
]


def _validar(path, kind):
    if kind == "xlsx":
        from openpyxl import load_workbook
        return "(%d abas)" % len(load_workbook(path).sheetnames)
    from docx import Document
    d = Document(path)
    txt = "\n".join(p.text for p in d.paragraphs)
    txt += "\n".join(c.text for t in d.tables for r in t.rows for c in r.cells)
    toks = len(re.findall(r"\{\{[a-z0-9_]+\}\}", txt))
    extra = (", %d TOKENS sobrando!" % toks) if toks else ""
    return "(%d parágrafos%s)" % (len(d.paragraphs), extra)


def check_geradores():
    os.makedirs(OUT, exist_ok=True)
    for modname, prefix, kind in GERADORES:
        try:
            mod = importlib.import_module(modname)
            with contextlib.redirect_stdout(io.StringIO()):
                mod.main()
            outs = [f for f in os.listdir(OUT) if f.startswith(prefix)]
            if not outs:
                reg("Geradores", modname, "FALHA", "sem arquivo de saída")
                continue
            f = os.path.join(OUT, sorted(outs)[-1])
            det = _validar(f, kind)
            status = "FALHA" if "TOKENS sobrando" in det else "OK"
            reg("Geradores", modname, status, det)
        except SystemExit as e:
            reg("Geradores", modname, "PULAR", str(e).splitlines()[0][:55])
        except Exception as e:
            reg("Geradores", modname, "FALHA", "%s: %s" % (type(e).__name__, str(e)[:70]))


# 4) Conversor verbal ---------------------------------------------------------
def check_verbal():
    try:
        import conversor_verbal as V
    except Exception as e:
        reg("Verbal", "import", "FALHA", str(e)[:60]); return
    casos = [("A empresa utiliza o módulo.", "utilizará"),
             ("O processo é integrado.", "será"),
             ("Possui controle e realiza a baixa.", "Possuirá"),
             ("Os relatórios são gerados.", "serão")]
    for entrada, esperado in casos:
        saida = V.converter(entrada)
        reg("Verbal", entrada[:32], "OK" if esperado in saida else "FALHA", "-> " + saida[:38])
    s = V.converter("Usa a Estrutura via Formulação.")
    ok = "estruturará" not in s.lower() and "formulará" not in s.lower()
    reg("Verbal", "termo protegido (Formulação/Estrutura)", "OK" if ok else "FALHA", s[:40])


# 5) Importador ---------------------------------------------------------------
def check_importador():
    cands = glob.glob(os.path.expanduser("~/Downloads/Mapeamento levantamento*.docx"))
    cands = [c for c in cands if "XXXX" not in c]
    if not cands:
        reg("Importador", "levantamento de exemplo", "PULAR", "nenhum .docx em ~/Downloads")
        return
    try:
        import importar_mapeamento as I
        data = I.extract(cands[0])
        reg("Importador", "extract()", "OK",
            "cliente=%s | areas=%s" % (data.get("client_name", "?"),
                                       ",".join(data.get("areas_incluidas", [])) or "-"))
    except Exception as e:
        reg("Importador", "extract()", "FALHA", str(e)[:70])


# 6) Agentes e Skills ---------------------------------------------------------
def check_estrutura():
    import yaml
    ag = glob.glob(os.path.join(REPO, ".claude", "agents", "*.md"))
    sk = glob.glob(os.path.join(REPO, ".claude", "skills", "*", "SKILL.md"))

    def fm(path):
        txt = open(path, encoding="utf-8").read()
        m = re.match(r"^---\n(.*?)\n---", txt, re.S)
        try:
            return yaml.safe_load(m.group(1)) if m else None
        except Exception:
            return None

    ruins = [os.path.basename(os.path.dirname(p)) or os.path.basename(p)
             for p in ag + sk
             if not (fm(p) and fm(p).get("name") and fm(p).get("description"))]
    if not ag and not sk:
        reg("Estrutura", "agentes/skills", "PULAR", "não encontrados (modo .exe)")
    else:
        reg("Estrutura", "agentes", "OK" if ag else "PULAR", "%d encontrados" % len(ag))
        reg("Estrutura", "skills", "OK" if sk else "PULAR", "%d encontradas" % len(sk))
        reg("Estrutura", "frontmatter (name+description)", "OK" if not ruins else "FALHA",
            "todos válidos" if not ruins else ("inválidos: " + ", ".join(ruins[:5])))


# Relatório -------------------------------------------------------------------
def main():
    for fn in (check_ambiente, check_templates, check_geradores,
               check_verbal, check_importador, check_estrutura):
        try:
            fn()
        except Exception as e:
            reg("(geral)", fn.__name__, "FALHA", str(e)[:70])

    print("\n" + "=" * 78)
    print(" VERIFICAÇÃO DA ESTRUTURA DE IMPLANTAÇÃO")
    print("=" * 78)
    area_atual = None
    for area, nome, status, det in R:
        if area != area_atual:
            print("\n[%s]" % area)
            area_atual = area
        marca = {"OK": "[ OK ]", "FALHA": "[FALHA]", "PULAR": "[PULAR]"}.get(status, status)
        print("  %s %-42s %s" % (marca, nome[:42], det))

    ok = sum(1 for r in R if r[2] == "OK")
    falha = sum(1 for r in R if r[2] == "FALHA")
    pular = sum(1 for r in R if r[2] == "PULAR")
    print("\n" + "-" * 78)
    print(" RESUMO: %d OK | %d FALHA | %d PULAR  (total %d)" % (ok, falha, pular, len(R)))
    print("-" * 78)
    return 1 if falha else 0


if __name__ == "__main__":
    sys.exit(main())
