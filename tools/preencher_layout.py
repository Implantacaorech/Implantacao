# -*- coding: utf-8 -*-
"""Preenchimento FIEL de layouts .docx/.xlsx.

Recebe o arquivo do layout (igual ao anexo) e substitui SOMENTE os placeholders
conhecidos (<...>) pelos dados do projeto, preservando 100% da estrutura e
formatação. Placeholders sem dado no projeto (conteúdo manual do consultor)
permanecem como estão, servindo de guia de preenchimento.
"""
import re


# ----------------------------- DOCX -----------------------------
def _iter_paragraphs(doc):
    """Todos os parágrafos do corpo, de tabelas e de cabeçalhos/rodapés."""
    def _tab(t):
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
                for tt in cell.tables:
                    yield from _tab(tt)
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        yield from _tab(t)
    for sec in doc.sections:
        for hf in (sec.header, sec.footer):
            for p in hf.paragraphs:
                yield p
            for t in hf.tables:
                yield from _tab(t)


def _aplica_no_paragrafo(p, novo):
    """Escreve `novo` mantendo a formatação do 1º run e zerando os demais."""
    if p.runs:
        p.runs[0].text = novo
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(novo)


def substituir_docx(doc, repl):
    """repl: lista de (literal, valor). Substitui em runs mesclados por parágrafo."""
    n = 0
    for p in _iter_paragraphs(doc):
        full = "".join(r.text for r in p.runs)
        if not full:
            continue
        novo = full
        for old, new in repl:
            if old and old in novo:
                novo = novo.replace(old, str(new))
        if novo != full:
            _aplica_no_paragrafo(p, novo)
            n += 1
    return n


def set_paragrafo_por_prefixo(doc, prefixo, texto):
    """Substitui o texto inteiro do 1º parágrafo que começa com `prefixo`."""
    for p in _iter_paragraphs(doc):
        if p.text.strip().startswith(prefixo):
            _aplica_no_paragrafo(p, texto)
            return True
    return False


_MARCADOR = re.compile(r"<[^<>]{0,300}>")


def remover_marcadores_docx(doc):
    """Remove TODOS os marcadores <...> restantes (não preenchidos) do documento —
    corpo, tabelas e cabeçalhos/rodapés. Aplicar por último, após os preenchimentos."""
    n = 0
    for p in _iter_paragraphs(doc):
        full = "".join(r.text for r in p.runs)
        if "<" in full and ">" in full:
            novo = _MARCADOR.sub("", full)
            novo = re.sub(r"[ \t]{2,}", " ", novo).rstrip()
            if novo != full:
                _aplica_no_paragrafo(p, novo)
                n += 1
    return n


def preencher_docx(src_path, repl=None, paragrafos=None):
    """Abre o layout, aplica substituições e devolve o Document (não salva)."""
    from docx import Document
    doc = Document(src_path)
    if repl:
        substituir_docx(doc, repl)
    for prefixo, texto in (paragrafos or []):
        set_paragrafo_por_prefixo(doc, prefixo, texto)
    return doc


# ----------------------------- XLSX -----------------------------
def preencher_xlsx(src_path, repl=None, celulas=None):
    """Abre a planilha, substitui literais em células de texto e grava valores
    diretos em coordenadas (celulas={'Aba!B6': valor}). Devolve o Workbook."""
    import openpyxl
    wb = openpyxl.load_workbook(src_path)
    repl = repl or []
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for c in row:
                if isinstance(c.value, str) and c.value:
                    novo = c.value
                    for old, new in repl:
                        if old and old in novo:
                            novo = novo.replace(old, str(new))
                    if novo != c.value:
                        c.value = novo
    for chave, valor in (celulas or {}).items():
        aba, coord = chave.split("!", 1)
        if aba in wb.sheetnames:
            wb[aba][coord] = valor
    return wb
