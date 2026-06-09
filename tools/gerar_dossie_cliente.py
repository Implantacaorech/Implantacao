# -*- coding: utf-8 -*-
"""Gera o Dossiê do Cliente (.docx): documento vivo onde 'mora' o estado
consolidado de cada implantação — identificação, escopo, status por etapa,
RNS vinculadas, artefatos gerados e links.

Uso:
    python gerar_dossie_cliente.py [data/exemplo_cliente.yaml] [data/dossie.yaml]
"""
import os
import sys
from docx import Document
from docx.shared import RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import _common as C


def _shade_header(row, fill="1F4E78"):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), fill)
        tcPr.append(shd)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r.font.bold = True


def _table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].paragraphs[0].add_run(h)
    _shade_header(t.rows[0])
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
    return t


def main(cliente_path="data/exemplo_cliente.yaml", dossie_path="data/dossie.yaml"):
    cli = C.load_yaml(os.path.basename(cliente_path))
    dos = C.load_yaml(os.path.basename(dossie_path))
    cliente = cli.get("cliente", {})
    projeto = cli.get("projeto", {})

    doc = Document()
    h = doc.add_heading(f"Dossiê de Implantação — {cliente.get('nome','')}", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Atualizado em {C.today()}").italic = True

    doc.add_heading("Identificação", level=1)
    cnpjs = "; ".join(f"{c.get('sigla')} — {c.get('cnpj')}" for c in cliente.get("cnpjs", []))
    _table(doc, ["Campo", "Valor"], [
        ["Cliente", cliente.get("nome", "")],
        ["Código SICLA", cliente.get("codigo_sicla", "")],
        ["CNPJ(s) / Sigla(s)", cnpjs],
        ["RNS de Implantação", cliente.get("rns_implantacao", "")],
        ["Consultor responsável", projeto.get("consultor_responsavel", "")],
        ["Área", projeto.get("area", "")],
        ["Usuário líder", cliente.get("usuario_lider", "")],
        ["Contato", cliente.get("contato_fone", "")],
        ["Virada prevista", cliente.get("data_virada_prevista", "")],
    ])

    doc.add_heading("Escopo (módulos)", level=1)
    for m in projeto.get("modulos", []):
        doc.add_paragraph(m, style="List Bullet")

    doc.add_heading("Status por etapa", level=1)
    _table(doc, ["Etapa", "Status"],
           [[e.get("etapa", ""), e.get("status", "")] for e in dos.get("status_etapas", [])])

    doc.add_heading("RNS vinculadas", level=1)
    _table(doc, ["Tipo", "Número", "Descrição", "Status"],
           [[r.get("tipo",""), r.get("numero",""), r.get("descricao",""), r.get("status","")]
            for r in dos.get("rns_vinculadas", [])])

    doc.add_heading("Artefatos gerados", level=1)
    for a in dos.get("artefatos", []):
        doc.add_paragraph(a, style="List Bullet")

    doc.add_heading("Links", level=1)
    _table(doc, ["Recurso", "Link"],
           [[l.get("nome",""), l.get("url","")] for l in dos.get("links", [])])

    doc.add_heading("Observações", level=1)
    doc.add_paragraph("________________________________________________________________")
    doc.add_paragraph("________________________________________________________________")

    C.ensure_out()
    fname = f"Dossie_{C.slug(cliente.get('nome'))}.docx"
    path = os.path.join(C.OUT, fname)
    doc.save(path)
    print(f"OK: {fname} -> {path}")


if __name__ == "__main__":
    args = sys.argv[1:]
    main(*args) if args else main()
