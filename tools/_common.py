# -*- coding: utf-8 -*-
"""Helpers compartilhados pelos geradores Office (.xlsx/.docx)."""
import os
import sys
import re
import unicodedata
import datetime
import yaml
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

FROZEN = getattr(sys, "frozen", False)
if FROZEN:
    # Empacotado (.exe): os defaults vêm do bundle (somente leitura); as saídas e
    # os arquivos de cliente vão para %LOCALAPPDATA%\PainelImplantacao (gravável).
    HERE = os.path.join(sys._MEIPASS, "tools")                          # noqa
    _APP = os.path.join(os.environ.get("LOCALAPPDATA") or os.path.expanduser("~"),
                        "PainelImplantacao")
    DATA = os.path.join(HERE, "data")          # defaults (leitura)
    DATA_WRITE = os.path.join(_APP, "data")    # cliente/projeto (gravável)
    OUT = os.path.join(_APP, "exemplos")
else:
    HERE = os.path.dirname(os.path.abspath(__file__))
    REPO = os.path.dirname(HERE)
    DATA = os.path.join(HERE, "data")
    DATA_WRITE = DATA
    OUT = os.path.join(REPO, "exemplos")
TEMPLATES = os.path.join(HERE, "templates")
os.makedirs(DATA_WRITE, exist_ok=True)

# Estilos
HEADER_FILL = PatternFill("solid", fgColor="1F4E78")
HEADER_FONT = Font(bold=True, color="FFFFFF", size=11)
TITLE_FONT = Font(bold=True, size=16, color="1F4E78")
SUB_FONT = Font(italic=True, color="595959")
_THIN = Side(style="thin", color="BFBFBF")
BORDER = Border(left=_THIN, right=_THIN, top=_THIN, bottom=_THIN)
WRAP = Alignment(wrap_text=True, vertical="top")
CENTER = Alignment(horizontal="center", vertical="center", wrap_text=True)


def load_yaml(name):
    name = os.path.basename(name)
    for d in (DATA_WRITE, DATA):          # gravável primeiro, depois os defaults do bundle
        p = os.path.join(d, name)
        if os.path.exists(p):
            with open(p, encoding="utf-8") as fh:
                return yaml.safe_load(fh)
    with open(os.path.join(DATA, name), encoding="utf-8") as fh:
        return yaml.safe_load(fh)


def ensure_out():
    os.makedirs(OUT, exist_ok=True)
    return OUT


def path_dentro(path, *raizes):
    """True se `path` está DENTRO de uma das raízes (comparação por componente —
    startswith puro deixaria C:\\Dados casar com C:\\DadosXyz)."""
    try:
        p = os.path.normcase(os.path.abspath(path))
        for r in raizes:
            r = os.path.normcase(os.path.abspath(r))
            if p == r or p.startswith(r.rstrip("\\/") + os.sep):
                return True
    except (TypeError, ValueError):
        pass
    return False


def today():
    return datetime.date.today().strftime("%d/%m/%Y")


def slug(text):
    text = unicodedata.normalize("NFKD", str(text)).encode("ascii", "ignore").decode()
    text = re.sub(r"[^A-Za-z0-9]+", "_", text).strip("_")
    return text or "cliente"


def norm_doc(s):
    """Normaliza texto para comparar cabeçalhos (sem acento, minúsculo, sem pontas)."""
    s = unicodedata.normalize("NFKD", s or "")
    s = "".join(c for c in s if not unicodedata.combining(c))
    s = re.sub(r"\s+", " ", s.strip().lower())
    return s.strip(" .:-")


def safe_sheet(name):
    """Sanitiza nome de aba do Excel (proíbe \\ / ? * [ ] : e máx. 31 chars)."""
    name = re.sub(r"[\\/?*\[\]:]", "-", str(name))
    return name[:31]


def header_row(ws, headers, row=1):
    for col, text in enumerate(headers, start=1):
        c = ws.cell(row=row, column=col, value=text)
        c.fill = HEADER_FILL
        c.font = HEADER_FONT
        c.alignment = CENTER
        c.border = BORDER
    ws.row_dimensions[row].height = 28
    ws.freeze_panes = ws.cell(row=row + 1, column=1)


def set_widths(ws, widths):
    from openpyxl.utils import get_column_letter
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w


def write_rows(ws, rows, start=2):
    for r, row in enumerate(rows, start=start):
        for c, val in enumerate(row, start=1):
            cell = ws.cell(row=r, column=c, value=val)
            cell.alignment = WRAP
            cell.border = BORDER


def title_block(ws, title, subtitle=None, span=6):
    from openpyxl.utils import get_column_letter
    ws.merge_cells(f"A1:{get_column_letter(span)}1")
    t = ws.cell(row=1, column=1, value=title)
    t.font = TITLE_FONT
    if subtitle:
        ws.merge_cells(f"A2:{get_column_letter(span)}2")
        s = ws.cell(row=2, column=1, value=subtitle)
        s.font = SUB_FONT


# --- Word (.docx) helpers --- (TEMPLATES definido no topo, ciente de frozen)


def style_base(tipo):
    """Document com cabeçalho/rodapé/estilos do template real da Rech
    (tools/templates/base_<tipo>.docx), com o corpo limpo. Se o template não
    existir, retorna um Document novo. Retorna (doc, usou_base)."""
    from docx import Document
    from docx.oxml.ns import qn
    path = os.path.join(TEMPLATES, "base_%s.docx" % tipo)
    if os.path.exists(path):
        doc = Document(path)
        body = doc.element.body
        for child in list(body):
            if child.tag != qn("w:sectPr"):   # preserva seção (margens, header/footer)
                body.remove(child)
        used = True
    else:
        doc, used = Document(), False
    try:
        doc.styles["Normal"].paragraph_format.line_spacing = 1.15   # espaçamento padrão Rech
    except Exception:
        pass
    return doc, used


def docx_heading(doc, txt, size=13, center=False):
    """Título em negrito (estilo Rech: parágrafo manual, não 'Heading')."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt)
    r.bold = True
    r.font.size = Pt(size)
    return p


def docx_bullet(doc, text):
    """Marcador resiliente: usa 'List Bullet' quando disponível (doc novo); senão
    'List Paragraph' (template Rech) com glifo manual. Cria UM parágrafo só
    (define o estilo antes do texto para não deixar parágrafo órfão)."""
    p = doc.add_paragraph()
    try:
        p.style = "List Bullet"
        p.add_run(str(text))
    except KeyError:
        try:
            p.style = "List Paragraph"
        except KeyError:
            pass
        p.add_run("•  " + str(text))
    return p
