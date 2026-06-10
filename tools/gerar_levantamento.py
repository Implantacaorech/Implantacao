# -*- coding: utf-8 -*-
"""Gera o Levantamento (Mapeamento de Processos) (.docx) fiel ao template Rech.

Uso:
    python gerar_levantamento.py [data/levantamento.yaml]
"""
import os
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import _common as C
import catalogo as CAT

NAVY = RGBColor(0x1F, 0x4E, 0x78)


def shade_header(row, fill="1F4E78"):
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), fill)
        tcPr.append(shd)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r.font.bold = True


def main(lev_path="data/levantamento.yaml"):
    d = C.load_yaml(os.path.basename(lev_path))
    # Automação: resolve os módulos contratados (códigos/abreviações) no catálogo
    # e agrupa por área para preencher o Resumo e os "Módulos Previstos".
    contratados, _faltam = CAT.resolve(d.get("modulos_contratados"))
    areas_auto = CAT.por_area(contratados) if contratados else None
    doc, based = C.style_base("levantamento")
    if not based:
        doc.styles["Normal"].font.name = "Calibri"
        doc.styles["Normal"].font.size = Pt(11)

    _HS = {1: 13, 2: 12, 3: 11}

    def H(txt, level=1):
        return C.docx_heading(doc, txt, size=_HS.get(level, 12))

    def P(txt="", bold=False):
        p = doc.add_paragraph(); p.add_run(txt).bold = bold
        return p

    def B(items):
        for it in items or []:
            C.docx_bullet(doc, it)

    def table(headers, rows, empty=0):
        t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
        for i, h in enumerate(headers):
            t.rows[0].cells[i].paragraphs[0].add_run(h)
        shade_header(t.rows[0])
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row): cells[i].text = str(v) if v is not None else ""
        for _ in range(empty):
            t.add_row()
        return t

    # Título
    C.docx_heading(doc, "Mapeamento de Processos", size=16, center=True)
    P("").add_run(d.get("cliente", "")).bold = True
    P(f"Data: {d.get('data','')}")
    P(f"Responsáveis: {d.get('responsaveis','')}")

    H("Revisões", 2)
    table(["Data", "Revisão", "Redator", "Motivo da Alteração"], [], empty=5)

    # Identificação
    H("Informações da empresa", 1)
    H("Identificação da Empresa", 2)
    idf = d.get("identificacao", {})
    for label, key in [("Razão Social", "razao_social"), ("Ramo Atividade", "ramo"),
                       ("Produto", "produto"), ("Fornecedor Atual Software", "fornecedor_atual"),
                       ("Localização / Filiais", "localizacao"),
                       ("Observações / Objetivos", "observacoes_objetivos")]:
        P(f"{label}: {idf.get(key,'')}")

    P("Quantidade usuários e identificação:", bold=True)
    table(["Nome", "E-mail", "Atribuições"],
          [[u.get("nome",""), u.get("email",""), u.get("atribuicoes","")] for u in d.get("usuarios", [])],
          empty=2)

    # Resumo dos Módulos e Adicionais Contratados (A)
    P("Resumo dos Módulos e Adicionais Contratados", bold=True)
    if contratados:
        rows_a = [[f"{m['abrev']} — {m['descricao']}", "X", "", ""] for m in contratados]
    else:
        rows_a = [[m.get("modulo", ""), m.get("necessidade", ""), "", m.get("obs", "")]
                  for m in d.get("modulos_previstos_antes", [])]
    table(["Módulos/Adicionais (A) — Previstos antes do Levantamento", "Sim", "Não", "Observações"], rows_a)
    table(["Módulos/Adicionais (B) — Identificados no Levantamento", "Sim", "Não", "Observações"],
          [[m.get("modulo", ""), m.get("necessidade", ""), "", m.get("obs", "")]
           for m in d.get("modulos_identificados", [])])

    # Horas
    H("Implantação/Treinamento", 2)
    h = d.get("horas", {})
    table(["Quantidade de horas Cobradas", "Quantidade de horas Bonificadas", "Total de Horas previstas"],
          [[h.get("cobradas",""), h.get("bonificadas",""), h.get("total","")]])

    # Conversões — bloco padrão fiel ao template (estimativas vêm do YAML/formulário)
    conv = d.get("conversoes", {})
    e = conv.get("estimativas", {}) if isinstance(conv.get("estimativas"), dict) else {}
    H(f"CONVERSÕES ({conv.get('horas','')} horas)", 2)
    P("Detalhamento e considerações levantadas:", bold=True)
    P("Considerações Gerais -> Reforçar que a conversão depende do acesso aos dados ou da "
      "exportação das informações necessárias para que se torne viável (exceto histórico de "
      "venda que pode ser feito por importação de XML).")
    P(f"Imp. Cad. clientes e fornecedores – Estimativa: {e.get('clientes_fornecedores','')}")
    P(f"Imp. Cad. produtos – Estimativa: {e.get('produtos','')}")
    P(f"Imp. Mov. Financeiro doc. em aberto – Estimativa: {e.get('financeiro','')}")
    P("Validar aspectos como:", bold=True)
    B(["Numero Bancário",
       "Comissão (se mais de um representante por documento)",
       "Conta de Planejamento Financeiro (se mais de uma conta por documento – idem com Centro de custo do planejamento)",
       "Conta Contábil (se mais de uma conta por documento – idem com Centro de custo da contabilidade)"])
    P(f"Imp. Notas Fiscais já emitidas – Estimativa: {e.get('notas_fiscais','')}")
    B(["Validar período desejado (impacto no tempo e na necessidade de ter os arquivos)",
       "Validar aspectos de mudança de códigos – necessidade de montar equivalência"])
    P("Importação de Histórico de Compras, por nota de entradas", bold=True)
    P("Não convertemos. (no máximo que temos é poder importar histórico de Ordens de Compra, "
      "mas não de notas de entrada em função das equivalências).")
    P(f"Importação de movimentos da Folha de Pagamento: {e.get('folha','')}")

    H("Desenvolvimentos Específicos", 2)
    P(d.get("desenvolvimentos", "A definir"))

    # Mapeamento por área — AUTOMÁTICO a partir dos módulos contratados (catálogo);
    # ou manual (campo 'areas') quando não houver 'modulos_contratados'.
    nao_processo = {"BI e Integrações", "Outros"}
    base_cadastros = ("Cliente/Fornecedor", "Produto")   # sempre presentes, sem "Módulos Previstos"
    try:
        perg_areas = (C.load_yaml("perguntas_levantamento.yaml") or {}).get("areas", {}) or {}
    except Exception:
        perg_areas = {}

    def map_area(area, mods=None):
        H(f"Mapeamento de processo – {area.upper()}", 2)
        if mods is not None:
            P("Módulos Previstos:", bold=True)
            B([m["descricao"] for m in mods])
        for asp in (perg_areas.get(area) or [{"subtitulo": ""}]):
            sub = asp.get("subtitulo", "")
            P("Aspectos identificados" + (f" – {sub}" if sub else ""), bold=True)
            perguntas = asp.get("perguntas") or []
            if perguntas:
                B(perguntas)
            else:
                P("<Colar aqui o quadro com as perguntas para ir preenchendo as respostas>")
            if asp.get("nota"):
                P(asp["nota"])
        P("Dúvidas e Observações", bold=True)

    if areas_auto:
        for area in base_cadastros:           # Cliente/Fornecedor e Produto (cadastros base)
            map_area(area)
        for area, mods in areas_auto:
            if area in nao_processo or area in base_cadastros:
                continue
            map_area(area, mods)
    else:
        for a in d.get("areas", []):
            H(f"Mapeamento de processo – {a.get('nome','')}", 2)
            P("Módulos Previstos:", bold=True)
            P(a.get("modulos_previstos", ""))
            P("Aspectos identificados", bold=True)
            B(a.get("aspectos"))
            P("Dúvidas e Observações", bold=True)
            B(a.get("duvidas"))

    # (O Roteiro/Check List por módulo é gerado em planilha separada —
    #  gerar_checklist_consultor.py — como guia do Consultor.)
    C.ensure_out()
    fname = f"Levantamento_{C.slug(d.get('cliente'))}.docx"
    path = os.path.join(C.OUT, fname)
    doc.save(path)
    print(f"OK: {fname} ({len(d.get('areas', []))} áreas) -> {path}")


if __name__ == "__main__":
    args = sys.argv[1:]
    main(*args) if args else main()
