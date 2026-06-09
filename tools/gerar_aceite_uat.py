# -*- coding: utf-8 -*-
"""Gera o Termo de Aceite de Testes (SIT/UAT) em .docx — documento de sign-off
que serve de gate para autorizar a virada oficial.

Uso:
    python gerar_aceite_uat.py [data/exemplo_cliente.yaml] [data/roteiros_teste.yaml]
"""
import os
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import _common as C


def _shade_header(row):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:fill"), "1F4E78")
        tcPr.append(shd)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.bold = True


def main(cliente_path="data/exemplo_cliente.yaml", roteiros_path="data/roteiros_teste.yaml"):
    cli = C.load_yaml(os.path.basename(cliente_path))
    rot = C.load_yaml(os.path.basename(roteiros_path))
    cliente = cli.get("cliente", {})
    projeto = cli.get("projeto", {})
    modulos = rot.get("modulos", [])

    doc = Document()

    h = doc.add_heading("TERMO DE ACEITE DE TESTES (SIT / UAT)", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(f"{cliente.get('nome','')} — Projeto nº {projeto.get('numero','')}")
    r.italic = True

    doc.add_paragraph()
    doc.add_paragraph(
        "Este termo registra a conclusão e o aceite dos testes do sistema SIGER® para o cliente "
        f"{cliente.get('nome','')}, abrangendo o Teste Integrado (SIT) e o Aceite do Usuário (UAT) "
        "dos módulos no escopo. O aceite formaliza que os processos foram validados e autoriza a "
        "preparação da virada oficial."
    )

    # Dados
    doc.add_heading("Identificação", level=1)
    t = doc.add_table(rows=0, cols=2); t.style = "Table Grid"
    for k, v in [
        ("Cliente", cliente.get("nome", "")),
        ("Código SICLA", cliente.get("codigo_sicla", "")),
        ("RNS de Implantação", cliente.get("rns_implantacao", "")),
        ("Consultor responsável", projeto.get("consultor_responsavel", "")),
        ("Virada prevista", cliente.get("data_virada_prevista", "")),
        ("Data do aceite", C.today()),
    ]:
        cells = t.add_row().cells
        cells[0].paragraphs[0].add_run(k).bold = True
        cells[1].text = str(v)

    # Resumo por módulo
    doc.add_heading("Resultado por módulo", level=1)
    tm = doc.add_table(rows=1, cols=5); tm.style = "Table Grid"
    hdr = tm.rows[0].cells
    for i, txt in enumerate(["Módulo", "Casos", "Aprovados", "Reprovados", "Pendentes"]):
        hdr[i].paragraphs[0].add_run(txt)
    _shade_header(tm.rows[0])
    for mod in modulos:
        n = len(mod.get("casos", []))
        cells = tm.add_row().cells
        cells[0].text = mod.get("nome", "")
        cells[1].text = str(n)
        cells[2].text = ""; cells[3].text = ""; cells[4].text = ""
    doc.add_paragraph(
        "Preencher Aprovados/Reprovados/Pendentes a partir da planilha de roteiros (aba “Resumo e Sign-off”)."
    ).italic = True

    # Critério de liberação
    doc.add_heading("Critério de liberação (gate da virada)", level=1)
    for item in [
        "≥ 95% dos casos UAT com status Aprovado.",
        "Zero defeitos de severidade Crítica em aberto.",
        "Defeitos de severidade Alta com plano de ação acordado.",
        "Remessas bancárias e integrações com terceiros homologadas (quando aplicável).",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    # Assinaturas
    doc.add_heading("Assinaturas", level=1)
    doc.add_paragraph()
    for papel in ["Consultor de Implantação — Rech", "Usuário Líder — " + cliente.get("nome", ""),
                  "Gerente do Projeto — Rech"]:
        doc.add_paragraph("__________________________________________")
        p = doc.add_paragraph(papel); p.runs[0].italic = True
        doc.add_paragraph()

    C.ensure_out()
    fname = f"Termo_Aceite_UAT_{C.slug(cliente.get('nome'))}.docx"
    path = os.path.join(C.OUT, fname)
    doc.save(path)
    print(f"OK: {fname} -> {path}")


if __name__ == "__main__":
    args = sys.argv[1:]
    main(*args) if args else main()
