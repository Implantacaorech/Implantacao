# -*- coding: utf-8 -*-
"""
Gera o Projeto de Implantação do SIGER® (.docx) — engine de TOKENS, portada do
gerador interno da Rech (GeradorProjetoSIGER/docgen.py). Usa o template
tokenizado `tools/templates/base_projeto_tokenizado.docx` (com `{{tokens}}`):

  1) remove as áreas (Detalhamento das Rotinas) não incluídas;
  2) preenche os tokens (campos "bloco" -> cada linha vira um bullet);
  3) reconstrói a Tabela de Usuários;
  4) preenche a Equipe (opcional);
  5) limpa os marcadores do modelo (vermelho/realce verde);
  6) corrige o typo "Da de Início" -> "Data de Início";
  7) força a grade completa das tabelas.

Dados em `tools/data/projeto.yaml` (nomes = tokens). Uso:
    python gerar_projeto_implantacao.py [data/projeto.yaml]
"""
import os
import re
import sys
import copy

import docx
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import _common as C
import schema_projeto as S

TOKEN_RE = re.compile(r"\{\{[a-z0-9_]+\}\}")
TEMPLATE = os.path.join(C.TEMPLATES, "base_projeto_tokenizado.docx")
TYPO_FIXES = {"Da de Início do Uso oficial": "Data de Início do Uso oficial"}


# --- helpers de runs/parágrafos (port docgen.py) ---------------------------
def iter_paragraphs_in_order(doc):
    for child in doc.element.body.iterchildren():
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            yield Paragraph(child, doc)
        elif tag == 'tbl':
            for row in Table(child, doc).rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p


def strip_color(run):
    rpr = run._r.find(qn('w:rPr'))
    if rpr is None:
        return
    for tag in ('w:color', 'w:highlight', 'w:shd'):
        for el in rpr.findall(qn(tag)):
            rpr.remove(el)


def _clean_rpr(rpr):
    if rpr is None:
        return
    for h in rpr.findall(qn('w:highlight')):
        rpr.remove(h)
    for c in rpr.findall(qn('w:color')):
        if (c.get(qn('w:val')) or '').lower() not in ('000000', 'auto'):
            rpr.remove(c)


def _clean_paragraph_markers(p):
    ppr = p._p.find(qn('w:pPr'))
    if ppr is not None:
        _clean_rpr(ppr.find(qn('w:rPr')))
    for r in p.runs:
        _clean_rpr(r._r.find(qn('w:rPr')))


def clean_markers(doc):
    for p in iter_paragraphs_in_order(doc):
        _clean_paragraph_markers(p)
    for sec in doc.sections:
        for hf in (sec.header, sec.footer, sec.first_page_header,
                   sec.first_page_footer, sec.even_page_header, sec.even_page_footer):
            for p in hf.paragraphs:
                _clean_paragraph_markers(p)
            for tbl in hf.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            _clean_paragraph_markers(p)


def ensure_table_borders(doc):
    for table in doc.tables:
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)
        for old in tblPr.findall(qn('w:tblBorders')):
            tblPr.remove(old)
        borders = OxmlElement('w:tblBorders')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            e = OxmlElement('w:' + edge)
            e.set(qn('w:val'), 'single'); e.set(qn('w:sz'), '4')
            e.set(qn('w:space'), '0'); e.set(qn('w:color'), 'auto')
            borders.append(e)
        tblPr.append(borders)
        for row in table.rows:
            for cell in row.cells:
                tcPr = cell._tc.get_or_add_tcPr()
                for old in tcPr.findall(qn('w:tcBorders')):
                    tcPr.remove(old)


def fix_known_typos(doc):
    for p in iter_paragraphs_in_order(doc):
        for wrong, right in TYPO_FIXES.items():
            if wrong in p.text:
                idx = p.text.find(wrong)
                replace_span_in_runs(p.runs, idx, idx + len(wrong), right)


def replace_span_in_runs(runs, s, e, nt, clear_color=False):
    pos, placed, placed_run = 0, False, None
    for r in runs:
        txt = r.text or ''
        rs, re_ = pos, pos + len(txt)
        pos = re_
        if re_ <= s or rs >= e:
            continue
        left = txt[: s - rs] if rs < s else ''
        right = txt[e - rs:] if re_ > e else ''
        if not placed:
            r.text = left + nt + right
            placed, placed_run = True, r
        else:
            r.text = right
    if not placed and runs:
        runs[-1].text = (runs[-1].text or '') + nt
        placed_run = runs[-1]
    if clear_color and placed_run is not None:
        strip_color(placed_run)
    return placed_run


def set_paragraph_text(p, text):
    runs = p.runs
    if runs:
        runs[0].text = text
        for r in runs[1:]:
            r._r.getparent().remove(r._r)
    else:
        p.add_run(text)
    if p.runs:
        strip_color(p.runs[0])


def set_cell_text(cell, text):
    paras = cell.paragraphs
    set_paragraph_text(paras[0], text)
    for extra in paras[1:]:
        extra._p.getparent().remove(extra._p)


# --- 1) remoção de áreas não incluídas -------------------------------------
def remove_unselected_areas(doc, included_ids):
    body = doc.element.body
    children = list(body.iterchildren())

    def el_text(el):
        return Paragraph(el, doc).text if el.tag.split('}')[-1] == 'p' else None

    norm_children = [(el, C.norm_doc(el_text(el)) if el_text(el) is not None else None)
                     for el in children]
    boundaries = set(C.norm_doc(x) for x in (S.GRUPOS + S.SUBAREAS + [S.SECTION_APOS_ROTINAS]))

    def find_index(target):
        for i in range(len(norm_children)):
            if norm_children[i][1] == target:
                return i
        return -1

    to_remove = []
    for area in S.AREAS:
        if area["id"] in included_ids:
            continue
        sidx = find_index(C.norm_doc(area["subarea"]))
        if sidx == -1:
            continue
        eidx = len(norm_children)
        for j in range(sidx + 1, len(norm_children)):
            if norm_children[j][1] in boundaries:
                eidx = j
                break
        for k in range(sidx, eidx):
            to_remove.append(norm_children[k][0])
    for grupo in S.GRUPOS:
        subs = [a for a in S.AREAS if a["grupo"] == grupo]
        if subs and all(a["id"] not in included_ids for a in subs):
            gidx = find_index(C.norm_doc(grupo))
            if gidx != -1:
                to_remove.append(norm_children[gidx][0])
    seen = set()
    for el in to_remove:
        if id(el) in seen:
            continue
        seen.add(id(el))
        if el.getparent() is not None:
            el.getparent().remove(el)


# --- 2) tokens -------------------------------------------------------------
def fill_block_token(doc, token_name, value):
    token = "{{%s}}" % token_name
    target = None
    for p in iter_paragraphs_in_order(doc):
        if token in p.text:
            target = p
            break
    if target is None:
        return
    lines = [ln.strip() for ln in (value or "").splitlines() if ln.strip()]
    if not lines:
        target._p.getparent().remove(target._p)
        return
    set_paragraph_text(target, lines[0])
    anchor = target._p
    for ln in lines[1:]:
        new_el = copy.deepcopy(target._p)
        anchor.addnext(new_el)
        anchor = new_el
        set_paragraph_text(Paragraph(new_el, target._parent), ln)


def fill_inline_token(doc, token_name, value):
    token = "{{%s}}" % token_name
    lines = [ln.rstrip() for ln in (value or "").split("\n")]
    while lines and not lines[0].strip():
        lines.pop(0)
    while lines and not lines[-1].strip():
        lines.pop()
    first = lines[0] if lines else ""
    rest = lines[1:]
    for p in iter_paragraphs_in_order(doc):
        guard = 0
        while token in p.text:
            idx = p.text.find(token)
            run = replace_span_in_runs(p.runs, idx, idx + len(token), first, clear_color=True)
            if run is not None and rest:
                for ln in rest:
                    run.add_break(); run.add_text(ln)
            guard += 1
            if guard > 50:
                break


def fill_tokens(doc, data):
    for token_name in list(S.BLOCK_TOKENS):
        fill_block_token(doc, token_name, data.get(token_name, ""))
    inline_names = set()
    for p in iter_paragraphs_in_order(doc):
        for m in TOKEN_RE.findall(p.text):
            inline_names.add(m[2:-2])
    for token_name in inline_names:
        fill_inline_token(doc, token_name, data.get(token_name, ""))


# --- 3) tabela de usuários / 4) equipe -------------------------------------
def find_users_table(doc):
    for tbl in doc.tables:
        if tbl.rows:
            header = " ".join(C.norm_doc(c.text) for c in tbl.rows[0].cells)
            if "assina protocolo" in header and "nome" in header:
                return tbl
    return None


def rebuild_users_table(doc, usuarios):
    tbl = find_users_table(doc)
    if tbl is None or len(tbl.rows) < 2:
        return
    template_tr = copy.deepcopy(tbl.rows[1]._tr)
    for r in tbl.rows[1:]:
        r._tr.getparent().remove(r._tr)
    lista = usuarios or [{"nome": "", "email": "", "area": "", "assina": ""}]
    for u in lista:
        new_tr = copy.deepcopy(template_tr)
        tbl._tbl.append(new_tr)
        valores = [u.get("nome", ""), u.get("email", ""), u.get("area", ""), u.get("assina", "")]
        for ci, cell in enumerate(tbl.rows[-1].cells):
            set_cell_text(cell, valores[ci] if ci < len(valores) else "")


def fill_equipe(doc, equipe):
    if not equipe:
        return
    targets = {C.norm_doc(label): key for key, label, _ in S.EQUIPE_FIELDS}
    for p in iter_paragraphs_in_order(doc):
        key = targets.get(C.norm_doc(p.text))
        if key and equipe.get(key):
            strip_color(p.add_run(" " + str(equipe[key]).strip()))


# --- dados a partir do YAML ------------------------------------------------
def build_data(y):
    data = {}
    special = {"usuarios", "equipe", "areas_incluidas"}
    for k, v in y.items():
        if k in special:
            continue
        data[k] = "\n".join(str(x) for x in v) if isinstance(v, list) else ("" if v is None else str(v))
    data["usuarios"] = y.get("usuarios", []) or []
    data["equipe"] = y.get("equipe", {}) or {}
    inc = y.get("areas_incluidas")
    if inc is None:
        inc = [a["id"] for a in S.AREAS
               if any(data.get(f"{a['id']}_{sf}", "").strip() for sf in a["subfields"])]
    data["areas_incluidas"] = inc
    return data


def generate(data, out_docx):
    if not os.path.exists(TEMPLATE):
        raise SystemExit(
            "Template tokenizado ausente: tools/templates/base_projeto_tokenizado.docx\n"
            "Copie o assets/template.docx do GeradorProjetoSIGER para esse caminho "
            "(ver tools/templates/README.md).")
    doc = docx.Document(TEMPLATE)
    included = set(data.get("areas_incluidas", [a["id"] for a in S.AREAS]))
    remove_unselected_areas(doc, included)
    fill_tokens(doc, data)
    rebuild_users_table(doc, data.get("usuarios", []))
    fill_equipe(doc, data.get("equipe", {}))
    clean_markers(doc)
    fix_known_typos(doc)
    ensure_table_borders(doc)
    os.makedirs(os.path.dirname(os.path.abspath(out_docx)), exist_ok=True)
    doc.save(out_docx)
    return out_docx


def main(yaml_path="data/projeto.yaml"):
    y = C.load_yaml(os.path.basename(yaml_path))
    data = build_data(y)
    C.ensure_out()
    nome = y.get("client_name") or "cliente"
    out = os.path.join(C.OUT, f"Projeto_Implantacao_{C.slug(nome)}.docx")
    generate(data, out)
    print(f"OK: {os.path.basename(out)} (áreas: {', '.join(data['areas_incluidas']) or '—'}) -> {out}")


if __name__ == "__main__":
    args = sys.argv[1:]
    main(*args) if args else main()
