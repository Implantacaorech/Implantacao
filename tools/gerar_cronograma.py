# -*- coding: utf-8 -*-
"""Gera o Cronograma de Implantação do SIGER® (.docx) — documento obrigatório.

Distribui as agendas conforme as horas (cobradas + bonificadas), com os macro
tópicos por etapa/treinamento, datas previstas (em dias úteis) e modalidade.
Pode receber agendas explícitas no YAML; caso contrário, monta o plano
automaticamente a partir dos módulos contratados.

Uso:
    python gerar_cronograma.py [data/cronograma.yaml]
"""
import os
import re
import sys
import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import _common as C

NAVY = RGBColor(0x1F, 0x4E, 0x78)
_MESES = ["janeiro", "fevereiro", "março", "abril", "maio", "junho", "julho",
          "agosto", "setembro", "outubro", "novembro", "dezembro"]


def shade_header(row, fill="1F4E78"):
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), fill)
        tcPr.append(shd)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r.font.bold = True


def _num(s):
    m = re.search(r"\d+(?:[.,]\d+)?", str(s or ""))
    return float(m.group(0).replace(",", ".")) if m else 0.0


def _fmt(x):
    x = float(x or 0)
    return str(int(x)) if x == int(x) else ("%.1f" % x)


def _parse_date(s):
    s = str(s or "").strip()
    for fmt in ("%Y-%m-%d", "%d/%m/%Y", "%d/%m/%y"):
        try:
            return datetime.datetime.strptime(s, fmt).date()
        except ValueError:
            pass
    return datetime.date.today()


def _prox_util(d):
    while d.weekday() >= 5:           # 5=sáb, 6=dom
        d += datetime.timedelta(days=1)
    return d


def _add_uteis(d, n):
    while n > 0:
        d += datetime.timedelta(days=1)
        if d.weekday() < 5:
            n -= 1
    return d


def _distribuir(total, pesos):
    """Reparte `total` horas pelos pesos (método do maior resto), em inteiros.
    Sem total informado, usa peso*2 como horas."""
    if total <= 0:
        return [max(1, int(round(p * 2))) for p in pesos]
    soma = sum(pesos) or 1
    exatos = [total * p / soma for p in pesos]
    base = [int(x) for x in exatos]
    resto = int(round(total)) - sum(base)
    ordem = sorted(range(len(pesos)), key=lambda i: exatos[i] - base[i], reverse=True)
    i = 0
    while resto > 0 and ordem:
        base[ordem[i % len(ordem)]] += 1
        resto -= 1; i += 1
    return base


def _plano_automatico(modulos):
    """Plano padrão da implantação SIGER®: lista de (etapa, macro tópicos, peso)."""
    import catalogo as CAT
    achados, _ = CAT.resolve(modulos or [])
    plano = [
        ("Abertura + Parametrização inicial",
         "Criação de empresas e siglas; parâmetros gerais (1.1.P) e por empresa (1.2.A); "
         "compartilhamento de cadastros (1.2.M); tabelas genéricas.", 2.0),
    ]
    for m in achados:
        desc = m.get("descricao") or m.get("abrev") or "Módulo"
        plano.append(("Treinamento — %s" % desc,
                      "Tabelas e cadastros do módulo; importações via layout do SIGER®; "
                      "rotinas do processo; relatórios.", 2.0))
    if not achados:
        plano.append(("Treinamento das rotinas",
                      "Tabelas e cadastros; importações via layout; rotinas dos processos.", 4.0))
    plano += [
        ("Simulação de microprocessos",
         "Teste das rotinas treinadas por processo; ajustes finos.", 1.5),
        ("Simulação do macroprocesso",
         "Ensaio do processo completo (cenário real); validação ponta a ponta.", 1.5),
        ("Conversão — prévia",
         "Carga de teste; reconciliação origem × destino; validação de amostras.", 1.0),
        ("Conversão — oficial / ponto de corte",
         "Conversão oficial; conferência de saldos; definição do ponto de corte.", 1.0),
        ("Virada oficial (go-live)",
         "Início do uso em produção; acompanhamento full time; primeiros lançamentos.", 1.5),
        ("Acompanhamento / Hypercare",
         "Estabilização; micro ajustes; primeiros fechamentos.", 1.0),
        ("Encerramento",
         "Revisão de pendências; Termo de Encerramento; transição ao Suporte.", 0.5),
    ]
    return plano


def main(crono_path="data/cronograma.yaml"):
    d = C.load_yaml(os.path.basename(crono_path))
    cliente = d.get("cliente", "")
    num = d.get("numero_projeto", "")
    consultor = d.get("consultor", "")
    modalidade = d.get("modalidade_padrao", "A combinar")
    cadencia = int(_num(d.get("cadencia_dias", 5)) or 5)     # dias úteis entre agendas

    horas = d.get("horas") or {}
    cob = _num(horas.get("cobradas"))
    bon = _num(horas.get("bonificadas"))
    total = cob + bon or _num(horas.get("total") or d.get("horas_total"))

    # Agendas explícitas têm prioridade; senão, plano automático pelos módulos.
    agendas = d.get("agendas")
    if not agendas:
        plano = _plano_automatico(d.get("modulos_contratados") or d.get("modulos") or [])
        hs = _distribuir(total, [p for _, _, p in plano])
        if total <= 0:
            total = sum(hs)
        agendas = [{"etapa": e, "topicos": t, "horas": h}
                   for (e, t, _), h in zip(plano, hs)]

    # Datas previstas (dias úteis), espaçadas pela cadência.
    dt0 = _prox_util(_parse_date(d.get("data_inicio") or d.get("data")))
    for i, ag in enumerate(agendas):
        if not ag.get("data"):
            ag["data"] = (dt0 if i == 0 else _add_uteis(dt0, cadencia * i)).strftime("%d/%m/%Y")
        ag.setdefault("modalidade", modalidade)
        ag.setdefault("horas", "")
    soma_horas = sum(_num(a.get("horas")) for a in agendas)

    # --- Documento ---
    doc, based = C.style_base("cronograma")
    if not based:
        doc.styles["Normal"].font.name = "Calibri"
        doc.styles["Normal"].font.size = Pt(11)

    def P(txt="", bold=False):
        p = doc.add_paragraph(); p.add_run(txt).bold = bold
        return p

    C.docx_heading(doc, "Cronograma de Implantação do SIGER®", size=16, center=True)
    P("").add_run("Cliente: %s" % cliente).bold = True
    linha = []
    if num:
        linha.append("Projeto nº %s" % num)
    if consultor:
        linha.append("Consultor: %s" % consultor)
    linha.append("Horas: %s contratadas + %s bonificadas = %s"
                 % (_fmt(cob), _fmt(bon), _fmt(total)))
    P(" · ".join(linha))

    P("Este cronograma organiza a distribuição das agendas (horas contratadas + bonificadas), "
      "com os macro tópicos de cada visita/treinamento. As datas são previstas e podem ser "
      "ajustadas em comum acordo; o limite de 5 dias úteis após a liberação do levantamento é "
      "balizador. As agendas serão incluídas no SICLA (no nome do cliente) e compartilhadas com "
      "o usuário líder.")

    tab = doc.add_table(rows=1, cols=6); tab.style = "Table Grid"
    cabec = ["#", "Data prevista", "Etapa / Treinamento", "Macro tópicos", "Horas", "Modalidade"]
    for i, txt in enumerate(cabec):
        tab.rows[0].cells[i].paragraphs[0].add_run(txt)
    shade_header(tab.rows[0])
    for i, ag in enumerate(agendas, 1):
        c = tab.add_row().cells
        c[0].text = str(i)
        c[1].text = ag.get("data", "")
        c[2].text = ag.get("etapa", "")
        c[3].text = ag.get("topicos", "")
        h = _num(ag.get("horas"))
        c[4].text = (_fmt(h) + " h") if h else ""
        c[5].text = ag.get("modalidade", "")
    tot = tab.add_row().cells
    tot[3].paragraphs[0].add_run("Total").bold = True
    tot[4].paragraphs[0].add_run("%s h" % _fmt(soma_horas)).bold = True

    doc.add_paragraph()
    C.docx_heading(doc, "Observações", size=12)
    for it in (d.get("observacoes") or [
        "Datas previstas, sujeitas a ajuste conforme a disponibilidade das partes.",
        "Agendas registradas no SICLA no nome do cliente; treinamento/parametrização interna como tipo 84.",
        "Cronograma compartilhado via Google Drive (link enviado ao usuário líder).",
        "Cliente de menor porte: cronograma simplificado, deixando claro o que será e o que não será atendido.",
    ]):
        C.docx_bullet(doc, it)

    doc.add_paragraph()
    hoje = datetime.date.today()
    P(d.get("cidade_data") or ("Novo Hamburgo, %d de %s de %d."
                               % (hoje.day, _MESES[hoje.month - 1], hoje.year)))

    C.ensure_out()
    fname = "Cronograma_Implantacao_%s.docx" % C.slug(cliente)
    path = os.path.join(C.OUT, fname)
    doc.save(path)
    print("OK: %s -> %s" % (fname, path))


if __name__ == "__main__":
    args = sys.argv[1:]
    main(*args) if args else main()
