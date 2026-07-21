# -*- coding: utf-8 -*-
"""
caracterizacao.py — testes de CARACTERIZAÇÃO dos geradores Office.

Existe para sustentar o porte dos geradores de Python para Node/TypeScript, exigido pelos
Padrões de Desenvolvimento da Rech (§4.2/§4.7 do PADRAO-RECH.md). A §4.7 é explícita: antes
de escrever a primeira linha na linguagem destino, cubra o original com testes que
documentam o que o sistema **faz** hoje — não o que deveria fazer.

O que este arquivo captura NÃO são os bytes do .docx/.xlsx: arquivos Office são ZIPs com
timestamps embutidos, então dois arquivos idênticos em conteúdo têm bytes diferentes.
O contrato observável que o porte precisa preservar é o CONTEÚDO:

  .docx -> texto de cada parágrafo, na ordem, e o texto de cada célula de cada tabela
  .xlsx -> nomes das abas, na ordem, e o valor de cada célula preenchida

Uso (a partir da pasta tools/):
    python caracterizacao.py --gravar      # (re)gera os snapshots dourados
    python caracterizacao.py               # verifica os geradores contra os snapshots
    python caracterizacao.py --gerador gerar_raid   # só um gerador

Código de saída != 0 se algum gerador divergir do snapshot — é o sinal de regressão.
Os snapshots ficam em tools/caracterizacao/*.json e SÃO versionados: eles são a
especificação executável que o código em Node/TS terá de cumprir.
"""
import argparse
import contextlib
import importlib
import io
import json
import os
import datetime
import sys
import time

HERE = os.path.dirname(os.path.abspath(__file__))
if HERE not in sys.path:
    sys.path.insert(0, HERE)
import _common as C  # noqa: E402

OUT = C.OUT
SNAP_DIR = os.path.join(HERE, "caracterizacao")

# Mesma lista de verificar.py — (módulo, prefixo do arquivo de saída, extensão).
GERADORES = [
    ("gerar_kit_mudanca", "Kit_Gestao_Mudanca_", "xlsx"),
    ("gerar_roteiros_teste", "Roteiros_SIT_UAT_", "xlsx"),
    ("gerar_aceite_uat", "Termo_Aceite_UAT_", "docx"),
    ("gerar_reconciliacao_conversao", "Reconciliacao_Conversao_", "xlsx"),
    ("gerar_painel_hypercare", "Painel_Hypercare_", "xlsx"),
    ("gerar_log_fitgap", "Log_FitGap_", "xlsx"),
    ("gerar_painel_kpi", "Painel_KPIs_", "xlsx"),
    ("gerar_raid", "RAID_", "xlsx"),
    ("gerar_dossie_cliente", "Dossie_", "docx"),
    ("gerar_projeto_implantacao", "Projeto_Implantacao_", "docx"),
    ("gerar_termo_encerramento", "Termo_Encerramento_", "docx"),
    ("gerar_cronograma", "Cronograma_Implantacao_", "docx"),
    ("gerar_levantamento", "Levantamento_", "docx"),
    ("gerar_checklist_consultor", "CheckList_Consultor_", "xlsx"),
]


# Vários geradores carimbam a data de hoje ("gerado em 21/07/2026"). Sem mascarar, o
# snapshot quebraria sozinho no dia seguinte e o teste viraria ruído que todo mundo ignora.
#
# Mascaramos APENAS a data de hoje — não qualquer data. Mascarar todas apagaria as datas de
# NEGÓCIO (data da virada, dias do hypercare, prazos), que vêm dos YAMLs e são estáveis; sem
# elas no contrato, um porte com aritmética de datas errada passaria despercebido.
_HOJE = datetime.date.today().strftime("%d/%m/%Y")

# O Cronograma fecha com "Novo Hamburgo, 21 de julho de 2026." — a data de hoje POR EXTENSO,
# que a máscara dd/mm/aaaa não pega. Sem mascarar as duas formas, aquele snapshot quebraria
# sozinho no dia seguinte.
_MESES = ["janeiro", "fevereiro", "março", "abril", "maio", "junho", "julho",
          "agosto", "setembro", "outubro", "novembro", "dezembro"]
_hoje = datetime.date.today()
_HOJE_EXTENSO = "%d de %s de %d" % (_hoje.day, _MESES[_hoje.month - 1], _hoje.year)


def _mascarar(valor):
    """Substitui a data de HOJE por <HOJE>, mantendo as demais datas no contrato."""
    if not isinstance(valor, str):
        return valor
    return valor.replace(_HOJE, "<HOJE>").replace(_HOJE_EXTENSO, "<HOJE_EXTENSO>")


def extrair_docx(caminho):
    """Conteúdo observável de um .docx: parágrafos e tabelas, na ordem."""
    from docx import Document

    d = Document(caminho)
    return {
        "tipo": "docx",
        "paragrafos": [_mascarar(p.text) for p in d.paragraphs],
        "tabelas": [
            [[_mascarar(c.text) for c in linha.cells] for linha in t.rows] for t in d.tables
        ],
    }


def extrair_xlsx(caminho):
    """Conteúdo observável de um .xlsx: abas e células preenchidas, na ordem."""
    from openpyxl import load_workbook

    wb = load_workbook(caminho, data_only=False)
    abas = {}
    for nome in wb.sheetnames:
        ws = wb[nome]
        linhas = []
        for linha in ws.iter_rows(values_only=True):
            # Normaliza para texto e descarta a cauda de células vazias da linha.
            valores = ["" if v is None else _mascarar(str(v)) for v in linha]
            while valores and valores[-1] == "":
                valores.pop()
            linhas.append(valores)
        # Descarta a cauda de linhas totalmente vazias.
        while linhas and not any(linhas[-1]):
            linhas.pop()
        abas[nome] = linhas
    return {"tipo": "xlsx", "abas_ordem": wb.sheetnames, "abas": abas}


def rodar_gerador(modname, prefix, kind):
    """Executa o gerador e devolve o conteúdo extraído do arquivo que ELE acabou de escrever.

    `exemplos/` acumula saídas antigas de vários clientes (Cronograma_ tem 10, Levantamento_
    tem 11). Escolher "o último em ordem alfabética" pegava um arquivo de outro cliente, de
    semanas atrás — o snapshot congelava um artefato morto e o teste passava para sempre sem
    olhar o gerador. Por isso a escolha é por data de modificação, e o arquivo TEM de ter sido
    escrito nesta execução: se não foi, o gerador não produziu nada e isso é uma falha.
    """
    os.makedirs(OUT, exist_ok=True)
    mod = importlib.import_module(modname)
    marco = time.time()
    with contextlib.redirect_stdout(io.StringIO()):
        mod.main()
    saidas = [f for f in os.listdir(OUT) if f.startswith(prefix)]
    if not saidas:
        raise RuntimeError("gerador não produziu arquivo com prefixo %r" % prefix)
    caminho = max((os.path.join(OUT, f) for f in saidas), key=os.path.getmtime)
    # Tolerância de 2s: FAT tem resolução de 2 segundos em mtime.
    if os.path.getmtime(caminho) < marco - 2:
        raise RuntimeError(
            "nenhum arquivo com prefixo %r foi escrito nesta execução (mais recente: %s)"
            % (prefix, os.path.basename(caminho))
        )
    return extrair_xlsx(caminho) if kind == "xlsx" else extrair_docx(caminho)


def caminho_snapshot(modname):
    return os.path.join(SNAP_DIR, "%s.json" % modname)


def gravar(modname, conteudo):
    os.makedirs(SNAP_DIR, exist_ok=True)
    with open(caminho_snapshot(modname), "w", encoding="utf-8") as fp:
        json.dump(conteudo, fp, ensure_ascii=False, indent=2, sort_keys=True)


def comparar(modname, atual):
    """Compara com o snapshot; devolve (ok, detalhe)."""
    caminho = caminho_snapshot(modname)
    if not os.path.exists(caminho):
        return False, "sem snapshot — rode com --gravar"
    with open(caminho, encoding="utf-8") as fp:
        esperado = json.load(fp)
    if esperado == atual:
        return True, "conteúdo idêntico ao snapshot"
    # Aponta a primeira diferença de forma útil, sem despejar o documento inteiro.
    if esperado.get("tipo") != atual.get("tipo"):
        return False, "tipo mudou: %s -> %s" % (esperado.get("tipo"), atual.get("tipo"))
    if atual["tipo"] == "docx":
        pe, pa = esperado["paragrafos"], atual["paragrafos"]
        if len(pe) != len(pa):
            return False, "nº de parágrafos: %d -> %d" % (len(pe), len(pa))
        for i, (a, b) in enumerate(zip(pe, pa)):
            if a != b:
                return False, "parágrafo %d: %r -> %r" % (i, a[:60], b[:60])
        return False, "tabelas divergiram"
    ae, aa = esperado["abas_ordem"], atual["abas_ordem"]
    if ae != aa:
        return False, "abas: %s -> %s" % (ae, aa)
    for nome in ae:
        if esperado["abas"][nome] != atual["abas"][nome]:
            le, la = esperado["abas"][nome], atual["abas"][nome]
            if len(le) != len(la):
                return False, "aba %r: %d -> %d linhas" % (nome, len(le), len(la))
            for i, (a, b) in enumerate(zip(le, la)):
                if a != b:
                    return False, "aba %r linha %d: %r -> %r" % (nome, i, a, b)
    return False, "divergência não localizada"


def main():
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--gravar", action="store_true", help="(re)grava os snapshots dourados")
    ap.add_argument("--gerador", help="roda só um gerador (nome do módulo)")
    args = ap.parse_args()

    alvos = [g for g in GERADORES if not args.gerador or g[0] == args.gerador]
    if not alvos:
        print("Gerador %r não está na lista." % args.gerador)
        return 2

    falhas = 0
    for modname, prefix, kind in alvos:
        try:
            conteudo = rodar_gerador(modname, prefix, kind)
        except SystemExit as e:
            print("PULAR  %-32s %s" % (modname, str(e).splitlines()[0][:60]))
            continue
        except Exception as e:  # noqa: BLE001 — queremos o relatório completo
            print("FALHA  %-32s %s: %s" % (modname, type(e).__name__, str(e)[:60]))
            falhas += 1
            continue

        if args.gravar:
            gravar(modname, conteudo)
            print("GRAVOU %-32s %s" % (modname, kind))
            continue

        ok, detalhe = comparar(modname, conteudo)
        print("%s %-32s %s" % ("OK    " if ok else "DIFERE", modname, detalhe))
        if not ok:
            falhas += 1

    print("\n%d gerador(es), %d falha(s)." % (len(alvos), falhas))
    return 1 if falhas else 0


if __name__ == "__main__":
    raise SystemExit(main())
