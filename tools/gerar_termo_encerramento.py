# -*- coding: utf-8 -*-
"""Gera o Termo de Encerramento do Projeto de Implantação (.docx) FIEL ao
template real (versão nova: "Rech® Sistemas de Gestão" / setor "Consultoria").

Uso:
    python gerar_termo_encerramento.py [data/termo.yaml]
"""
import os
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import _common as C

NAVY = RGBColor(0x1F, 0x4E, 0x78)


def shade_header(row, fill="1F4E78"):
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), fill)
        tcPr.append(shd)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF); r.font.bold = True


def main(termo_path="data/termo.yaml"):
    d = C.load_yaml(os.path.basename(termo_path))
    cliente = d.get("cliente", "")
    num = d.get("numero_projeto", "")
    num_txt = f" (nº {num})" if num else ""

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    def H(txt, level=1):
        h = doc.add_heading(txt, level=level)
        for r in h.runs:
            r.font.color.rgb = NAVY
        return h

    def P(txt="", bold=False):
        p = doc.add_paragraph(); p.add_run(txt).bold = bold
        return p

    def B(items):
        for it in items or []:
            doc.add_paragraph(str(it), style="List Bullet")

    t = doc.add_heading("Termo de encerramento do projeto de implantação do SIGER®", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    P("").add_run(f"Cliente: {cliente}").bold = True

    H("Processo de transição e finalização da implantação", 1)
    P(f"O presente termo tem por finalidade oficializar o encerramento da implantação do ERP "
      f"SIGER® junto à empresa {cliente}, cujo escopo do projeto original{num_txt} fora "
      f"estabelecido e validado entre as partes na data de {d.get('data_validacao','____')}.")
    P("Diante do encerramento do projeto referenciado neste termo, a empresa supracitada "
      "desvincula-se do atendimento do setor de implantação da Rech, passando a ter como canal "
      "principal de atendimento o setor de suporte, cujos meios de contatos são indicados a seguir:")
    B([
        "Dúvidas pontuais de uso e situações que necessitem suporte técnico deverão ser reportadas ao setor de Suporte da Rech®:",
        "Os contatos receberão prioridade de retorno na seguinte ordem:",
        "1º Contatos via telefone (51-3582.4001) – (atendimento prioritário)",
        "2º Contatos via e-mail (suporte@rech.com.br)",
        "3º Contatos via Teams",
        "Horário de funcionamento do suporte será conforme previsto em contrato.",
        "O Sistema SIGER® também dispõe de um manual próprio que pode ser consultado pela tecla de atalho “F1”, quando estiver acessando qualquer módulo.",
        "Quando detectadas demandas de natureza diferente ao propósito do suporte Rech®, estes serão redirecionados para um contato do setor de “Consultoria” da Rech® Sistemas de Gestão.",
    ])

    H("Do encerramento", 1)
    P("Considerações para encerramento do projeto de implantação.")
    P("As rotinas aplicadas atenderam as necessidades previstas com a implantação do ERP SIGER® "
      "descritas no projeto o qual está referenciada neste termo.")
    P("No ato de formalização deste termo, ratificamos a inexistência de pendências ou restrições "
      "ao processo de implantação por parte da contratante.")

    H("Das alterações e incrementos tratados fora do escopo original", 1)
    P("Neste tópico serão apresentados aspectos que refletiram em incrementos, mudanças de "
      "definições e alterações do processo inicialmente mapeado e da sua efetiva aderência.")
    alts = d.get("alteracoes") or []
    if alts:
        for a in alts:
            if a.get("modulo"): P(a["modulo"], bold=True)
            if a.get("rotina"): P(a["rotina"], bold=True)
            if a.get("descricao"): P(a["descricao"])
    else:
        P("Não houve alterações significativas no projeto original, tais como substituição, "
          "cancelamento ou incrementos de módulos e mudanças de rotinas que impactassem no "
          "objetivo final de uso do SIGER®.")

    H("Do resumo geral da implantação", 1)
    tab = doc.add_table(rows=1, cols=5); tab.style = "Table Grid"
    for i, txt in enumerate(["Módulo", "Adicional", "Processo de Implantação", "Status de Uso", "Obs."]):
        tab.rows[0].cells[i].paragraphs[0].add_run(txt)
    shade_header(tab.rows[0])
    for m in d.get("resumo_modulos", []):
        cells = tab.add_row().cells
        cells[0].text = m.get("modulo", ""); cells[1].text = m.get("adicional", "")
        cells[2].text = m.get("processo", ""); cells[3].text = m.get("status_uso", "")
        cells[4].text = m.get("obs", "")
    for i, obs in enumerate(d.get("observacoes") or [], 1):
        P(f"Observação {i}: {obs}")

    H("Dos pontos que serão sequenciados e entregues pelo setor de implantação mesmo após o encerramento do projeto", 1)
    pend = d.get("pendencias") or []
    if pend:
        for p in pend:
            P(f"Pendência: {p.get('pendencia','')}", bold=True)
            P(f"Técnico responsável: {p.get('tecnico','')}")
            P(f"Detalhamento da ação: {p.get('detalhamento','')}")
    else:
        P("Não há pendências a serem sequenciadas após o encerramento do projeto.")

    doc.add_paragraph()
    P(d.get("cidade_data", ""))
    doc.add_paragraph()
    P("Assinatura Rech\t\t\tAssinatura Cliente")

    C.ensure_out()
    fname = f"Termo_Encerramento_{C.slug(cliente)}.docx"
    path = os.path.join(C.OUT, fname)
    doc.save(path)
    print(f"OK: {fname} -> {path}")


if __name__ == "__main__":
    args = sys.argv[1:]
    main(*args) if args else main()
