# -*- coding: utf-8 -*-
"""Geração FIEL das documentações das fases a partir dos layouts cadastrados.

Pega o arquivo VIGENTE do modelo (Cadastro → Modelos de Documentos), troca só os
placeholders conhecidos pelos dados do projeto e salva um novo arquivo, com a
mesma estrutura do anexo. Os placeholders manuais (estimativas, quadros de
perguntas, <XX> por área) permanecem como guia para o consultor.
"""
import os
import datetime

import _common as C
import db
import preencher_layout as PL

_MESES = ["", "janeiro", "fevereiro", "março", "abril", "maio", "junho", "julho",
          "agosto", "setembro", "outubro", "novembro", "dezembro"]


def _data_iso(s):
    s = (s or "").strip()
    for fmt in ("%Y-%m-%d", "%d/%m/%Y"):
        try:
            return datetime.datetime.strptime(s, fmt).date()
        except ValueError:
            pass
    return None


def _por_extenso(d):
    return "%d de %s de %d" % (d.day, _MESES[d.month], d.year)


def _hoje():
    return datetime.date.today()


def _num(v):
    v = str(v or "").strip()
    return v


# ---- mapas de preenchimento por modelo (literal -> valor) -------------------
def _repl_termo(p):
    cli = (p.get("cliente") or "").strip()
    repl, paras = [], []
    if cli:
        repl.append(("<Razão Social Longa>", cli))
    if (p.get("numero_projeto") or "").strip():
        repl.append(("<Número do projeto quando se aplicar>", p["numero_projeto"].strip()))
    d = _data_iso(p.get("data_encerramento")) or _hoje()
    paras.append(("Novo Hamburgo", "Novo Hamburgo, %s." % _por_extenso(d)))
    return repl, paras


def _repl_projeto(p):
    val = _conteudo(p, "projeto")
    cli = (p.get("cliente") or "").strip()
    repl, paras = [], []
    if cli:
        # ocorrências no corpo (literais ASCII) + linha-cabeçalho via prefixo
        repl += [("<Nome do Cliente >", cli), ("<Nome do Cliente>", cli)]
        paras.append(("Nome do Cliente:", "Nome do Cliente: %s" % cli))
    if val("cnpj", "cnpj"):
        paras.append(("CNPJ", "CNPJ: %s" % val("cnpj", "cnpj")))
    if val("objetivos", "observacoes"):
        repl.append(("<(preencher)>", val("objetivos", "observacoes")))
    # Equipes (telas de edição estruturada)
    if val("gerente_contas", "gci"):
        paras.append(("Gerente de Contas do Projeto", "Gerente de Contas do Projeto: %s" % val("gerente_contas", "gci")))
    if val("redator"):
        paras.append(("Redator do Projeto", "Redator do Projeto: %s" % val("redator")))
    if val("consultor", "consultor"):
        paras.append(("Consultor/Implantador", "Consultor/Implantador: %s" % val("consultor", "consultor")))
    hb, hc = _num(p.get("horas_bonificadas")), _num(p.get("horas_cobradas"))
    if hb:
        repl.append(("<XX horas bonificadas>", "%s horas bonificadas" % hb))
    if hc:
        repl.append(("<XX horas cobradas>", "%s horas cobradas" % hc))
    paras.append(("Novo Hamburgo", "Novo Hamburgo, %s." % _por_extenso(_hoje())))
    return repl, paras


def _conteudo(p, doc):
    """Valores estruturados (DocConteudo) do documento; val(campo, origem_projeto)."""
    cont = db.doc_conteudo(p.get("id"), doc) if p.get("id") else {}

    def val(campo, orig=None):
        return (cont.get(campo) or (p.get(orig) if orig else "") or "").strip()
    return val


def _repl_levantamento(p):
    val = _conteudo(p, "levantamento")
    cli = (p.get("cliente") or "").strip()
    repl, paras = [], []
    if cli:
        repl.append(("< Nome Cliente >", cli))
        paras.append(("<Razão Social:", "Razão Social: %s" % cli))
    d = _data_iso(p.get("data_levantamento"))
    if d:
        repl.append(("<xx/xx/xxxxx>", d.strftime("%d/%m/%Y")))
    resp = " / ".join(x for x in (p.get("gci"), p.get("consultor")) if x)
    if resp:
        repl.append(("<xxxxxxxxxxxxx>", resp))
    # Identificação da empresa (telas de edição estruturada)
    if val("ramo", "ramo"):
        paras.append(("Ramo Atividade", "Ramo Atividade: %s" % val("ramo", "ramo")))
    if val("produto"):
        paras.append(("Produto:", "Produto: %s" % val("produto")))
    if val("software_atual"):
        paras.append(("Fornecedor Atual Software", "Fornecedor Atual Software: %s" % val("software_atual")))
    if val("filiais"):
        repl.append(("<Localização / Filiais:>", "Localização / Filiais: %s" % val("filiais")))
    if val("objetivos", "observacoes"):
        paras.append(("Observações / Objetivos", "Observações / Objetivos: %s" % val("objetivos", "observacoes")))
    if val("qtd_usuarios"):
        repl.append(("<Quantidade usuários e identificação: >", "Quantidade usuários e identificação: %s" % val("qtd_usuarios")))
    return repl, paras


_GERADORES_DOCX = {"termo": _repl_termo, "projeto": _repl_projeto, "levantamento": _repl_levantamento}


def _topicos_por_modulo(modulos_str):
    """Para cada módulo contratado (sigla em projeto.modulos), busca os tópicos do
    cadastro Índice de Tópicos. Devolve [{sigla, nome, topicos:[linhas]}] na ordem informada."""
    import re as _re
    sigs = [m.strip().upper() for m in _re.split(r"[,;\n]+", modulos_str or "") if m.strip()]
    nomes = {m["sigla"].upper(): m["nome"] for m in db.indice_modulos()}
    out, vistos = [], set()
    for sig in sigs:
        if sig in vistos:
            continue
        vistos.add(sig)
        linhas, _ = db.indice_listar(modulo=sig)
        if linhas:
            out.append({"sigla": sig, "nome": nomes.get(sig, ""), "topicos": linhas})
    return out


def _anexar_topicos_levantamento(doc, modulos_str):
    """Acrescenta ao Levantamento, por módulo contratado, as perguntas/tópicos do Índice
    de Tópicos a serem respondidas. Não depende de estilos do template (robusto)."""
    grupos = _topicos_por_modulo(modulos_str)
    if not grupos:
        return 0

    def linha(txt="", bold=False):
        p = doc.add_paragraph()
        if txt:
            p.add_run(txt).bold = bold
        return p

    linha()
    linha("Tópicos a levantar por módulo contratado", bold=True)
    linha("Itens do Índice de Tópicos a responder no Levantamento, por módulo contratado.")
    total = 0
    for g in grupos:
        linha()
        linha("%s — %s" % (g["sigla"], g["nome"] or ""), bold=True)
        atual = None
        for l in g["topicos"]:
            adic = (l.get("adicional") or "").strip()
            if adic and adic != atual:
                linha(adic, bold=True)
                atual = adic
            doc.add_paragraph("•  " + (l.get("topico") or "").strip())
            total += 1
    return total


def _anexar_respostas_projeto(doc, projeto_id):
    """Acrescenta ao Projeto o detalhamento do Levantamento (tópicos JÁ respondidos no
    painel), por módulo contratado — liga as respostas do Levantamento ao Projeto."""
    if not projeto_id:
        return 0
    rs = [r for r in db.levantamento_respostas(projeto_id) if (r.get("resposta") or "").strip()]
    if not rs:
        return 0

    def linha(txt="", bold=False):
        p = doc.add_paragraph()
        if txt:
            p.add_run(txt).bold = bold
        return p

    linha()
    linha("Detalhamento do Levantamento por módulo", bold=True)
    linha("Rotinas e particularidades identificadas no Levantamento (respostas registradas no painel).")
    mod_atual = adic_atual = None
    total = 0
    for r in rs:
        sig = r.get("modulo_sigla", "")
        if sig != mod_atual:
            mod_atual, adic_atual = sig, None
            linha()
            linha("%s — %s" % (sig, r.get("modulo", "") or ""), bold=True)
        adic = (r.get("adicional") or "").strip()
        if adic and adic != adic_atual:
            adic_atual = adic
            linha(adic, bold=True)
        p = doc.add_paragraph()
        p.add_run(((r.get("topico") or "").strip() + ": ")).bold = True
        p.add_run((r.get("resposta") or "").strip())
        total += 1
    return total


def _norm(s):
    """Normaliza rótulo: colapsa espaços, tira pontuação final e baixa caixa."""
    return " ".join(str(s or "").split()).strip().rstrip(":. ").lower()


def _preencher_cronograma_xlsx(wb, projeto):
    """Preenche o cabeçalho (Consultor/Horas) por rótulo e as linhas de visita a partir
    do cronograma do projeto, sem tocar nas colunas com fórmula (Total/Horário)."""
    ws = wb["Cronograma de visitas"] if "Cronograma de visitas" in wb.sheetnames else wb.worksheets[0]
    alvos = {
        "consultor": (projeto.get("consultor") or "").strip(),
        "horas do planejamento": str(projeto.get("horas_cobradas") or "").strip(),
        "hrs previstas bonificadas": str(projeto.get("horas_bonificadas") or "").strip(),
    }
    for row in ws.iter_rows(min_row=1, max_row=8):       # cabeçalho: valor à direita do rótulo
        for c in row:
            if isinstance(c.value, str):
                v = alvos.get(_norm(c.value))
                if v:
                    ws.cell(row=c.row, column=c.column + 1, value=v)
    itens = db.cronograma_do_projeto(projeto.get("id")) if projeto.get("id") else []
    if not itens:
        return 0
    hdr, cols = None, {}
    for row in ws.iter_rows(min_row=1, max_row=15):       # acha a linha de cabeçalho da tabela
        vals = {_norm(c.value): c.column for c in row if isinstance(c.value, str)}
        if "data" in vals and "o que será abordado" in vals:
            hdr, cols = row[0].row, vals
            break
    if not hdr:
        return 0
    cons = (projeto.get("consultor") or "").strip()
    for i, it in enumerate(itens):
        r = hdr + 1 + i

        def setc(label, val):
            col = cols.get(label)
            if col and val:
                ws.cell(row=r, column=col, value=val)

        abordado = (it.get("etapa") or "").strip()
        if it.get("topicos"):
            abordado = (abordado + " — " + it["topicos"]).strip(" —")
        setc("data", it.get("data", ""))
        setc("local", it.get("modalidade", ""))
        setc("técnico", cons)
        setc("o que será abordado", abordado)
    return len(itens)


def _preencher_termo_grade(doc, modulos_str):
    """Preenche a grade 'Resumo Geral' do Termo (Módulo/Adicional/Processo/Status) a
    partir dos módulos contratados. Usa as linhas existentes e cria novas se faltar."""
    import re as _re
    sigs = [m.strip() for m in _re.split(r"[,;\n]+", modulos_str or "") if m.strip()]
    if not sigs or not doc.tables:
        return 0
    t = doc.tables[0]          # Resumo Geral: Módulo | Adicional | Processo | Status de Uso | Obs.
    if len(t.columns) < 4:
        return 0
    base = t.rows[1:]          # exclui o cabeçalho
    for i, sig in enumerate(sigs):
        row = base[i] if i < len(base) else t.add_row()
        cells = row.cells
        cells[0].text = sig
        cells[2].text = "Implantado"
        cells[3].text = "Sim"
    return len(sigs)


# Mapa sigla -> palavras-chave dos blocos do Levantamento que ela mantém (AJUSTÁVEL).
_SIGLA_BLOCOS = {
    "FAT": ["vendas e faturamento"], "PDV": ["vendas e faturamento"],
    "OSE": ["vendas e faturamento"], "SAC": ["vendas e faturamento"],
    "GIN": ["produção"], "GCA": ["produção"],
    "EST": ["compras/estoque"], "COM": ["compras/estoque"], "TLO": ["compras/estoque"],
    "FIN": ["gestão financeira"], "GCO": ["gestão financeira"],
    "CTB": ["gestão fiscal"], "LFI": ["gestão fiscal"], "GPA": ["gestão fiscal"], "AUE": ["gestão fiscal"],
    "FPA": ["folha de pagamento"],
    "PWC": ["portal de funcion", "portal de vagas"], "PGP": ["portal de funcion", "portal de vagas"],
    "RHU": ["recrutamento", "treinamen", "saúde ocupacional", "segurança do trabalho",
            "avaliação", "cargos e sal"],
}
_BLOCOS_FIXOS = ["cliente/fornecedor", "produto"]   # blocos fundacionais, sempre mantidos

# Nome de exibição da área (bloco) p/ agrupar a tela do Levantamento.
_BLOCO_DISPLAY = {
    "vendas e faturamento": "Vendas e Faturamento", "produção": "Produção",
    "compras/estoque": "Compras / Estoque", "gestão financeira": "Gestão Financeira",
    "gestão fiscal": "Gestão Fiscal, Contábil e Patrimonial", "folha de pagamento": "Folha de Pagamento",
    "portal de funcion": "Portais", "portal de vagas": "Portais",
    "recrutamento": "RHU", "treinamen": "RHU", "saúde ocupacional": "RHU",
    "segurança do trabalho": "RHU", "avaliação": "RHU", "cargos e sal": "RHU",
}


def area_do_modulo(sigla):
    """Nome de exibição da área (bloco) de um módulo (ou '' se não tiver bloco próprio)."""
    for kw in _SIGLA_BLOCOS.get((sigla or "").upper(), []):
        if kw in _BLOCO_DISPLAY:
            return _BLOCO_DISPLAY[kw]
    return ""


def _inserir_textos_depois(anchor_p, textos):
    """Insere parágrafos de texto simples logo após `anchor_p` (preservando ordem)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    el = anchor_p._p
    for txt in textos:
        p = OxmlElement("w:p")
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = txt
        r.append(t)
        p.append(r)
        el.addnext(p)
        el = p


def _montar_blocos_levantamento(doc, projeto):
    """Mantém apenas os blocos 'Mapeamento de processo – ÁREA' dos módulos contratados
    (remove os demais) e injeta as perguntas/respostas em cada bloco mantido. Módulos
    contratados sem bloco próprio são acrescentados ao final. Devolve (mantidos, removidos)."""
    import re as _re
    sigs = [m.strip().upper() for m in _re.split(r"[,;\n]+", projeto.get("modulos", "") or "") if m.strip()]
    nomes = {m["sigla"].upper(): m["nome"] for m in db.indice_modulos()}
    por_sigla = {}
    for r in (db.levantamento_respostas(projeto.get("id")) if projeto.get("id") else []):
        por_sigla.setdefault((r.get("modulo_sigla") or "").upper(), []).append(r)

    def itens_da_sigla(sig):
        its = por_sigla.get(sig)
        if its:
            return its
        tops, _ = db.indice_listar(modulo=sig)   # sem respostas semeadas: usa o Índice
        return tops

    # 1) varre os blocos (heading "Mapeamento de processo – X")
    blocos, atual = [], None
    for p in doc.paragraphs:
        t = p.text.strip()
        if _norm(t).startswith("mapeamento de processo") and ("-" in t or "–" in t):
            area = _norm(_re.split(r"[-–]", t, 1)[1])
            atual = {"area": area, "ps": [p], "colar": None, "modval": None, "viu_mod": False}
            blocos.append(atual)
        elif atual is not None:
            atual["ps"].append(p)
            tl = _norm(t)
            if "colar aqui" in tl and atual["colar"] is None:
                atual["colar"] = p
            if tl.startswith("módulos previsto") or tl.startswith("modulos previsto"):
                atual["viu_mod"] = True
            elif atual["viu_mod"] and atual["modval"] is None and t.startswith("<") and t.endswith(">"):
                atual["modval"] = p   # o "<xxxxx>" logo após "Módulos Previstos:"

    # 2) decide manter/remover e injeta perguntas
    remover, injetadas, removidos = [], set(), 0
    for b in blocos:
        area = b["area"]
        fixo = any(k in area for k in _BLOCOS_FIXOS)
        siglas_bloco = [s for s in sigs if any(k in area for k in _SIGLA_BLOCOS.get(s, []))]
        if not fixo and not siglas_bloco:
            remover.extend(b["ps"])               # bloco não contratado -> remover
            removidos += 1
            continue
        if siglas_bloco:
            if b["modval"] is not None:           # "<xxxxx>" -> sigla(s) do bloco
                rot = ", ".join(("%s — %s" % (s, nomes[s])) if nomes.get(s) else s for s in siglas_bloco)
                PL._aplica_no_paragrafo(b["modval"], rot)
            if b["colar"] is not None:            # injeta perguntas (com respostas) no bloco
                textos = []
                for s in siglas_bloco:
                    for it in itens_da_sigla(s):
                        top = (it.get("topico") or "").strip()
                        resp = (it.get("resposta") or "").strip()
                        textos.append("•  " + top + (": " + resp if resp else ""))
                        injetadas.add(s)
                _inserir_textos_depois(b["colar"], textos)

    # 3) remove os parágrafos dos blocos não contratados
    for p in remover:
        el = p._p
        if el.getparent() is not None:
            el.getparent().remove(el)

    # 4) módulos contratados sem bloco próprio -> acrescenta ao final
    sobrando = [s for s in sigs if s not in injetadas]
    if sobrando:
        def linha(txt="", bold=False):
            par = doc.add_paragraph()
            if txt:
                par.add_run(txt).bold = bold
        linha(); linha("Outros módulos contratados", bold=True)
        for s in sobrando:
            linha(); linha("%s — %s" % (s, nomes.get(s, "")), bold=True)
            for it in itens_da_sigla(s):
                top = (it.get("topico") or "").strip()
                resp = (it.get("resposta") or "").strip()
                doc.add_paragraph("•  " + top + (": " + resp if resp else ""))
    return (len(blocos) - removidos), removidos


def _preencher_levantamento_tabelas(doc, projeto):
    """Preenche no Levantamento a tabela 'Módulos/Adicionais (A)' (módulos contratados)
    e a tabela de horas (Cobradas / Bonificadas / Total) com os dados do fechamento."""
    import re as _re
    sigs = [m.strip() for m in _re.split(r"[,;\n]+", projeto.get("modulos", "") or "") if m.strip()]
    nomes = {m["sigla"].upper(): m["nome"] for m in db.indice_modulos()}
    cob = (projeto.get("horas_cobradas") or "").strip()
    bon = (projeto.get("horas_bonificadas") or "").strip()
    for t in doc.tables:
        h0 = (t.rows[0].cells[0].text or "").strip().lower() if t.rows else ""
        # Tabela de horas: Cobradas | Bonificadas | Total
        if "horas cobradas" in h0 and len(t.rows) >= 2 and len(t.columns) >= 3:
            t.rows[1].cells[0].text = cob
            t.rows[1].cells[1].text = bon
            def _n(v):
                m = _re.search(r"\d+(?:[.,]\d+)?", v or "")
                return float(m.group(0).replace(",", ".")) if m else 0.0
            tot = _n(cob) + _n(bon)
            if tot:
                t.rows[1].cells[2].text = ("%g" % tot)
        # Tabela 'Módulos/Adicionais (A)': uma linha por módulo contratado
        if "módulos/adicionais (a)" in h0 and sigs:
            base = t.rows[2:] if len(t.rows) > 2 else []
            for i, sig in enumerate(sigs):
                row = base[i] if i < len(base) else t.add_row()
                nome = nomes.get(sig.upper(), "")
                row.cells[0].text = ("%s — %s" % (sig, nome)) if nome else sig
                if len(row.cells) > 1:
                    row.cells[1].text = "X"   # Necessidade: Sim
    return True


# Grupos do 'Detalhamento das Rotinas' do layout do Projeto -> áreas que contêm (AJUSTÁVEL).
# Um grupo é removido quando NENHUMA das suas áreas foi contratada.
_PROJ_GRUPOS = {
    "gestão comercial": {"vendas e faturamento"},
    "gestão de materiais": {"controle de estoque", "controle de compras"},
    "gestão da produção": {"gestão industrial"},
    "gestão financeira": {"controle financeiro"},
    "gestão de controladoria": {"livros fiscais"},
}


def _eh_marcador(t):
    """True se o parágrafo é um placeholder a substituir: '<...>' ou um 'XX' solto."""
    s = (t or "").strip()
    return (s.startswith("<") and s.endswith(">")) or s.upper() == "XX"


def _preencher_detalhamento_projeto(doc, projeto, guia=False):
    """No 'Detalhamento das Rotinas' do Projeto: mantém SÓ as áreas dos módulos
    contratados (remove as demais e os grupos que ficarem vazios) e preenche cada
    bloco com os dados do Levantamento — 'Módulos Previstos' = módulos contratados
    da área; 'Detalhamento' = respostas do Levantamento da área. A tela do Projeto
    (DocConteudo det_<area>_*) sobrepõe quando preenchida. Com `guia=True` (modelo
    para preenchimento manual), sem respostas usa as PERGUNTAS do Índice como guia."""
    import doc_edit, re as _re
    pid = projeto.get("id")
    sigs = {m.strip().upper() for m in _re.split(r"[,;\n]+", projeto.get("modulos", "") or "") if m.strip()}
    nomes = {m["sigla"].upper(): m["nome"] for m in db.indice_modulos()}
    cont = db.doc_conteudo(pid, "projeto") if pid else {}

    # área (nome normalizado) -> (k, siglas contratadas) e sigla -> área k
    area_info, sigla_area, inter_by_k = {}, {}, {}
    for (k, nome, ss) in doc_edit._PROJ_AREAS:
        inter = [s for s in ss if s in sigs]
        area_info[_norm(nome)] = (k, inter)
        inter_by_k[k] = inter
        for s in ss:
            sigla_area[s] = k
    contratadas = {a for a, (k, inter) in area_info.items() if inter}

    # respostas do Levantamento por área (k) — alimentam o 'Detalhamento'
    resp_por_k = {}
    for r in (db.levantamento_respostas(pid) if pid else []):
        k = sigla_area.get((r.get("modulo_sigla") or "").upper())
        resp = (r.get("resposta") or "").strip()
        if k and resp:
            resp_por_k.setdefault(k, []).append("%s: %s" % ((r.get("topico") or "").strip(), resp))

    def _valor(k, campo):
        v = (cont.get("det_%s_%s" % (k, campo)) or "").strip()
        if v:
            return v
        if campo == "modulos":
            return ", ".join(("%s — %s" % (s, nomes[s])) if nomes.get(s) else s for s in inter_by_k.get(k, []))
        if campo == "detalhamento":
            base = resp_por_k.get(k)
            if base:
                return "  ·  ".join(base)
            if guia:                       # modelo manual: usa as perguntas do Índice como guia
                qs = []
                for sig in inter_by_k.get(k, []):
                    tops, _ = db.indice_listar(modulo=sig)
                    qs += [(t.get("topico") or "").strip() for t in tops if (t.get("topico") or "").strip()]
                return "  ·  ".join(qs)
        return ""

    rotulos = [("módulos previsto", "modulos"), ("detalhamento das rotinas", "detalhamento"),
               ("particularidade", "particularidade"), ("não está previsto", "naoprevisto")]

    paras = doc.paragraphs
    ini = next((i for i, p in enumerate(paras) if _norm(p.text) == "detalhamento das rotinas"), None)
    if ini is None:
        return 0
    fim = next((i for i in range(ini + 1, len(paras)) if _norm(paras[i].text).startswith("responsabilidades")), len(paras))

    # segmenta o trecho em grupos e áreas (com seus parágrafos de conteúdo)
    segs, cur = [], None
    for i in range(ini + 1, fim):
        p = paras[i]
        tl = _norm(p.text)
        if tl in _PROJ_GRUPOS:
            cur = {"kind": "group", "name": tl, "head": p}
            segs.append(cur)
        elif tl in area_info:
            cur = {"kind": "area", "name": tl, "head": p, "paras": []}
            segs.append(cur)
        elif cur and cur["kind"] == "area":
            cur["paras"].append(p)

    remover, n = [], 0
    for seg in segs:
        if seg["kind"] == "group":
            if not (_PROJ_GRUPOS[seg["name"]] & contratadas):
                remover.append(seg["head"])
            continue
        if seg["name"] not in contratadas:          # área não contratada -> remove o bloco inteiro
            remover.append(seg["head"])
            remover.extend(seg["paras"])
            continue
        k, campo = area_info[seg["name"]][0], None    # área contratada -> preenche os <XX>
        for p in seg["paras"]:
            tl = _norm(p.text)
            lab = next((c for (kw, c) in rotulos if kw in tl), None) if tl else None
            if lab:
                campo = lab
            elif campo and _eh_marcador(p.text):
                val = _valor(k, campo)
                if val:
                    PL._aplica_no_paragrafo(p, val)
                    n += 1
                campo = None

    for p in remover:
        el = p._p
        if el.getparent() is not None:
            el.getparent().remove(el)
    return n


def _preencher_projeto_tabelas(doc, projeto):
    """Preenche no Projeto a Tabela de Usuários e o Cronograma Macro a partir do DocConteudo."""
    cont = db.doc_conteudo(projeto.get("id"), "projeto") if projeto.get("id") else {}
    if not cont:
        return
    crono = [("levantamento de requisitos", "crono_levantamento"),
             ("elaboração do cronograma", "crono_cronograma"), ("elaboracao do cronograma", "crono_cronograma"),
             ("parametriz", "crono_parametrizacao"), ("treinamento", "crono_treinamento"),
             ("simula", "crono_simulacao"), ("início do uso", "crono_inicio"), ("inicio do uso", "crono_inicio"),
             ("finaliza", "crono_finalizacao")]
    for t in doc.tables:
        hdr = [(c.text or "").strip().lower() for c in t.rows[0].cells] if t.rows else []
        if not hdr:
            continue
        # Cronograma Macro: Fase | Etapa | Período previsto
        if hdr[0] == "fase" and any("previsto" in h for h in hdr):
            col_per = next((j for j, h in enumerate(hdr) if "previsto" in h), len(hdr) - 1)
            col_et = next((j for j, h in enumerate(hdr) if h == "etapa"), 1)
            for row in t.rows[1:]:
                et = (row.cells[col_et].text or "").strip().lower()
                key = next((k for (kw, k) in crono if kw in et), None)
                if key and cont.get(key):
                    row.cells[col_per].text = cont[key]
        # Tabela de Usuários: Nome | E-mail | Área | Assina
        if hdr[0] == "nome" and any("assina" in h for h in hdr):
            base = t.rows[1:]
            usuarios = []
            for i in range(4):
                nome = (cont.get("usu_%d_nome" % i) or "").strip()
                if nome:
                    usuarios.append([nome, cont.get("usu_%d_email" % i, ""),
                                     cont.get("usu_%d_area" % i, ""), cont.get("usu_%d_assina" % i, "")])
            for idx, u in enumerate(usuarios):
                row = base[idx] if idx < len(base) else t.add_row()
                for j, v in enumerate(u):
                    if j < len(row.cells):
                        row.cells[j].text = v


def _preencher_levantamento_usuarios(doc, projeto):
    """Preenche a tabela de Usuários-chave (Nome | E-mail | Atribuições) do Levantamento."""
    cont = db.doc_conteudo(projeto.get("id"), "levantamento") if projeto.get("id") else {}
    if not cont:
        return
    for t in doc.tables:
        hdr = [(c.text or "").strip().lower() for c in t.rows[0].cells] if t.rows else []
        if hdr and hdr[0] == "nome" and any(("atribuiç" in h or "atribuic" in h) for h in hdr):
            base = t.rows[1:]
            usuarios = []
            for i in range(5):
                nome = (cont.get("usu_%d_nome" % i) or "").strip()
                if nome:
                    usuarios.append([nome, cont.get("usu_%d_email" % i, ""), cont.get("usu_%d_atrib" % i, "")])
            for idx, u in enumerate(usuarios):
                row = base[idx] if idx < len(base) else t.add_row()
                for j, v in enumerate(u):
                    if j < len(row.cells):
                        row.cells[j].text = v
            break


def _saida(slug, cliente, ext):
    nome = "%s_%s.%s" % (slug, C.slug(cliente or "cliente"), ext)
    os.makedirs(C.OUT, exist_ok=True)
    return os.path.join(C.OUT, nome)


def gerar_agenda_xlsx(projeto, atividades):
    """Gera o cronograma de visitas (.xlsx) a partir das atividades ALOCADAS (data+turno)
    do agendador. Devolve o caminho do arquivo. Substitui o cronograma linear pela agenda."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    import datetime as _dt
    DIAS = ["Seg", "Ter", "Qua", "Qui", "Sex", "Sáb", "Dom"]
    TURNOS = {"manha": "Manhã", "tarde": "Tarde"}
    aloc = [a for a in atividades if (a.get("data") and a.get("turno"))]
    aloc.sort(key=lambda a: (a["data"], 0 if a["turno"] == "manha" else 1, a["modulo"], a["seq"]))

    wb = Workbook()
    ws = wb.active
    ws.title = "Cronograma de Visitas"
    ws.append(["Cronograma de Visitas — %s" % (projeto.get("cliente") or "")])
    ws.append([])
    cab = ["Data", "Dia", "Turno", "Módulo", "Visita", "Atividade", "Tipo", "Técnico", "Status"]
    ws.append(cab)
    azul = PatternFill("solid", fgColor="1F4E79")
    borda = Border(*(Side(style="thin", color="D0D7E2"),) * 4)
    hdr_row = ws.max_row
    for c in ws[hdr_row]:
        c.font = Font(bold=True, color="FFFFFF")
        c.fill = azul
        c.alignment = Alignment(horizontal="center")
    for a in aloc:
        try:
            d = _dt.date.fromisoformat(a["data"])
            data_lbl, dia_lbl = d.strftime("%d/%m/%Y"), DIAS[d.weekday()]
        except ValueError:
            data_lbl, dia_lbl = a["data"], ""
        ws.append([data_lbl, dia_lbl, TURNOS.get(a["turno"], a["turno"]), a["modulo"],
                   "V%s" % a["seq"], a.get("descricao", ""), a.get("tipo", ""),
                   a.get("tecnico", ""), a.get("status", "")])
    for row in ws.iter_rows(min_row=hdr_row + 1, max_row=ws.max_row):
        for c in row:
            c.border = borda
            c.alignment = Alignment(vertical="top", wrap_text=True)
    larg = [12, 6, 9, 9, 8, 52, 16, 20, 12]
    for i, w in enumerate(larg):
        ws.column_dimensions[chr(ord("A") + i)].width = w
    ws.freeze_panes = "A%d" % (hdr_row + 1)

    os.makedirs(C.OUT, exist_ok=True)
    dest = os.path.join(C.OUT, "cronograma_visitas_%s.xlsx" % C.slug(projeto.get("cliente") or "cliente"))
    wb.save(dest)
    return dest


def gerar(slug, projeto, modo="auto"):
    """Gera o documento fiel da fase `slug` para o `projeto` (dict). Devolve o caminho.
    `modo='modelo'` (só Projeto) preenche o Detalhamento pelas perguntas do Índice
    quando não há respostas, para preenchimento manual."""
    modelo = next((m for m in db.modelos_documento_listar() if m["slug"] == slug), None)
    if not modelo:
        raise ValueError("Modelo '%s' não cadastrado." % slug)
    base = db.modelo_documento_arquivo_path(modelo["id"])
    if not base or not os.path.exists(base):
        raise FileNotFoundError("Arquivo do modelo '%s' não encontrado." % slug)
    destino = _saida(slug, projeto.get("cliente"), modelo["tipo"])

    if modelo["tipo"] == "docx":
        repl, paras = _GERADORES_DOCX.get(slug, lambda p: ([], []))(projeto)
        doc = PL.preencher_docx(base, repl, paras)
        if slug == "levantamento":   # blocos contratados + perguntas + tabelas (módulos/horas/usuários)
            _montar_blocos_levantamento(doc, projeto)
            _preencher_levantamento_tabelas(doc, projeto)
            _preencher_levantamento_usuarios(doc, projeto)
        elif slug == "projeto":      # detalhamento por área + tabelas (usuários/cronograma) + respostas
            _preencher_detalhamento_projeto(doc, projeto, guia=(modo == "modelo"))
            _preencher_projeto_tabelas(doc, projeto)
            _anexar_respostas_projeto(doc, projeto.get("id"))
        elif slug == "termo":        # preenche a grade Resumo Geral com os módulos contratados
            _preencher_termo_grade(doc, projeto.get("modulos", ""))
        PL.remover_marcadores_docx(doc)   # remove todos os marcadores <...> restantes
        doc.save(destino)
    else:  # xlsx (cronograma)
        repl = []
        cli = (projeto.get("cliente") or "").strip()
        if cli:
            repl.append(("XXXX - RAZÃO SOCIAL LONGA", cli))
        wb = PL.preencher_xlsx(base, repl)
        _preencher_cronograma_xlsx(wb, projeto)   # cabeçalho (consultor/horas) + linhas de visita
        wb.save(destino)
    return destino
