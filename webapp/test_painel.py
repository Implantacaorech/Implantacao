# -*- coding: utf-8 -*-
"""Suíte de testes do painel web (pytest + Flask test_client).
Roda com SQLite temporário, independente do Postgres. Uso:  pytest webapp/test_painel.py
"""
import os
import re
import sys
import tempfile

os.environ.pop("PAINEL_DB_URL", None)
os.environ.pop("PAINEL_SENHA", None)   # ignora a senha mestra do ambiente (login desativado nos testes)
os.environ["PAINEL_DB"] = os.path.join(tempfile.gettempdir(), "painel_pytest.db")

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)
sys.path.insert(0, os.path.join(os.path.dirname(HERE), "tools"))

import pytest          # noqa: E402
import app as A        # noqa: E402
import db              # noqa: E402
import fluxo           # noqa: E402


@pytest.fixture(autouse=True)
def _suite_hermetica(monkeypatch):
    """Isola a suíte de serviços REAIS da máquina (rede): sem envio de e-mail de verdade
    (Gmail/SMTP) e sem consultar o banco externo de disponibilidade (Oracle). Testes que
    exercitam esses caminhos re-mockam por cima (monkeypatch do teste prevalece)."""
    import mailer
    import disponibilidade as D
    monkeypatch.setattr(mailer, "enviar", lambda *a, **k: (True, None))
    monkeypatch.setattr(D, "configurado", lambda cfg=None: False)
    D._CACHE.clear()


@pytest.fixture
def client():
    A.app.config["TESTING"] = True
    with A.app.test_client() as c:
        yield c


def _novo(client, **dados):
    r = client.post("/projetos/novo", data=dados)
    return re.search(r"/projetos/(\d+)", r.headers["Location"]).group(1)


def test_health(client):
    j = client.get("/health").get_json()
    assert j["status"] in ("ok", "degraded")


def test_paginas_principais(client):
    for url in ("/", "/projetos", "/coordenacao", "/monitoramento", "/atividade", "/fluxo"):
        assert client.get(url).status_code == 200


def test_crud_projeto(client):
    pid = _novo(client, cliente="Teste Pytest", situacao="Em andamento")
    assert "Teste Pytest" in client.get("/projetos/%s" % pid).get_data(as_text=True)
    client.post("/projetos/%s/excluir" % pid)
    with db.Session() as s:
        assert s.get(db.Projeto, int(pid)) is None


def test_gate_bloqueia_avanco(client):
    pid = _novo(client, cliente="Gate PT", etapa="Levantamento", modulos="FAT")
    client.post("/projetos/%s/avancar" % pid)   # sem documento -> bloqueia
    with db.Session() as s:
        assert s.get(db.Projeto, int(pid)).etapa == "Levantamento"
    client.post("/projetos/%s/excluir" % pid)


def test_defaults_nao_zeram(client):
    pid = _novo(client, cliente="Defaults PT")   # sem etapa/situacao no form
    with db.Session() as s:
        p = s.get(db.Projeto, int(pid))
        assert p.etapa == "Agendamento" and p.situacao == "Em andamento"
    client.post("/projetos/%s/excluir" % pid)


def test_fluxo_parser():
    d = fluxo.parse_fechamento(
        "Cliente (Razão Social): ACME\nCNPJ: 1\nMódulos contratados (siglas): FAT, CTB\nHoras cobradas: 40\n")
    assert d["cliente"] == "ACME" and d["modulos"] == "FAT, CTB" and d["horas_cobradas"] == "40"


def test_fluxo_para_projeto_contato_separado():
    txt = ("Cliente (Razão Social): ACME\nCNPJ: 1\nContato (nome): Maria\n"
           "E-mail do contato: maria@acme.com\nTelefone: (51) 99999-0000\n"
           "Módulos contratados (siglas): FAT\n")
    p = fluxo.para_projeto(fluxo.parse_fechamento(txt))
    assert p["contato_nome"] == "Maria"
    assert p["contato_email"] == "maria@acme.com"
    assert p["contato_tel"] == "(51) 99999-0000"


def test_fluxo_parse_preenche_contato_na_tela(client):
    txt = ("Cliente (Razão Social): ACME\nContato (nome): Maria\n"
           "E-mail do contato: maria@acme.com\nTelefone: 51999000\n"
           "Módulos contratados (siglas): FAT\nHoras cobradas: 40\n")
    body = client.post("/fluxo/parse", data={"texto": txt}).get_data(as_text=True)
    assert 'value="Maria"' in body
    assert 'value="maria@acme.com"' in body
    assert 'value="51999000"' in body


def test_metricas_alertas():
    proj = [{"id": 1, "etapa": "Projeto", "situacao": "Em risco", "horas_cobradas": "10"}]
    assert db.metricas(proj, {})["n_risco"] == 1
    assert any(a["tipo"] == "risco" for a in db.alertas(proj, {}))


def test_monitoramento_renderiza_paineis_setoriais(client):
    body = client.get("/monitoramento").get_data(as_text=True)
    assert "Centro de Monitoramento Operacional" in body
    assert "Comercial" in body and "Desenvolvimento" in body
    assert "pixel-worker" in body


def test_d_etapa_bloqueia_geracao():
    # D: a geração só é liberada na etapa do documento (ou depois)
    assert A._etapa_permite_gerar("levantamento", "Levantamento")
    assert not A._etapa_permite_gerar("projeto", "Levantamento")
    assert not A._etapa_permite_gerar("termo", "Levantamento")
    assert A._etapa_permite_gerar("projeto", "Projeto")
    assert A._etapa_permite_gerar("cronograma", "Cronograma e Check-list")
    assert A._etapa_permite_gerar("levantamento", "Encerramento")  # docs anteriores sempre ok


def test_c_auto_avanca_com_gate(client):
    # com docs + GCI + consultor, avança Projeto → Designação → Cronograma e Check-list
    pid = int(_novo(client, cliente="Auto Avanca", etapa="Projeto", modulos="FAT",
                    gci="Beto", consultor="Ana"))
    with db.Session() as s:
        for t in ("levantamento", "projeto"):
            s.add(db.Documento(projeto_id=pid, tipo=t, arquivo=t + ".docx", caminho=t + ".docx"))
        s.commit()
    A._auto_avancar(pid)
    with db.Session() as s:
        assert s.get(db.Projeto, pid).etapa == "Cronograma e Check-list"
    client.post("/projetos/%s/excluir" % pid)


def test_etapas_seis_e_gates_acao():
    assert db.ETAPAS[0] == "Agendamento" and "Designação" in db.ETAPAS
    assert db.proxima_etapa("Agendamento") == "Levantamento"
    assert db.proxima_etapa("Projeto") == "Designação"
    # Sem GCI e sem data: não pode avançar
    assert not db.acao_entrada_ok("Levantamento", {"gci": "", "data_levantamento": ""})
    # Só com GCI (sem data): não pode avançar
    assert not db.acao_entrada_ok("Levantamento", {"gci": "João", "data_levantamento": ""})
    # Só com data (sem GCI): não pode avançar
    assert not db.acao_entrada_ok("Levantamento", {"gci": "", "data_levantamento": "2026-07-01"})
    # Com ambos (GCI + data): pode avançar
    assert db.acao_entrada_ok("Levantamento", {"gci": "João", "data_levantamento": "2026-07-01"})
    assert not db.acao_entrada_ok("Cronograma e Check-list", {"consultor": ""})
    assert db.acao_entrada_ok("Cronograma e Check-list", {"consultor": "Ana"})
    # Sem GCI: próxima ação deve ser definir_gci
    cab = db.cabecalho({"etapa": "Agendamento", "gci": "", "data_levantamento": "", "situacao": "Em andamento"}, [])
    assert cab["proxima"] and cab["proxima"]["tipo"] == "acao:definir_gci"
    # Com GCI mas sem data: próxima ação deve ser data_levantamento
    cab2 = db.cabecalho({"etapa": "Agendamento", "gci": "João", "data_levantamento": "", "situacao": "Em andamento"}, [])
    assert cab2["proxima"] and cab2["proxima"]["tipo"] == "acao:data_levantamento"


def test_agendamento_define_gci_e_data_e_avanca(client):
    """Testa o fluxo de duas etapas: (1) definir GCI, (2) definir data, depois avança."""
    _login_como(client, "Administrativo")
    pid = int(_novo(client, cliente="Agendar Co", modulos="FAT"))   # inicia em Agendamento
    # Etapa 1: definir GCI
    r1 = client.post("/projetos/%s/definir_gci" % pid, data={"gci": "GCI Teste"})
    assert r1.status_code == 302
    with db.Session() as s:
        p = s.get(db.Projeto, pid)
        assert p.gci == "GCI Teste" and p.etapa == "Agendamento"  # ainda em Agendamento
    # Etapa 2: definir data (com GCI já definido)
    r2 = client.post("/projetos/%s/agendar" % pid, data={"data_levantamento": "2026-07-15"})
    assert r2.status_code == 302
    with db.Session() as s:
        p = s.get(db.Projeto, pid)
        assert p.data_levantamento == "2026-07-15" and p.etapa == "Levantamento"
    client.post("/projetos/%s/excluir" % pid)
    with client.session_transaction() as sess:
        sess.clear()


def test_agendamento_sem_gci_nao_acessa_data(client):
    """Testa que a rota de data redireciona para definir_gci se GCI não estiver definido."""
    _login_como(client, "Administrativo")
    pid = int(_novo(client, cliente="Sem GCI Co", modulos="FAT"))
    # Tenta acessar diretamente a rota de data sem GCI definido
    r = client.get("/projetos/%s/agendar" % pid)
    assert r.status_code == 302
    assert b"definir_gci" in r.data or "/definir_gci" in r.headers.get("Location", "")
    client.post("/projetos/%s/excluir" % pid)
    with client.session_transaction() as sess:
        sess.clear()


def test_definir_multiplos_gci(client):
    """O GCI pode ter mais de um responsável no levantamento (checkboxes)."""
    _login_como(client, "Administrativo")
    pid = int(_novo(client, cliente="Multi GCI Co", modulos="FAT"))
    r = client.post("/projetos/%s/definir_gci" % pid, data={"gci": ["GCI Um", "GCI Dois"]})
    assert r.status_code == 302
    with db.Session() as s:
        assert s.get(db.Projeto, pid).gci == "GCI Um, GCI Dois"
    client.post("/projetos/%s/excluir" % pid)
    with client.session_transaction() as sess:
        sess.clear()


def test_designacao_consultores_avanca(client):
    _login_como(client, "GCI")
    pid = int(_novo(client, cliente="Consult Co", etapa="Designação", modulos="FAT, CTB"))
    with db.Session() as s:
        for t in ("levantamento", "projeto"):
            s.add(db.Documento(projeto_id=pid, tipo=t, arquivo=t, caminho=t))
        s.commit()
    r = client.post("/projetos/%s/consultores" % pid, data={"mod_0": "Ana C", "mod_1": "Ana C"})
    assert r.status_code == 302
    with db.Session() as s:
        p = s.get(db.Projeto, pid)
        assert "Ana C" in (p.consultor or "") and p.etapa == "Cronograma e Check-list"
    client.post("/projetos/%s/excluir" % pid)
    with client.session_transaction() as sess:
        sess.clear()


def test_c_nao_avanca_no_levantamento(client):
    pid = int(_novo(client, cliente="Fica Levant", etapa="Levantamento", modulos="FAT"))
    with db.Session() as s:
        s.add(db.Documento(projeto_id=pid, tipo="levantamento", arquivo="l.docx", caminho="l.docx"))
        s.commit()
    A._auto_avancar(pid)   # Levantamento é confirmado pelo humano -> não avança sozinho
    with db.Session() as s:
        assert s.get(db.Projeto, pid).etapa == "Levantamento"
    client.post("/projetos/%s/excluir" % pid)


def test_a_cria_projeto_de_fechamento():
    corpo = ("Cliente (Razão Social): ROBO LTDA\nCNPJ: 99\n"
             "Módulos contratados (siglas): FAT\nHoras cobradas: 20\n")
    pid = A._criar_projeto_de_fechamento(corpo, "[IMPLANTACAO] ROBO")
    with db.Session() as s:
        p = s.get(db.Projeto, pid)
        assert p and "ROBO" in p.cliente
        s.delete(p)
        s.commit()


def test_e_docview_docx(tmp_path):
    from docx import Document
    import docview
    f = tmp_path / "amostra.docx"
    doc = Document()
    doc.add_heading("Projeto de Implantação", level=0)
    doc.add_paragraph("Cliente: ACME")
    t = doc.add_table(rows=2, cols=2)
    t.rows[0].cells[0].text = "Etapa"; t.rows[0].cells[1].text = "Horas"
    t.rows[1].cells[0].text = "Abertura"; t.rows[1].cells[1].text = "2"
    doc.save(str(f))
    h = docview.to_html(str(f))
    assert "Projeto de Implantação" in h
    assert "<table" in h and "Abertura" in h


def test_f_cronograma_seed_edita_e_historia(client):
    pid = int(_novo(client, cliente="Plano PT", modulos="FAT, CTB", horas_cobradas="20"))
    client.post("/projetos/%s/cronograma/seed" % pid)
    itens = db.cronograma_do_projeto(pid)
    assert len(itens) >= 3
    data = {("r_" + c): [(("Concluído" if (c == "status" and i == 0) else it[c]))
                         for i, it in enumerate(itens)] for c in db.CRONO_CAMPOS}
    client.post("/projetos/%s/cronograma" % pid, data=data)
    hist = db.modificacoes_do_projeto(pid, "cronograma")
    assert any("status" in h["campo"] for h in hist)
    assert db.cronograma_do_projeto(pid)[0]["status"] == "Concluído"
    client.post("/projetos/%s/excluir" % pid)


def test_f_checklist_salva(client):
    pid = int(_novo(client, cliente="Check PT", modulos="FAT"))
    data = {"r_modulo": ["FAT", "FAT"], "r_item": ["Cadastro", "Pedido"],
            "r_responsavel": ["Ana", "Bia"], "r_status": ["Pendente", "Concluído"],
            "r_obs": ["", "ok"]}
    client.post("/projetos/%s/checklist" % pid, data=data)
    itens = db.checklist_do_projeto(pid)
    assert len(itens) == 2 and itens[1]["status"] == "Concluído"
    client.post("/projetos/%s/excluir" % pid)


def _achar_usuario(email):
    from sqlalchemy import func
    with db.Session() as s:
        return s.query(db.Usuario).filter(func.lower(db.Usuario.email) == email.lower()).first()


def test_cadastro_valida_codigo():
    db.salvar_pendente("Fulano", "fulano@x.com", "fulano@x.com", "segredo123", "123456")
    u, e = db.confirmar_pendente("fulano@x.com", "000000")
    assert u is None and e            # código errado
    u, e = db.confirmar_pendente("fulano@x.com", "123456")
    assert u and u["perfil"] == "Consultor"
    achado = _achar_usuario("fulano@x.com")
    assert achado and achado.email == "fulano@x.com" and achado.ativo == 1
    with db.Session() as s:
        s.delete(s.get(db.Usuario, achado.id)); s.commit()


def test_cadastro_paginas(client):
    assert client.get("/cadastro").status_code == 200
    r = client.get("/cadastro/confirmar")   # sem sessão -> volta ao cadastro
    assert r.status_code == 302 and "/cadastro" in r.headers["Location"]


def test_cadastro_fluxo_completo(client, monkeypatch):
    import mailer
    monkeypatch.setattr(mailer, "configurado", lambda: True)
    monkeypatch.setattr(mailer, "enviar", lambda dest, asn, corpo: (True, None))
    r = client.post("/cadastro", data={"nome": "Beltrano", "email": "beltrano@x.com",
                                       "senha": "segredo1", "codigo_sicla": "B99"})
    assert r.status_code == 302 and "confirmar" in r.headers["Location"]
    p = db.pendente_por_email("beltrano@x.com")
    assert p and p["codigo"]
    r2 = client.post("/cadastro/confirmar", data={"codigo": p["codigo"]})
    assert r2.status_code == 302
    u = _achar_usuario("beltrano@x.com")
    assert u and u.perfil == "Consultor" and u.codigo_sicla == "B99"
    with db.Session() as s:
        s.delete(s.get(db.Usuario, u.id)); s.commit()


def _login_como(client, perfil):
    with client.session_transaction() as sess:
        sess["auth"] = True
        sess["perfil"] = perfil
        sess["perfil_nome"] = perfil


def test_acesso_por_perfil(client):
    _login_como(client, "Consultor")          # só Operação
    assert client.get("/coordenacao").status_code == 403
    assert client.get("/monitoramento").status_code == 403
    assert client.get("/atividade").status_code == 403
    assert client.get("/usuarios").status_code == 403
    h = client.get("/").get_data(as_text=True)
    assert ">Gestão<" not in h and ">Sistema<" not in h

    for papel in ("GCI", "Administrativo"):     # Operação + Gestão, sem Sistema
        _login_como(client, papel)
        assert client.get("/coordenacao").status_code == 200
        assert client.get("/monitoramento").status_code == 200
        assert client.get("/usuarios").status_code == 403
        h = client.get("/").get_data(as_text=True)
        assert ">Gestão<" in h and ">Sistema<" not in h

    _login_como(client, "ADM")                 # tudo
    assert client.get("/coordenacao").status_code == 200
    assert client.get("/monitoramento").status_code == 200
    assert client.get("/usuarios").status_code == 200
    assert ">Sistema<" in client.get("/").get_data(as_text=True)
    with client.session_transaction() as sess:
        sess.clear()


def test_dedup_fechamento():
    corpo = ("Cliente (Razão Social): DEDUP LTDA\nCNPJ: 11.222.333/0001-44\n"
             "Módulos contratados (siglas): FAT\n")
    pid1 = A._criar_projeto_de_fechamento(corpo, "[IMPLANTACAO] DEDUP")
    pid2 = A._criar_projeto_de_fechamento(corpo, "[IMPLANTACAO] DEDUP")   # mesmo cliente/CNPJ
    assert pid1 == pid2                                                   # não duplicou
    corpo2 = "Cliente (Razão Social): OUTRA RAZAO\nCNPJ: 11222333000144\n"  # mesmo CNPJ, sem máscara
    assert A._criar_projeto_de_fechamento(corpo2, "x") == pid1
    assert db.projeto_existe("qualquer", "11.222.333/0001-44") == pid1
    with db.Session() as s:
        s.delete(s.get(db.Projeto, pid1)); s.commit()


def test_notificar_assincrono_nao_bloqueia(client, monkeypatch):
    """A notificação por e-mail roda em segundo plano: um envio lento NÃO trava a chamada,
    e o evento é gravado na timeline pela thread."""
    import time, threading, mailer
    pid = int(_novo(client, cliente="Async LTDA", modulos="FAT"))
    monkeypatch.setitem(A.app.config, "TESTING", False)        # força o caminho assíncrono
    monkeypatch.setattr(mailer, "configurado", lambda: True)
    monkeypatch.setattr(mailer, "enviar", lambda d, a, c, anexos=None: (time.sleep(1.5) or (True, None)))
    assunto = "AsyncTest-%d" % time.time()                     # assunto único (evita colisão por reuso de id)
    t0 = time.time()
    A._notificar(pid, ["x@y.com"], assunto, "Corpo")
    assert time.time() - t0 < 0.8                              # retornou sem esperar o envio (1.5s)
    for t in threading.enumerate():                            # aguarda a thread concluir
        if t.name == "notificar":
            t.join(5)
    with db.Session() as s:
        n = s.query(db.Evento).filter(db.Evento.projeto_id == pid, db.Evento.tipo == "email",
                                      db.Evento.descricao.like("%" + assunto + "%")).count()
    assert n == 1
    monkeypatch.setitem(A.app.config, "TESTING", True)
    client.post("/projetos/%s/excluir" % pid)


def test_usuarios_grava_email_e_perfil(client):
    client.post("/usuarios", data={"nome": "Coord", "email": "coord@x.com", "perfil": "Coordenador",
                                   "codigo_sicla": "C77", "ativo": "on"})
    u = _achar_usuario("coord@x.com")
    assert u and u.login == "coord@x.com" and u.perfil == "Coordenador" and u.codigo_sicla == "C77"
    with db.Session() as s:
        s.delete(s.get(db.Usuario, u.id)); s.commit()


def test_usuarios_exige_codigo_sicla(client):
    """Código SICLA é obrigatório — sem ele, o usuário não é gravado."""
    r = client.post("/usuarios", data={"nome": "SemCod", "email": "semcod@x.com", "perfil": "Consultor", "ativo": "on"})
    assert r.status_code == 200 and "obrigat" in r.get_data(as_text=True).lower()
    assert _achar_usuario("semcod@x.com") is None


# ---- Cadastros de referência: Checklist e Índice de Tópicos ----

def test_cadastros_seed_e_paginas(client):
    db.init_db()  # idempotente: garante o seed dos catálogos
    _, ck = db.checklist_modelo_listar()
    _, idx = db.indice_listar()
    assert ck > 100                       # checklist do modelo entregue
    assert idx >= 351                     # índice da planilha (351 tópicos)
    assert len(db.indice_modulos()) >= 23
    assert client.get("/cadastros/checklist").status_code == 200
    assert client.get("/cadastros/indice").status_code == 200


def test_cadastro_indice_crud(client):
    _, antes = db.indice_listar()
    r = client.post("/cadastros/indice/salvar", data={
        "modulo_num": "99", "modulo_sigla": "TST", "modulo": "Módulo Teste",
        "adicional": "", "topico": "Tópico PYTEST único"})
    assert r.status_code == 302
    achados, _ = db.indice_listar(q="Tópico PYTEST único")
    assert len(achados) == 1
    tid = achados[0]["id"]
    client.post("/cadastros/indice/salvar", data={"id": tid, "modulo_sigla": "TST",
        "modulo": "Módulo Teste", "topico": "Tópico PYTEST editado"})
    assert db.indice_listar(q="editado")[1] >= 1
    client.post("/cadastros/indice/%s/excluir" % tid)
    _, depois = db.indice_listar()
    assert depois == antes


def test_cadastro_checklist_crud(client):
    r = client.post("/cadastros/checklist/salvar", data={
        "modulo": "ZZZ", "adicional": "ZZZ", "tipo": "Teste",
        "item": "Item PYTEST", "menu": "9.9", "golive": "Sim", "seq": "1"})
    assert r.status_code == 302
    add, _ = db.checklist_modelo_listar(q="Item PYTEST")
    assert len(add) == 1
    client.post("/cadastros/checklist/%s/excluir" % add[0]["id"])
    assert db.checklist_modelo_listar(q="Item PYTEST")[1] == 0


def test_cadastros_bloqueio_por_perfil(client):
    _login_como(client, "Consultor")
    assert client.get("/cadastros/checklist").status_code == 403
    assert client.get("/cadastros/indice").status_code == 403
    with client.session_transaction() as sess:
        sess.clear()


# ---- Cadastro de Modelos de Documentos (layouts fiéis das fases) ----

def test_modelos_documento_seed(client):
    db.init_db()
    mods = db.modelos_documento_listar()
    assert {m["slug"] for m in mods} == {"levantamento", "projeto", "cronograma", "termo"}
    for m in mods:
        assert m["n_versoes"] >= 1 and m["n_campos"] > 0
        p = db.modelo_documento_arquivo_path(m["id"])
        assert p and os.path.exists(p)        # layout fiel copiado como v1
    assert client.get("/cadastros/modelos").status_code == 200


def test_modelo_detalhe_e_download(client):
    mid = db.modelos_documento_listar()[0]["id"]
    assert client.get("/cadastros/modelos/%s" % mid).status_code == 200
    r = client.get("/cadastros/modelos/%s/baixar" % mid)
    assert r.status_code == 200 and len(r.data) > 1000


def test_modelo_campo_crud(client):
    mid = db.modelos_documento_listar()[0]["id"]
    n0 = len(db.modelo_documento_campos(mid))
    client.post("/cadastros/modelos/%s/campo/salvar" % mid, data={
        "secao": "X", "rotulo": "Campo PYTEST", "placeholder": "<P>",
        "origem": "manual", "obrigatorio": "on"})
    novos = [c for c in db.modelo_documento_campos(mid) if c["rotulo"] == "Campo PYTEST"]
    assert len(novos) == 1 and novos[0]["obrigatorio"] == 1
    client.post("/cadastros/modelos/%s/campo/%s/excluir" % (mid, novos[0]["id"]))
    assert len(db.modelo_documento_campos(mid)) == n0


def test_modelos_bloqueio_por_perfil(client):
    _login_como(client, "Consultor")
    assert client.get("/cadastros/modelos").status_code == 403
    with client.session_transaction() as sess:
        sess.clear()


def test_gerar_layout_fiel(client):
    pid = _novo(client, cliente="Cliente Layout LTDA", cnpj="11.222.333/0001-44",
                numero_projeto="PRJ-99", horas_cobradas="80", horas_bonificadas="10")
    r = client.post("/projetos/%s/gerar-layout/termo" % pid)
    assert r.status_code == 302
    with db.Session() as s:
        docs = s.query(db.Documento).filter_by(projeto_id=int(pid), tipo="termo").all()
        assert len(docs) == 1
        path = docs[0].caminho
    assert path and os.path.exists(path)
    from docx import Document
    txt = "\n".join(x.text for x in Document(path).paragraphs)
    assert "Cliente Layout LTDA" in txt          # cliente preenchido
    assert "<Razão Social Longa>" not in txt      # placeholder trocado
    assert "PRJ-99" in txt                         # número do projeto
    try:
        os.remove(path)
    except OSError:
        pass
    client.post("/projetos/%s/excluir" % pid)


def test_gerar_layout_slug_invalido(client):
    pid = _novo(client, cliente="X LTDA")
    assert client.post("/projetos/%s/gerar-layout/inexistente" % pid).status_code == 404
    client.post("/projetos/%s/excluir" % pid)


def test_ficha_caminho_unico_andamento(client):
    """A ficha abre no 'Andamento' (caminho único) e não tem mais o painel paralelo de docs."""
    pid = _novo(client, cliente="Caminho LTDA", modulos="FAT")
    h = client.get("/projetos/%s" % pid).get_data(as_text=True)
    assert "Andamento" in h                                 # aba do fluxo
    assert "Documentos oficiais (layout fiel)" not in h      # painel paralelo removido
    assert 'id="t-execucao" class="painel-aba ativa' in h    # Andamento é a aba ativa por padrão
    client.post("/projetos/%s/excluir" % pid)


def test_home_fila_proximas_acoes(client):
    """A Home lista 'Minhas próximas ações' com link direto para a ação certa."""
    import datetime
    passado = (datetime.date.today() - datetime.timedelta(days=5)).isoformat()
    pid = _novo(client, cliente="Fila LTDA ZZZ", numero_projeto="F-1", modulos="FAT",
                cnpj="00.000.000/0001-00", horas_cobradas="10", data_uso_oficial=passado)
    h = client.get("/").get_data(as_text=True)
    assert "Fila LTDA ZZZ" in h                          # projeto aparece na fila
    assert ("/projetos/%s/definir_gci" % pid) in h        # botão leva direto à ação (Definir GCI)
    assert "Minhas próximas ações" in h
    client.post("/projetos/%s/excluir" % pid)


def test_levantamento_respostas_no_painel(client):
    """As perguntas do Índice viram campos respondíveis (seed) e as respostas persistem."""
    sig = db.indice_modulos()[0]["sigla"]
    _linhas, total_idx = db.indice_listar(modulo=sig)
    pid = _novo(client, cliente="Resp LTDA", modulos=sig)
    r = client.get("/projetos/%s/levantamento" % pid)
    assert r.status_code == 200
    rs = db.levantamento_respostas(pid)
    assert len(rs) == total_idx              # uma pergunta por tópico do módulo contratado
    rid = rs[0]["id"]
    client.post("/projetos/%s/levantamento" % pid, data={"resposta_%d" % rid: "Resposta teste X"})
    achou = [x for x in db.levantamento_respostas(pid) if x["id"] == rid][0]
    assert achou["resposta"] == "Resposta teste X"
    resp, total = db.levantamento_resumo(pid)
    assert resp == 1 and total == total_idx
    client.post("/projetos/%s/excluir" % pid)


def test_levantamento_bloqueio_perfil(client):
    pid = _novo(client, cliente="Bloq LTDA", modulos="FAT")
    _login_como(client, "Consultor")
    assert client.get("/projetos/%s/levantamento" % pid).status_code == 403
    with client.session_transaction() as sess:
        sess.clear()
    client.post("/projetos/%s/excluir" % pid)


def test_levantamento_tabelas_e_sem_marcadores(client):
    """Levantamento: tabelas de módulos e horas preenchidas e SEM marcadores <...>."""
    import gerar_layout
    import re as _re
    pid = _novo(client, cliente="Marc LTDA", numero_projeto="M-1", modulos="FAT",
                horas_cobradas="120", horas_bonificadas="20")
    with db.Session() as s:
        proj = db.to_dict(s.get(db.Projeto, int(pid)))
    path = gerar_layout.gerar("levantamento", proj)
    from docx import Document
    d = Document(path)
    txt = "\n".join(p.text for p in d.paragraphs)
    blob = " | ".join(c.text for t in d.tables for row in t.rows for c in row.cells)
    assert "120" in blob and "140" in blob       # horas: cobradas + total (120+20)
    assert "FAT" in blob                          # tabela 'Módulos/Adicionais (A)'
    assert not _re.search(r"<[^>]*>", txt + blob)  # nenhum marcador <...> restante
    try:
        os.remove(path)
    except OSError:
        pass
    client.post("/projetos/%s/excluir" % pid)


def test_doc_editar_levantamento_alimenta(client):
    """A tela de edição estruturada do Levantamento persiste e a geração consome."""
    import gerar_layout
    pid = _novo(client, cliente="Editar LTDA", numero_projeto="ED-1", cnpj="00.000.000/0001-00",
                modulos="FAT", horas_cobradas="10")
    r = client.get("/projetos/%s/editar/levantamento" % pid)
    assert r.status_code == 200 and "Identificação da empresa" in r.get_data(as_text=True)
    client.post("/projetos/%s/editar/levantamento" % pid,
                data={"ramo": "Metalurgia", "produto": "Software ERP",
                      "software_atual": "Sistema X", "filiais": "Matriz + 2 filiais",
                      "objetivos": "Modernizar", "qtd_usuarios": "30 usuários"})
    assert db.doc_conteudo(int(pid), "levantamento").get("produto") == "Software ERP"
    with db.Session() as s:
        proj = db.to_dict(s.get(db.Projeto, int(pid)))
    path = gerar_layout.gerar("levantamento", proj)
    from docx import Document
    txt = "\n".join(p.text for p in Document(path).paragraphs)
    assert "Software ERP" in txt and "Matriz + 2 filiais" in txt   # consumido no .docx
    try:
        os.remove(path)
    except OSError:
        pass
    client.post("/projetos/%s/excluir" % pid)


def test_doc_editar_projeto_alimenta(client):
    """A tela de edição estruturada do Projeto persiste e a geração consome."""
    import gerar_layout
    pid = _novo(client, cliente="Edit Proj LTDA", numero_projeto="EP-1",
                cnpj="11.111.111/0001-11", modulos="FAT", horas_cobradas="10")
    client.post("/projetos/%s/editar/projeto" % pid,
                data={"cnpj": "11.111.111/0001-11", "objetivos": "OBJETIVO EDITADO XYZ",
                      "redator": "Maria Redatora", "gerente_contas": "Beto GCI"})
    with db.Session() as s:
        proj = db.to_dict(s.get(db.Projeto, int(pid)))
    path = gerar_layout.gerar("projeto", proj)
    from docx import Document
    txt = "\n".join(p.text for p in Document(path).paragraphs)
    assert "OBJETIVO EDITADO XYZ" in txt and "Maria Redatora" in txt
    try:
        os.remove(path)
    except OSError:
        pass
    client.post("/projetos/%s/excluir" % pid)


def test_doc_editar_projeto_detalhamento_area(client):
    """O Detalhamento de Rotinas por área (edição) aparece no Projeto gerado."""
    import gerar_layout
    pid = _novo(client, cliente="Det LTDA", numero_projeto="DT-1", cnpj="00.000.000/0001-00",
                modulos="FAT", horas_cobradas="10")
    h = client.get("/projetos/%s/editar/projeto" % pid).get_data(as_text=True)
    assert "Detalhamento de Rotinas — Vendas e Faturamento" in h   # seção dinâmica (FAT->Vendas)
    client.post("/projetos/%s/editar/projeto" % pid,
                data={"det_vendas_modulos": "FAT", "det_vendas_detalhamento": "DETALHE VENDAS XYZ"})
    with db.Session() as s:
        proj = db.to_dict(s.get(db.Projeto, int(pid)))
    path = gerar_layout.gerar("projeto", proj)
    from docx import Document
    txt = "\n".join(p.text for p in Document(path).paragraphs)
    assert "DETALHE VENDAS XYZ" in txt               # consumido no bloco da área no .docx
    try:
        os.remove(path)
    except OSError:
        pass
    client.post("/projetos/%s/excluir" % pid)


def test_doc_editar_projeto_tabelas(client):
    """Tabela de Usuários e Cronograma Macro (edição) aparecem no Projeto gerado."""
    import gerar_layout
    pid = _novo(client, cliente="Tab LTDA", numero_projeto="TB-1", cnpj="00.000.000/0001-00",
                modulos="FAT", horas_cobradas="10")
    client.post("/projetos/%s/editar/projeto" % pid, data={
        "usu_0_nome": "Joao Usuario", "usu_0_email": "joao@x.com",
        "usu_0_area": "Vendas", "usu_0_assina": "Sim",
        "crono_treinamento": "PERIODO TREINO 99"})
    with db.Session() as s:
        proj = db.to_dict(s.get(db.Projeto, int(pid)))
    path = gerar_layout.gerar("projeto", proj)
    from docx import Document
    blob = " | ".join(c.text for t in Document(path).tables for row in t.rows for c in row.cells)
    assert "Joao Usuario" in blob and "joao@x.com" in blob   # tabela de usuários
    assert "PERIODO TREINO 99" in blob                        # cronograma macro
    try:
        os.remove(path)
    except OSError:
        pass
    client.post("/projetos/%s/excluir" % pid)


def test_doc_editar_levantamento_usuarios(client):
    """A tabela de Usuários-chave (edição) aparece no Levantamento gerado."""
    import gerar_layout
    pid = _novo(client, cliente="LevUsu LTDA", numero_projeto="LU-1", modulos="FAT")
    client.post("/projetos/%s/editar/levantamento" % pid, data={
        "usu_0_nome": "Ana Chave", "usu_0_email": "ana@x.com", "usu_0_atrib": "Vendas"})
    with db.Session() as s:
        proj = db.to_dict(s.get(db.Projeto, int(pid)))
    path = gerar_layout.gerar("levantamento", proj)
    from docx import Document
    blob = " | ".join(c.text for t in Document(path).tables for row in t.rows for c in row.cells)
    assert "Ana Chave" in blob and "ana@x.com" in blob
    try:
        os.remove(path)
    except OSError:
        pass
    client.post("/projetos/%s/excluir" % pid)


def test_levantamento_tela_por_area(client):
    """A tela de perguntas do Levantamento agrupa por bloco de ÁREA (como no documento)."""
    pid = _novo(client, cliente="Area LTDA", modulos="FAT")
    h = client.get("/projetos/%s/levantamento" % pid).get_data(as_text=True)
    assert "Vendas e Faturamento" in h        # FAT -> área Vendas e Faturamento
    client.post("/projetos/%s/excluir" % pid)


def test_doc_editar_bloqueio_perfil(client):
    pid = _novo(client, cliente="Bloq Edit LTDA", modulos="FAT")
    _login_como(client, "Consultor")
    assert client.get("/projetos/%s/editar/levantamento" % pid).status_code == 403
    with client.session_transaction() as sess:
        sess.clear()
    client.post("/projetos/%s/excluir" % pid)


def test_termo_grade_com_modulos(client):
    """O Termo gerado preenche a grade Resumo Geral com os módulos contratados."""
    import gerar_layout
    pid = _novo(client, cliente="Termo Grade LTDA", numero_projeto="T-1", modulos="FAT, EST")
    with db.Session() as s:
        proj = db.to_dict(s.get(db.Projeto, int(pid)))
    path = gerar_layout.gerar("termo", proj)
    from docx import Document
    celulas = [c.text.strip() for row in Document(path).tables[0].rows for c in row.cells]
    assert "FAT" in celulas and "EST" in celulas and "Implantado" in celulas
    try:
        os.remove(path)
    except OSError:
        pass
    client.post("/projetos/%s/excluir" % pid)


def test_cronograma_xlsx_cabecalho_e_linhas(client):
    """O Cronograma .xlsx preenche cabeçalho (consultor/horas) e linhas de visita."""
    import gerar_layout
    import openpyxl
    pid = _novo(client, cliente="Crono LTDA", numero_projeto="C-1", modulos="FAT",
                consultor="Ana", horas_cobradas="40", horas_bonificadas="8")
    db.salvar_linhas(int(pid), "cronograma", [
        {"etapa": "Treinamento Vendas", "topicos": "Pedidos", "horas": "4",
         "data": "01/08/2026", "modalidade": "Presencial", "status": "Previsto"}], "tester")
    with db.Session() as s:
        proj = db.to_dict(s.get(db.Projeto, int(pid)))
    path = gerar_layout.gerar("cronograma", proj)
    wb = openpyxl.load_workbook(path)
    blob = " | ".join(str(c.value) for ws in wb.worksheets for row in ws.iter_rows()
                      for c in row if c.value is not None)
    assert "Ana" in blob                     # consultor no cabeçalho/técnico
    assert "Crono LTDA" in blob              # cliente
    assert "Treinamento Vendas" in blob      # linha de visita (O que será abordado)
    try:
        os.remove(path)
    except OSError:
        pass
    client.post("/projetos/%s/excluir" % pid)


def test_projeto_exige_levantamento_realizado(client):
    """O Projeto só é gerado a partir de um Levantamento realizado (respostas preenchidas)."""
    pid = _novo(client, cliente="Gate Lev LTDA", numero_projeto="G-1", cnpj="00.000.000/0001-00",
                horas_cobradas="10", modulos="FAT", etapa="Projeto")
    db.levantamento_seed(int(pid), "FAT")
    r = client.post("/projetos/%s/gerar/projeto" % pid)      # sem responder -> vai ao GATE de origem
    assert r.status_code == 302 and "/projeto/origem" in r.headers["Location"]
    with db.Session() as s:
        assert s.query(db.Documento).filter_by(projeto_id=int(pid), tipo="projeto").count() == 0
    rs = db.levantamento_respostas(int(pid))
    client.post("/projetos/%s/levantamento" % pid, data={"resposta_%d" % rs[0]["id"]: "ok"})
    client.post("/projetos/%s/projeto/origem" % pid, data={"fonte": "tela"})   # usa dados da tela -> gera
    with db.Session() as s:
        docs = s.query(db.Documento).filter_by(projeto_id=int(pid), tipo="projeto").all()
        assert len(docs) == 1
        path = docs[0].caminho
    try:
        os.remove(path)
    except OSError:
        pass
    client.post("/projetos/%s/excluir" % pid)


def test_projeto_puxa_respostas_do_levantamento(client):
    """O Projeto gerado consome as respostas do Levantamento (liga as fases)."""
    import gerar_layout
    sig = "FAT"                                          # mapeia para a área Vendas e Faturamento
    pid = _novo(client, cliente="Liga LTDA", cnpj="00.000.000/0001-00", numero_projeto="L-1",
                modulos=sig, horas_cobradas="10")
    db.levantamento_seed(int(pid), sig)
    rs = db.levantamento_respostas(int(pid))
    rid, topico = rs[0]["id"], rs[0]["topico"]
    client.post("/projetos/%s/levantamento" % pid,
                data={"resposta_%d" % rid: "RESPOSTA QUE LIGA AO PROJETO"})
    with db.Session() as s:
        proj = db.to_dict(s.get(db.Projeto, int(pid)))
    path = gerar_layout.gerar("projeto", proj)
    from docx import Document
    txt = "\n".join(p.text for p in Document(path).paragraphs)
    # a resposta aparece no Detalhamento das Rotinas (não há mais bloco duplicado após a assinatura)
    assert "RESPOSTA QUE LIGA AO PROJETO" in txt
    assert topico[:25] in txt
    assert "Detalhamento do Levantamento por módulo" not in txt   # bloco redundante removido
    try:
        os.remove(path)
    except OSError:
        pass
    client.post("/projetos/%s/excluir" % pid)


def test_levantamento_inclui_topicos_do_indice(client):
    """O Levantamento gerado injeta as perguntas/tópicos do Índice por módulo contratado."""
    import gerar_layout
    mods = db.indice_modulos()
    assert mods, "Índice de Tópicos deve estar seedado"
    sig = mods[0]["sigla"]
    linhas, _ = db.indice_listar(modulo=sig)
    assert linhas
    topico = linhas[0]["topico"]
    pid = _novo(client, cliente="Topicos LTDA", modulos=sig, gci="Beto")
    with db.Session() as s:
        proj = db.to_dict(s.get(db.Projeto, int(pid)))
    path = gerar_layout.gerar("levantamento", proj)
    from docx import Document
    txt = "\n".join(p.text for p in Document(path).paragraphs)
    assert sig in txt                  # módulo aparece (Módulos Previstos / seção)
    assert topico[:30] in txt          # a pergunta do Índice foi injetada no documento
    try:
        os.remove(path)
    except OSError:
        pass
    client.post("/projetos/%s/excluir" % pid)


def test_levantamento_so_blocos_contratados(client):
    """O Levantamento mantém só os blocos dos módulos contratados e injeta as respostas."""
    import gerar_layout
    pid = _novo(client, cliente="Blocos LTDA", numero_projeto="B-1", modulos="FAT")
    db.levantamento_seed(int(pid), "FAT")
    rs = db.levantamento_respostas(int(pid))
    client.post("/projetos/%s/levantamento" % pid, data={"resposta_%d" % rs[0]["id"]: "RESP BLOCO"})
    with db.Session() as s:
        proj = db.to_dict(s.get(db.Projeto, int(pid)))
    import re as _re
    path = gerar_layout.gerar("levantamento", proj)
    from docx import Document
    d = Document(path)
    heads = [p.text.strip().upper() for p in d.paragraphs
             if p.text.strip().lower().startswith("mapeamento de processo")
             and _re.search(r"[-–—]", p.text)]
    # só sobram os blocos contratados (FAT) + fundacionais (CLIENTE/FORNECEDOR, PRODUTO)
    assert len(heads) == 3
    junto = " | ".join(heads)
    assert "VENDAS E FATURAMENTO" in junto and "CLIENTE" in junto and "PRODUTO" in junto
    assert "FOLHA" not in junto                  # bloco da folha removido
    txt = "\n".join(p.text for p in d.paragraphs)
    assert "RESP BLOCO" in txt                   # resposta injetada no bloco
    try:
        os.remove(path)
    except OSError:
        pass
    client.post("/projetos/%s/excluir" % pid)


def test_projeto_avanca_sem_consultor(client):
    """Fim do deadlock: sair do Projeto não exige mais 'consultor' (definido na Designação)."""
    pid = _novo(client, cliente="SemConsultor LTDA", numero_projeto="S-1", modulos="FAT",
                cnpj="00.000.000/0001-00", horas_cobradas="10", gci="Beto")
    with db.Session() as s:
        p = s.get(db.Projeto, int(pid)); p.etapa = "Projeto"; s.commit()
        s.add(db.Documento(projeto_id=int(pid), tipo="levantamento", arquivo="l", caminho="l"))
        s.add(db.Documento(projeto_id=int(pid), tipo="projeto", arquivo="p", caminho="p"))
        s.commit()
    client.post("/projetos/%s/avancar" % pid)
    with db.Session() as s:
        assert s.get(db.Projeto, int(pid)).etapa == "Designação"
    client.post("/projetos/%s/excluir" % pid)


def test_fluxo_e2e_continuidade(client):
    """Robô de fluxo ponta-a-ponta: percorre as 6 fases validando gates e avanços
    (Agendamento → Levantamento → Projeto → Designação → Cronograma/Check-list → Encerramento)."""
    import datetime
    futuro = (datetime.date.today() + datetime.timedelta(days=7)).isoformat()
    pid = _novo(client, cliente="E2E Fluxo LTDA", cnpj="00.000.000/0001-00",
                numero_projeto="E2E-1", modulos="FAT, EST", horas_cobradas="40",
                data_uso_oficial=futuro)

    def etapa():
        with db.Session() as s:
            return s.get(db.Projeto, int(pid)).etapa

    def n_doc(tipo):
        with db.Session() as s:
            return s.query(db.Documento).filter_by(projeto_id=int(pid), tipo=tipo).count()

    assert etapa() == "Agendamento"
    # Agendamento: GCI (sub-etapa 1) + Data (sub-etapa 2) -> auto-avança p/ Levantamento
    client.post("/projetos/%s/definir_gci" % pid, data={"gci": "Beto"})
    assert etapa() == "Agendamento"
    client.post("/projetos/%s/agendar" % pid, data={"data_levantamento": futuro})
    assert etapa() == "Levantamento"

    # Levantamento: responde (base do Projeto), gera Mapeamento e confirma avanço
    db.levantamento_seed(int(pid), "FAT, EST")
    _rs = db.levantamento_respostas(int(pid))
    client.post("/projetos/%s/levantamento" % pid, data={"resposta_%d" % _rs[0]["id"]: "ok"})
    client.post("/projetos/%s/gerar/levantamento" % pid)
    assert n_doc("levantamento") == 1
    client.post("/projetos/%s/avancar" % pid)
    assert etapa() == "Projeto"

    # Projeto: gerado A PARTIR do Levantamento realizado (via gate, dados da tela) -> auto-avança
    client.post("/projetos/%s/projeto/origem" % pid, data={"fonte": "tela"})
    assert n_doc("projeto") == 1
    assert etapa() == "Designação"

    # Designação: designa SÓ consultores por módulo (GCI já foi no Agendamento)
    client.post("/projetos/%s/consultores" % pid, data={"mod_0": "Ana", "mod_1": "Ana"})
    client.post("/projetos/%s/avancar" % pid)
    assert etapa() == "Cronograma e Check-list"

    # Cronograma e Check-list: gera Cronograma (fiel) + anexa Check List -> avança p/ Encerramento
    client.post("/projetos/%s/gerar/cronograma" % pid)
    assert n_doc("cronograma") == 1
    with db.Session() as s:   # não há gerador fiel de checklist; anexa p/ satisfazer o gate
        s.add(db.Documento(projeto_id=int(pid), tipo="checklist", arquivo="cl.xlsx", caminho="cl.xlsx"))
        s.commit()
    client.post("/projetos/%s/avancar" % pid)
    assert etapa() == "Encerramento"

    # Encerramento: gera Termo
    client.post("/projetos/%s/gerar/termo" % pid)
    assert n_doc("termo") == 1

    # limpeza (remove arquivos gerados + projeto)
    with db.Session() as s:
        caminhos = [d.caminho for d in s.query(db.Documento).filter_by(projeto_id=int(pid)).all()]
    for cp in caminhos:
        try:
            if cp and os.path.exists(cp):
                os.remove(cp)
        except OSError:
            pass
    client.post("/projetos/%s/excluir" % pid)


def test_projeto_gerar_usa_layout_fiel(client):
    """A rota de geração da fase (projeto_gerar) agora produz o layout FIEL preenchido."""
    pid = _novo(client, cliente="Faithful Proj LTDA", cnpj="22.333.444/0001-55",
                horas_cobradas="50", etapa="Projeto")
    r = client.post("/projetos/%s/projeto/origem" % pid, data={"fonte": "tela"})
    assert r.status_code == 302
    with db.Session() as s:
        docs = s.query(db.Documento).filter_by(projeto_id=int(pid), tipo="projeto").all()
        assert len(docs) == 1
        path = docs[0].caminho
    from docx import Document
    txt = "\n".join(x.text for x in Document(path).paragraphs)
    assert "Faithful Proj LTDA" in txt and "22.333.444/0001-55" in txt
    try:
        os.remove(path)
    except OSError:
        pass
    client.post("/projetos/%s/excluir" % pid)


def test_rota_antiga_gerar_projeto_delega(client):
    """A rota antiga /gerar_projeto foi aposentada e delega ao GATE de origem do Projeto."""
    pid = _novo(client, cliente="Delega LTDA", etapa="Projeto")
    r = client.post("/projetos/%s/gerar_projeto" % pid)   # delega ao gate
    assert r.status_code == 302 and "/projeto/origem" in r.headers["Location"]
    client.post("/projetos/%s/excluir" % pid)


def test_projeto_origem_gate_modelo_e_importacao(client):
    """Gate de origem: GET mostra opções; 'modelo' gera p/ preenchimento manual;
    importar um .docx popula respostas e gera o Projeto fiel."""
    pid = _novo(client, cliente="Gate Origem LTDA", cnpj="00.000.000/0001-00",
                numero_projeto="GO-1", horas_cobradas="10", modulos="FAT", etapa="Projeto")
    db.levantamento_seed(int(pid), "FAT")
    # GET mostra a tela do gate
    r = client.get("/projetos/%s/projeto/origem" % pid)
    assert r.status_code == 200 and "Gerar Projeto" in r.get_data(as_text=True)
    # 'modelo' (preenchimento manual) gera mesmo sem respostas
    client.post("/projetos/%s/projeto/origem" % pid, data={"fonte": "modelo"})
    with db.Session() as s:
        docs = s.query(db.Documento).filter_by(projeto_id=int(pid), tipo="projeto").all()
        assert len(docs) == 1
        for d in docs:
            try:
                os.remove(d.caminho)
            except OSError:
                pass
    # Importar um Levantamento .docx -> popula respostas e gera
    rs = db.levantamento_respostas(int(pid))
    topico = rs[0]["topico"]
    import io
    from docx import Document as _Docx
    doc = _Docx()
    doc.add_paragraph("•  %s: RESPOSTA IMPORTADA" % topico)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    r = client.post("/projetos/%s/projeto/origem" % pid,
                    data={"fonte": "importar", "arquivo": (buf, "mapa.docx")},
                    content_type="multipart/form-data")
    assert r.status_code == 302
    assert any((x["resposta"] or "").strip() == "RESPOSTA IMPORTADA"
               for x in db.levantamento_respostas(int(pid)))
    assert db.levantamento_importado(int(pid)) is not None
    with db.Session() as s:
        for d in s.query(db.Documento).filter_by(projeto_id=int(pid)).all():
            try:
                if d.caminho and os.path.exists(d.caminho):
                    os.remove(d.caminho)
            except OSError:
                pass
    client.post("/projetos/%s/excluir" % pid)


def test_projeto_detalhamento_uma_linha_por_topico(client):
    """Detalhamento das Rotinas do Projeto: cada tópico do Levantamento vira UMA linha
    (não emendado numa única com '·'), respeitando a formatação do modelo."""
    import gerar_layout
    pid = _novo(client, cliente="ListaDet LTDA", cnpj="00.000.000/0001-00", numero_projeto="LD-1",
                modulos="FAT", horas_cobradas="10")
    db.levantamento_seed(int(pid), "FAT")
    rs = db.levantamento_respostas(int(pid))
    assert len(rs) >= 2
    client.post("/projetos/%s/levantamento" % pid, data={
        "resposta_%d" % rs[0]["id"]: "AAA_PRIMEIRO",
        "resposta_%d" % rs[1]["id"]: "BBB_SEGUNDO"})
    with db.Session() as s:
        proj = db.to_dict(s.get(db.Projeto, int(pid)))
    path = gerar_layout.gerar("projeto", proj)
    from docx import Document
    paras = [p.text for p in Document(path).paragraphs]
    # ambas aparecem no Detalhamento das Rotinas, cada uma na SUA linha (não emendadas com '·')
    assert any("AAA_PRIMEIRO" in t for t in paras)
    assert any("BBB_SEGUNDO" in t for t in paras)
    assert not any(("AAA_PRIMEIRO" in t and "BBB_SEGUNDO" in t) for t in paras)
    try:
        os.remove(path)
    except OSError:
        pass
    client.post("/projetos/%s/excluir" % pid)


# ---- Protocolos de treinamento (vídeo -> transcrição -> IA -> revisão) ----

def _video_fake(tmp_path, nome="treino.mp4", conteudo=b"VIDEOFAKE"):
    p = tmp_path / nome
    p.write_bytes(conteudo)
    return str(p)


def test_protocolo_criar_dedup_e_decisao(client, tmp_path):
    """Cria o registro do vídeo, deduplica pelo hash e aprova/reprova com histórico."""
    v = _video_fake(tmp_path)
    pid, novo = db.protocolo_criar("treino.mp4", v, "sharepoint", "Tester")
    assert novo is True
    pid2, novo2 = db.protocolo_criar("treino.mp4", v, "sharepoint", "Tester")
    assert pid2 == pid and novo2 is False                 # dedup por hash
    assert db.protocolo_atualizar_status(pid, "Em revisão", autor="Tester")
    assert db.protocolo_decidir(pid, True, "Chefe")
    p = db.protocolo_get(pid)
    assert p["status"] == "Aprovado" and p["aprovador"] == "Chefe" and p["aprovado_em"]
    assert "APROVADO" in p["historico"]
    with db.Session() as s:
        s.delete(s.get(db.Protocolo, pid)); s.commit()


def test_protocolo_pipeline_mock(client, tmp_path, monkeypatch):
    """Pipeline completo (mocks): transcreve -> analisa -> Em revisão, campos preenchidos."""
    import protocolos as P
    import transcritor, protocolo_ia
    v = _video_fake(tmp_path, "rotina_fiscal.mp4")
    monkeypatch.setattr(transcritor, "transcrever_isolado",
                        lambda path, timeout=0, progress_file=None:
                        {"texto": "[0:05] Configurando a regra fiscal.",
                         "duracao": 65, "idioma": "pt"})
    monkeypatch.setattr(protocolo_ia, "analisar", lambda t, n="": (
        {"titulo": "Configuração de Regra Fiscal", "modulo": "Fiscal", "menu": "2.6-R",
         "assunto": "Regras de tributação", "resumo": "Mostra a configuração.",
         "objetivo": "Configurar regra.", "quando_utilizar": "Na implantação.",
         "pre_requisitos": "- Cadastro de produtos", "passo_a_passo": "1. Acessar o menu 2.6-R",
         "configuracoes": "- Regra padrão", "dependencias": "- Fiscal x Estoque",
         "regras_negocio": "- NCM define aliquota", "pontos_atencao": "- Conferir CFOP",
         "exemplos": "Informação não detalhada no vídeo",
         "assuntos_removidos": "- Conversa sobre agenda interna",
         "pendencias": "- Menu citado rapidamente"}, "{json bruto}"))
    pid, _ = db.protocolo_criar("rotina_fiscal.mp4", v, "upload", "Tester")
    ok, msg = P.processar(pid, "Tester")
    assert ok, msg
    p = db.protocolo_get(pid)
    assert p["status"] == "Em revisão" and p["modulo"] == "Fiscal" and p["menu"] == "2.6-R"
    assert p["duracao_seg"] == 65 and "regra fiscal" in p["transcricao"].lower()
    assert "agenda interna" in p["assuntos_removidos"]     # auditoria do removido
    # upload NÃO move o arquivo (segue no lugar)
    assert os.path.exists(v)
    with db.Session() as s:
        s.delete(s.get(db.Protocolo, pid)); s.commit()


def test_protocolo_pipeline_erro(client, tmp_path, monkeypatch):
    """Falha na transcrição -> status Erro com log; vídeo de upload não é movido."""
    import protocolos as P
    import transcritor
    v = _video_fake(tmp_path, "corrompido.mp4")
    monkeypatch.setattr(transcritor, "transcrever_isolado",
                        lambda path, timeout=0, progress_file=None:
                        (_ for _ in ()).throw(RuntimeError("audio ilegivel")))
    pid, _ = db.protocolo_criar("corrompido.mp4", v, "upload", "Tester")
    ok, _msg = P.processar(pid, "Tester")
    assert ok is False
    p = db.protocolo_get(pid)
    assert p["status"] == "Erro" and "audio ilegivel" in p["log_erro"]
    with db.Session() as s:
        s.delete(s.get(db.Protocolo, pid)); s.commit()


def test_protocolo_varredura_pasta(client, tmp_path, monkeypatch):
    """Robô: varre 'Videos Pendentes', registra vídeos novos e ignora repetidos/extensões."""
    import protocolos as P
    raiz = tmp_path / "Treinamentos"
    (raiz / "Videos Pendentes").mkdir(parents=True)
    (raiz / "Videos Pendentes" / "aula1.mp4").write_bytes(b"AULA1")
    (raiz / "Videos Pendentes" / "notas.txt").write_text("nao é video")
    monkeypatch.setenv("PROTOCOLOS_DIR", str(raiz))
    assert P.configurado()
    assert P.varrer_pasta("robô") == []               # recém-copiado = instável, aguarda
    monkeypatch.setattr(P, "ESTAVEL_SEG", 0)          # simula arquivo já estável
    novos = P.varrer_pasta("robô")
    assert len(novos) == 1
    assert P.varrer_pasta("robô") == []                    # 2ª varredura: nada novo
    p = db.protocolo_get(novos[0])
    assert p["video_origem"] == "sharepoint" and p["status"] == "Pendente"
    with db.Session() as s:
        s.delete(s.get(db.Protocolo, novos[0])); s.commit()


def test_protocolo_reprocessa_sem_retranscrever(client, tmp_path, monkeypatch):
    """Reprocessar com transcrição já existente: NÃO chama o whisper de novo — vai direto
    para a análise IA (ex.: recarga de créditos após falha na análise)."""
    import protocolos as P
    import transcritor, protocolo_ia
    v = _video_fake(tmp_path, "ja_transcrito.mp4")
    pid, _ = db.protocolo_criar("ja_transcrito.mp4", v, "upload", "Tester")
    with db.Session() as s:
        obj = s.get(db.Protocolo, pid)
        obj.transcricao = "[0:05] Conteúdo já transcrito antes da falha."
        s.commit()
    def _nao_pode(*a, **k):
        raise AssertionError("não deveria transcrever de novo")
    monkeypatch.setattr(transcritor, "transcrever_isolado", _nao_pode)
    monkeypatch.setattr(protocolo_ia, "analisar", lambda t, n="": (
        {c: ("Fiscal" if c == "modulo" else "x") for c in db.PROTO_CAMPOS_TEXTO}, "{}"))
    ok, _msg = P.processar(pid, "Tester")
    assert ok and db.protocolo_get(pid)["status"] == "Em revisão"
    with db.Session() as s:
        s.delete(s.get(db.Protocolo, pid)); s.commit()


def test_protocolo_erro_amigavel():
    """Erros comuns da API viram mensagens claras (créditos, chave, sobrecarga)."""
    import protocolos as P
    assert "Créditos da API" in P._erro_amigavel(RuntimeError(
        "Error code: 400 - Your credit balance is too low to access the Anthropic API."))
    assert "Chave da API" in P._erro_amigavel(RuntimeError("authentication_error: invalid x-api-key"))
    assert "sobrecarregada" in P._erro_amigavel(RuntimeError("Error code: 529 - overloaded_error"))
    assert "ValueError" in P._erro_amigavel(ValueError("outra coisa"))


def test_subtelas_tem_contexto_do_projeto(client):
    """Uniformidade do fluxo: TODA subtela do projeto carrega a faixa de contexto
    (cliente + fase + volta à ficha) — o usuário nunca perde onde está."""
    pid = int(_novo(client, cliente="Contexto Ltda", numero_projeto="CTX-1"))
    with db.Session() as s:
        s.get(db.Projeto, pid).etapa = "Projeto"      # etapa que libera o gate da origem
        s.commit()
    telas = ["projeto_cronograma", "projeto_checklist", "projeto_levantamento",
             "projeto_email", "projeto_origem", "projeto_designar"]
    with A.app.test_request_context():
        from flask import url_for
        urls = [url_for(ep, pid=pid) for ep in telas]
    for ep, u in zip(telas, urls):
        r = client.get(u, follow_redirects=True)
        assert r.status_code == 200, "%s -> %s" % (ep, r.status_code)
        html = r.get_data(as_text=True)
        assert "Ficha do projeto" in html and "Contexto Ltda" in html, \
            "%s sem a faixa de contexto do projeto" % ep
    with db.Session() as s:
        s.delete(s.get(db.Projeto, pid)); s.commit()


# ---- Endurecimento (auditoria 2026-07-06) ----

def test_path_dentro_sem_bypass_de_prefixo(tmp_path):
    """Validação de diretório por componente: C:\\Dados não pode casar com C:\\DadosXyz."""
    import _common as C
    raiz = str(tmp_path / "Dados")
    quase = str(tmp_path / "DadosSecretos" / "x.txt")
    dentro = str(tmp_path / "Dados" / "sub" / "x.txt")
    assert C.path_dentro(dentro, raiz)
    assert C.path_dentro(raiz, raiz)                   # a própria raiz
    assert not C.path_dentro(quase, raiz)              # prefixo parecido NÃO passa
    assert not C.path_dentro("", raiz) and not C.path_dentro(None, raiz)


def test_config_endurecida(client):
    """Teto de upload e cookie de sessão endurecido ficam ativos."""
    assert A.app.config["MAX_CONTENT_LENGTH"] == 4096 * 1024 * 1024
    assert A.app.config["SESSION_COOKIE_SAMESITE"] == "Lax"
    assert A.app.config["SESSION_COOKIE_HTTPONLY"] is True
    assert db.engine.pool._pre_ping is True            # conexões mortas não derrubam requests


def test_protocolo_aceita_audio(client, tmp_path):
    """Áudio gravado (.mp3/.wav/...) entra no MESMO fluxo; a revisão mostra player de áudio."""
    import protocolos as P
    assert P.eh_audio("reuniao.mp3") and P.eh_audio("call.WAV")
    assert not P.eh_audio("treino.mp4")
    for e in (".mp3", ".wav", ".m4a", ".ogg", ".opus", ".flac"):
        assert e in P.EXTS                            # formatos de áudio aceitos
    a = tmp_path / "gravacao.mp3"
    a.write_bytes(b"AUDIOFAKE")
    pid, _ = db.protocolo_criar("gravacao.mp3", str(a), "upload", "Tester")
    html = client.get("/protocolos/%s" % pid).get_data(as_text=True)
    assert "<audio" in html and "Áudio original" in html and "<video" not in html
    with db.Session() as s:
        s.delete(s.get(db.Protocolo, pid)); s.commit()


def test_protocolo_status_e_progresso(client, tmp_path, monkeypatch):
    """Linha do tempo: /status devolve o andamento; durante a transcrição expõe o % gravado
    pelo subprocesso no arquivo de progresso; a tela mostra o stepper."""
    import json as _json
    import protocolos as P
    v = _video_fake(tmp_path, "andamento.mp4")
    pid, _ = db.protocolo_criar("andamento.mp4", v, "upload", "Tester")
    # Pendente: status sem pct
    r = client.get("/protocolos/%s/status" % pid)
    assert r.status_code == 200 and r.get_json()["status"] == "Pendente"
    # Transcrevendo + arquivo de progresso -> pct/pos/dur no JSON
    db.protocolo_atualizar_status(pid, "Transcrevendo")
    with open(P._progresso_path(pid), "w", encoding="utf-8") as f:
        _json.dump({"pos": 90, "dur": 300, "pct": 30}, f)
    j = client.get("/protocolos/%s/status" % pid).get_json()
    assert j["status"] == "Transcrevendo" and j["pct"] == 30 and j["dur"] == 300
    # a tela de revisão traz a linha do tempo
    html = client.get("/protocolos/%s" % pid).get_data(as_text=True)
    assert "pt-timeline" in html and "Transcrição" in html
    assert client.get("/protocolos/999999/status").status_code == 404
    os.remove(P._progresso_path(pid))
    with db.Session() as s:
        s.delete(s.get(db.Protocolo, pid)); s.commit()


def test_protocolo_rotas_e_revisao(client, tmp_path):
    """Telas: consulta renderiza; revisão mostra o protocolo; salvar edição e aprovar."""
    v = _video_fake(tmp_path, "tela.mp4")
    pid, _ = db.protocolo_criar("tela.mp4", v, "upload", "Tester")
    db.protocolo_atualizar_status(pid, "Em revisão")
    r = client.get("/protocolos")
    assert r.status_code == 200 and "Protocolos de Treinamento" in r.get_data(as_text=True)
    r = client.get("/protocolos?modulo=Fiscal&q=regra")     # filtros não quebram
    assert r.status_code == 200
    r = client.get("/protocolos/%s" % pid)
    assert r.status_code == 200 and "tela.mp4" in r.get_data(as_text=True)
    r = client.post("/protocolos/%s/salvar" % pid,
                    data={"titulo": "Título Editado", "modulo": "Estoque", "menu": "14-I"})
    assert r.status_code == 302
    p = db.protocolo_get(pid)
    assert p["titulo"] == "Título Editado" and p["modulo"] == "Estoque"
    r = client.post("/protocolos/%s/aprovar" % pid)
    assert r.status_code == 302 and db.protocolo_get(pid)["status"] == "Aprovado"
    assert client.get("/protocolos/999999").status_code == 404
    with db.Session() as s:
        s.delete(s.get(db.Protocolo, pid)); s.commit()


def test_doc_ver_espelho_pdf(client, monkeypatch, tmp_path):
    """O 'Ver' mostra o PDF FIEL (iframe) quando disponível; sem PDF, cai na visão HTML."""
    import os as _os, docview, _common as C
    from docx import Document as _D
    pid = int(_novo(client, cliente="VerPDF LTDA", cnpj="00.000.000/0001-00", numero_projeto="VP-1",
                    modulos="FAT", horas_cobradas="10"))
    doc_path = _os.path.join(C.DATA_WRITE, "ver_teste_%s.docx" % pid)
    _d = _D(); _d.add_paragraph("conteudo do documento"); _d.save(doc_path)
    with db.Session() as s:
        d = db.Documento(projeto_id=pid, tipo="projeto", arquivo=_os.path.basename(doc_path), caminho=doc_path)
        s.add(d); s.commit(); did = d.id
    # (a) com PDF disponível -> iframe apontando para a rota /pdf
    pdf_fake = _os.path.join(str(tmp_path), "f.pdf"); open(pdf_fake, "wb").write(b"%PDF-1.4\n1 0 obj<<>>endobj\n%%EOF\n")
    monkeypatch.setattr(docview, "to_pdf", lambda p: pdf_fake)
    r = client.get("/projetos/%s/doc/%s/ver" % (pid, did))
    body = r.get_data(as_text=True)
    assert r.status_code == 200 and "<iframe" in body and "/doc/%s/pdf" % did in body
    r2 = client.get("/projetos/%s/doc/%s/pdf" % (pid, did))
    assert r2.status_code == 200 and r2.mimetype == "application/pdf"
    # (b) sem PDF (None) -> cai no HTML (sem iframe, mostra o conteúdo)
    monkeypatch.setattr(docview, "to_pdf", lambda p: None)
    r3 = client.get("/projetos/%s/doc/%s/ver" % (pid, did))
    assert r3.status_code == 200 and "<iframe" not in r3.get_data(as_text=True)
    try:
        _os.remove(doc_path)
    except OSError:
        pass
    client.post("/projetos/%s/excluir" % pid)


def test_excluir_documento_respeita_fluxo(client):
    """Excluir documento gerado: não exclui um documento se existe outro que depende dele.
    Primeiro exclui o Projeto, depois o Levantamento."""
    pid = int(_novo(client, cliente="DelDoc LTDA", cnpj="00.000.000/0001-00", numero_projeto="DD-1",
                    modulos="FAT", horas_cobradas="10"))
    with db.Session() as s:
        lev = db.Documento(projeto_id=pid, tipo="levantamento", arquivo="lev.docx", caminho="")
        proj = db.Documento(projeto_id=pid, tipo="projeto", arquivo="proj.docx", caminho="")
        s.add_all([lev, proj]); s.commit()
        lev_id, proj_id = lev.id, proj.id
    # regra de dependência
    ok, motivo = db.pode_excluir_documento(pid, "levantamento")
    assert ok is False and "Projeto" in motivo
    assert db.pode_excluir_documento(pid, "projeto")[0] is True
    # 1) excluir o Levantamento com o Projeto existente -> bloqueado (documento permanece)
    r = client.post("/projetos/%s/doc/%s/excluir" % (pid, lev_id))
    assert r.status_code == 302
    with db.Session() as s:
        assert s.get(db.Documento, lev_id) is not None
    # 2) excluir o Projeto -> ok
    r = client.post("/projetos/%s/doc/%s/excluir" % (pid, proj_id))
    assert r.status_code == 302
    with db.Session() as s:
        assert s.get(db.Documento, proj_id) is None
    # 3) agora o Levantamento pode ser excluído
    r = client.post("/projetos/%s/doc/%s/excluir" % (pid, lev_id))
    assert r.status_code == 302
    with db.Session() as s:
        assert s.get(db.Documento, lev_id) is None
    client.post("/projetos/%s/excluir" % pid)


def test_detalhamento_so_areas_contratadas(client):
    """No Projeto, o Detalhamento das Rotinas mantém SÓ as áreas dos módulos contratados."""
    import gerar_layout
    pid = _novo(client, cliente="Areas LTDA", cnpj="00.000.000/0001-00", numero_projeto="A-1",
                modulos="FAT", horas_cobradas="10")
    db.levantamento_seed(int(pid), "FAT")
    with db.Session() as s:
        proj = db.to_dict(s.get(db.Projeto, int(pid)))
    path = gerar_layout.gerar("projeto", proj)
    from docx import Document
    txt = "\n".join(p.text for p in Document(path).paragraphs)
    assert "Vendas e Faturamento" in txt        # área contratada (FAT) permanece
    assert "Livros Fiscais" not in txt           # área não contratada foi removida
    assert "Controle de Estoque" not in txt      # idem
    try:
        os.remove(path)
    except OSError:
        pass
    client.post("/projetos/%s/excluir" % pid)


def test_cronograma_atividades_seed_do_checklist(client):
    """Agendador de visitas: deriva as Visitas (módulo+seq) do Check List dos módulos
    contratados, agrupa por visita e permite alocar/desalocar uma atividade."""
    pid = _novo(client, cliente="Agenda LTDA", cnpj="00.000.000/0001-00", numero_projeto="AG-1",
                modulos="FAT", horas_cobradas="10")
    n = db.cronograma_atividades_seed(int(pid), "FAT")
    assert n > 0
    assert db.cronograma_atividades_seed(int(pid), "FAT") == n      # idempotente
    ats = db.cronograma_atividades(int(pid))
    assert ats and all(a["modulo"] == "FAT" and a["seq"] >= 1 for a in ats)
    visitas = db.cronograma_visitas(int(pid))
    assert visitas and visitas[0]["titulo"].startswith("Visita V")
    assert all(v["atividades"] for v in visitas)                    # toda visita tem atividade
    aid = ats[0]["id"]
    upd = db.cronograma_alocar(aid, data="2026-07-01", turno="manha", tecnico="Ana")
    assert upd["data"] == "2026-07-01" and upd["turno"] == "manha" and upd["tecnico"] == "Ana"
    db.cronograma_alocar(aid, data="", turno="")                    # desaloca
    assert next(a for a in db.cronograma_atividades(int(pid)) if a["id"] == aid)["data"] == ""
    client.post("/projetos/%s/excluir" % pid)


def test_agenda_visitas_render_e_aloca(client):
    """Agenda de Visitas: a tela renderiza e o endpoint aloca uma atividade num slot;
    outro projeto não consegue mover a atividade (guarda de propriedade)."""
    pid = _novo(client, cliente="Visita Cal LTDA", cnpj="00.000.000/0001-00", numero_projeto="VC-1",
                modulos="FAT", horas_cobradas="10", etapa="Cronograma e Check-list")
    import datetime
    d1 = (datetime.date.today() + datetime.timedelta(days=7)).isoformat()
    d2 = (datetime.date.today() + datetime.timedelta(days=8)).isoformat()
    r = client.get("/projetos/%s/agenda" % pid)
    assert r.status_code == 200 and "Agenda de Visitas" in r.get_data(as_text=True)
    ats = db.cronograma_atividades(int(pid))
    assert ats
    aid = ats[0]["id"]
    r = client.post("/projetos/%s/agenda/alocar" % pid,
                    data={"atividade_id": aid, "data": d1, "turno": "manha"})
    assert r.status_code == 200 and r.get_json()["ok"] is True
    upd = next(a for a in db.cronograma_atividades(int(pid)) if a["id"] == aid)
    assert upd["data"] == d1 and upd["turno"] == "manha"
    pid2 = _novo(client, cliente="Outro Proj LTDA", modulos="FAT", etapa="Cronograma e Check-list")
    r = client.post("/projetos/%s/agenda/alocar" % pid2,
                    data={"atividade_id": aid, "data": d2, "turno": "tarde"})
    assert r.status_code == 404                                       # não pertence ao projeto 2
    client.post("/projetos/%s/excluir" % pid)
    client.post("/projetos/%s/excluir" % pid2)


def test_agenda_polimentos(client):
    """Polimentos: atualizar só o técnico preserva a alocação; postergar pula o fim de
    semana; o filtro de fim de semana mostra Sáb/Dom."""
    pid = _novo(client, cliente="Polish LTDA", cnpj="00.000.000/0001-00", numero_projeto="PO-1",
                modulos="FAT", horas_cobradas="10", etapa="Cronograma e Check-list")
    import datetime
    pol_d = (datetime.date.today() + datetime.timedelta(days=7)).isoformat()
    post_d = (datetime.date.today() + datetime.timedelta(days=14)).isoformat()
    db.cronograma_atividades_seed(int(pid), "FAT")
    aid = db.cronograma_atividades(int(pid))[0]["id"]
    client.post("/projetos/%s/agenda/alocar" % pid,            # dia útil futuro, de manhã
                data={"atividade_id": aid, "data": pol_d, "turno": "manha"})
    r = client.post("/projetos/%s/agenda/alocar" % pid,        # só técnico (sem data/turno)
                    data={"atividade_id": aid, "tecnico": "Fulano"})
    assert r.status_code == 200 and r.get_json()["ok"] is True
    a = next(x for x in db.cronograma_atividades(int(pid)) if x["id"] == aid)
    assert a["data"] == pol_d and a["turno"] == "manha" and a["tecnico"] == "Fulano"
    r = client.post("/projetos/%s/agenda/postergar" % pid,     # posterga este assunto p/ data futura
                    data={"atividade_id": aid, "nova_data": post_d, "novo_turno": "tarde"})
    assert r.status_code == 302
    a = next(x for x in db.cronograma_atividades(int(pid)) if x["id"] == aid)
    assert a["status"] == "Postergada" and a["data"] == pol_d   # original fica no lugar (histórico)
    novos = [x for x in db.cronograma_atividades(int(pid)) if x["data"] == post_d and x["turno"] == "tarde"]
    assert novos and novos[0]["status"] == "Agendada"                  # nova ocorrência criada no destino
    r = client.get("/projetos/%s/agenda?fds=1&ref=%s" % (pid, post_d))
    assert r.status_code == 200 and "Sáb" in r.get_data(as_text=True)
    client.post("/projetos/%s/excluir" % pid)


def test_agenda_bloqueia_data_passada(client):
    """Montagem do cronograma: não permite alocar assunto em data já passada."""
    import datetime
    pid = _novo(client, cliente="Passado LTDA", cnpj="00.000.000/0001-00", numero_projeto="PA-1",
                modulos="FAT", horas_cobradas="10", etapa="Cronograma e Check-list")
    db.cronograma_atividades_seed(int(pid), "FAT")
    aid = db.cronograma_atividades(int(pid))[0]["id"]
    ontem = (datetime.date.today() - datetime.timedelta(days=1)).isoformat()
    r = client.post("/projetos/%s/agenda/alocar" % pid,
                    data={"atividade_id": aid, "data": ontem, "turno": "manha"})
    assert r.status_code == 409 and r.get_json()["ok"] is False
    assert "passada" in (r.get_json().get("erro") or "").lower()
    a = next(x for x in db.cronograma_atividades(int(pid)) if x["id"] == aid)
    assert a["data"] == ""                      # permaneceu não alocada
    client.post("/projetos/%s/excluir" % pid)


def test_agenda_bloqueia_por_ocupacao_sicla(client, monkeypatch):
    """Casa o Código SICLA do consultor (cadastro) com a coluna 'tecnico' do SELECT e
    bloqueia o slot ocupado; um turno livre é liberado."""
    import datetime
    import disponibilidade as D
    # projeto + designação + seed enquanto o login está inativo (ainda sem usuários)
    pid = _novo(client, cliente="Ocup LTDA", cnpj="00.000.000/0001-00", numero_projeto="OC-1",
                modulos="FAT", horas_cobradas="10", etapa="Cronograma e Check-list")
    db.cronograma_atividades_seed(int(pid), "FAT")
    aid = db.cronograma_atividades(int(pid))[0]["id"]
    with db.Session() as s:
        s.add(db.Designacao(projeto_id=int(pid), modulo="FAT", consultor="Tecnico Ze"))
        s.add(db.Usuario(nome="Tecnico Ze", login="ze@x.com", email="ze@x.com",
                         perfil="Consultor", codigo_sicla="Z9", ativo=1))
        s.commit()
    # com usuário no banco o login fica ativo -> autentica o client como ADM
    with client.session_transaction() as sess:
        sess["auth"] = True; sess["perfil"] = "ADM"; sess["perfil_nome"] = "ADM"
    dia = (datetime.date.today() + datetime.timedelta(days=3)).isoformat()
    monkeypatch.setattr(D, "configurado", lambda: True)
    monkeypatch.setattr(D, "ocupacao_por_slot",
                        lambda i, f, tecnicos=None, cfg=None: {("z9", dia, "manha"): True})
    r = client.post("/projetos/%s/agenda/alocar" % pid,
                    data={"atividade_id": aid, "data": dia, "turno": "manha"})
    assert r.status_code == 409 and "ocupado" in (r.get_json().get("erro") or "").lower()
    r2 = client.post("/projetos/%s/agenda/alocar" % pid,        # mesmo dia, turno livre
                     data={"atividade_id": aid, "data": dia, "turno": "tarde"})
    assert r2.status_code == 200 and r2.get_json()["ok"] is True
    # cleanup: remove o usuário (restaura login inativo) e o projeto
    with db.Session() as s:
        u = s.query(db.Usuario).filter(db.Usuario.email == "ze@x.com").first()
        if u:
            s.delete(u)
        s.commit()
    client.post("/projetos/%s/excluir" % pid)


def test_agenda_status_acompanhamento_e_xlsx(client):
    """Fase 3/4: marcar 'Realizada', ver acompanhamento e gerar o cronograma .xlsx das alocações."""
    pid = _novo(client, cliente="Visita XLS LTDA", cnpj="00.000.000/0001-00", numero_projeto="VX-1",
                modulos="FAT", horas_cobradas="10", etapa="Cronograma e Check-list")
    db.cronograma_atividades_seed(int(pid), "FAT")
    # sem alocação -> gerar é bloqueado (redireciona à agenda, sem documento)
    r = client.post("/projetos/%s/agenda/gerar" % pid)
    assert r.status_code == 302 and "/agenda" in r.headers["Location"]
    with db.Session() as s:
        assert s.query(db.Documento).filter_by(projeto_id=int(pid), tipo="cronograma").count() == 0
    # aloca uma atividade e marca como realizada
    aid = db.cronograma_atividades(int(pid))[0]["id"]
    db.cronograma_alocar(aid, projeto_id=int(pid), data="2026-07-06", turno="manha")
    r = client.post("/projetos/%s/agenda/status" % pid, data={"atividade_id": aid, "status": "Realizada"})
    assert r.status_code == 200 and r.get_json()["ok"] is True
    assert next(a for a in db.cronograma_atividades(int(pid)) if a["id"] == aid)["status"] == "Realizada"
    r = client.get("/projetos/%s/agenda/acompanhamento" % pid)
    assert r.status_code == 200 and "Acompanhamento" in r.get_data(as_text=True)
    # gera o .xlsx e anexa como Documento de cronograma
    r = client.post("/projetos/%s/agenda/gerar" % pid)
    assert r.status_code == 302
    with db.Session() as s:
        docs = s.query(db.Documento).filter_by(projeto_id=int(pid), tipo="cronograma").all()
        assert len(docs) == 1
        path = docs[0].caminho
    assert path.endswith(".xlsx") and os.path.exists(path)
    from openpyxl import load_workbook
    blob = "\n".join(str(c.value) for row in load_workbook(path).active.iter_rows()
                     for c in row if c.value)
    assert "FAT" in blob and "06/07/2026" in blob
    try:
        os.remove(path)
    except OSError:
        pass
    client.post("/projetos/%s/excluir" % pid)


def test_agenda_horario_global_status_e_tecnico_modulo(client):
    """Horário GLOBAL por turno; contagem/filtro por status; técnico por módulo sincroniza Designação."""
    pid = _novo(client, cliente="Horario LTDA", cnpj="00.000.000/0001-00", numero_projeto="HR-1",
                modulos="FAT", horas_cobradas="10", etapa="Cronograma e Check-list")
    db.cronograma_atividades_seed(int(pid), "FAT")
    ats = db.cronograma_atividades(int(pid))
    db.cronograma_alocar(ats[0]["id"], projeto_id=int(pid), data="2026-07-06", turno="manha")
    db.cronograma_alocar(ats[1]["id"], projeto_id=int(pid), data="2026-07-07", turno="tarde")
    # horário GLOBAL do turno manhã (um só para todas as visitas)
    r = client.post("/projetos/%s/agenda/horario" % pid,
                    data={"turno": "manha", "hora_inicio": "09:00", "hora_fim": "11:30"})
    assert r.status_code == 200 and r.get_json()["ok"] is True
    assert db.cronograma_horarios(int(pid))["manha"] == ("09:00", "11:30")
    html = client.get("/projetos/%s/agenda/acompanhamento" % pid).get_data(as_text=True)
    assert "09:00" in html and "11:30" in html and html.count('class="ac-status"') == 2
    # status: marca uma Realizada e filtra
    r = client.post("/projetos/%s/agenda/status" % pid, data={"atividade_id": ats[0]["id"], "status": "Realizada"})
    assert r.status_code == 200 and r.get_json()["ok"] is True
    html = client.get("/projetos/%s/agenda/acompanhamento?status=Realizada" % pid).get_data(as_text=True)
    assert html.count('class="ac-status"') == 1                # só a Realizada
    # status inválido é rejeitado
    assert client.post("/projetos/%s/agenda/status" % pid,
                       data={"atividade_id": ats[0]["id"], "status": "Inexistente"}).status_code == 400
    # técnico por módulo: aplica aos cartões e sincroniza a Designação
    r = client.post("/projetos/%s/agenda/tecnico_modulo" % pid, data={"modulo": "FAT", "tecnico": "Chico"})
    assert r.status_code == 302
    assert all(a["tecnico"] == "Chico" for a in db.cronograma_atividades(int(pid)) if a["modulo"] == "FAT")
    assert {d["modulo"]: d["consultor"] for d in db.designacoes_do_projeto(int(pid))}.get("FAT") == "Chico"
    client.post("/projetos/%s/excluir" % pid)


def test_agenda_alocar_visita_inteira(client):
    """Alocar a visita inteira: todas as atividades pendentes de (modulo, seq) vão ao slot;
    a sidebar renderiza o acordeão (recolher por módulo/visita)."""
    pid = _novo(client, cliente="Bloco LTDA", cnpj="00.000.000/0001-00", numero_projeto="BL-1",
                modulos="FAT", horas_cobradas="10", etapa="Cronograma e Check-list")
    db.cronograma_atividades_seed(int(pid), "FAT")
    g = db.cronograma_visitas(int(pid))[0]
    seq, nat = g["seq"], len(g["atividades"])
    r = client.post("/projetos/%s/agenda/alocar_visita" % pid,
                    data={"modulo": "FAT", "seq": seq, "data": "2026-07-06", "turno": "manha"})
    assert r.status_code == 200 and r.get_json()["ok"] is True and r.get_json()["n"] == nat
    alocadas = [a for a in db.cronograma_atividades(int(pid)) if a["modulo"] == "FAT" and a["seq"] == seq]
    assert alocadas and all(a["data"] == "2026-07-06" and a["turno"] == "manha"
                            and a["status"] == "Agendada" for a in alocadas)
    html = client.get("/projetos/%s/agenda" % pid).get_data(as_text=True)
    assert "Recolher tudo" in html and "ag-visita-d" in html      # acordeão na sidebar
    client.post("/projetos/%s/excluir" % pid)


def test_config_disponibilidade(client):
    """Tela de Disponibilidade (ADM): monta a URL pelos campos, prioriza URL completa,
    e a página abre."""
    import disponibilidade as D
    url = D._build_url({"tipo": "postgresql", "host": "h", "porta": "5432",
                        "banco": "b", "usuario": "u", "senha": "p@ss"})
    assert url.startswith("postgresql+psycopg2://u:") and "@h:5432/b" in url
    assert D._build_url({"url": "sqlite:///x.db", "tipo": "postgresql"}) == "sqlite:///x.db"
    assert D.configurado() in (True, False)
    assert client.get("/config/disponibilidade").status_code == 200


def test_disponibilidade_filtra_por_tecnico(monkeypatch):
    """Detecta se o SELECT usa o filtro :tecnicos (define se amplia a janela de datas)."""
    import disponibilidade as D
    monkeypatch.setattr(D, "load_cfg", lambda: {"select": "SELECT c AS tecnico FROM a WHERE c IN :tecnicos"})
    assert D.filtra_por_tecnico() is True
    monkeypatch.setattr(D, "load_cfg", lambda: {"select": "SELECT c AS tecnico FROM a"})
    assert D.filtra_por_tecnico() is False


def test_disponibilidade_ocupacao_repassa_tecnicos(monkeypatch):
    """ocupacao_por_slot repassa os códigos dos consultores para consultar (filtro no banco)."""
    import disponibilidade as D
    capturado = {}
    def fake_consultar(di, df, tecnicos=None, cfg=None):
        capturado["tecnicos"] = tecnicos
        return [{"tecnico": "Z9", "data": "2027-01-05", "turno": "manha"}]
    monkeypatch.setattr(D, "consultar", fake_consultar)
    ocup = D.ocupacao_por_slot("2027-01-01", "2027-06-01", ["Z9", "A1"])
    assert capturado["tecnicos"] == ["Z9", "A1"]
    assert ocup.get(("z9", "2027-01-05", "manha")) is True


def test_disponibilidade_cache_ttl(monkeypatch):
    """ocupacao_por_slot_cache: mesma janela+técnicos dentro do TTL consulta o banco 1 vez."""
    import disponibilidade as D
    chamadas = {"n": 0}
    def fake(di, df, tecnicos=None, cfg=None):
        chamadas["n"] += 1
        return {("z9", di, "manha"): True}
    monkeypatch.setattr(D, "ocupacao_por_slot", fake)
    D._CACHE.clear()
    a = D.ocupacao_por_slot_cache("2027-02-01", "2027-08-01", ["Z9"])
    b = D.ocupacao_por_slot_cache("2027-02-01", "2027-08-01", ["z9"])   # mesma chave (case-insensitive)
    assert a == b and chamadas["n"] == 1
    D.ocupacao_por_slot_cache("2027-02-01", "2027-08-01", ["A1"])       # técnicos diferentes -> nova consulta
    assert chamadas["n"] == 2
    D._CACHE.clear()


def test_docview_preaquecer(monkeypatch, tmp_path):
    """preaquecer() converte o PDF em segundo plano para .docx e ignora outros tipos."""
    import threading, docview
    feito = threading.Event()
    monkeypatch.setattr(docview, "to_pdf", lambda p: feito.set())
    monkeypatch.setattr(docview, "_limpar_cache_antigo", lambda dias=30: None)
    docview.preaquecer(str(tmp_path / "doc.docx"))
    assert feito.wait(5)                                # converteu em background
    feito.clear()
    docview.preaquecer(str(tmp_path / "planilha.xlsx"))  # não-docx: não dispara
    assert not feito.wait(0.3)


def test_agenda_disponibilidade_modos(client, monkeypatch):
    """Disponibilidade: análise CONJUNTA bloqueia se qualquer envolvido está ocupado;
    INDIVIDUAL bloqueia só conforme o técnico escolhido."""
    import disponibilidade as D
    import datetime as _dt
    pid = _novo(client, cliente="Disp LTDA", cnpj="00.000.000/0001-00", numero_projeto="DP-1",
                modulos="FAT, EST", horas_cobradas="10", etapa="Cronograma e Check-list")
    db.cronograma_atividades_seed(int(pid), "FAT, EST")
    client.post("/projetos/%s/agenda/tecnico_modulo" % pid, data={"modulo": "FAT", "tecnico": "Ana"})
    client.post("/projetos/%s/agenda/tecnico_modulo" % pid, data={"modulo": "EST", "tecnico": "Beto"})
    # vínculo por Código SICLA: Ana e Beto precisam de cadastro com código (casa com o SELECT)
    with db.Session() as s:
        s.add(db.Usuario(nome="Ana", login="ana@x.com", email="ana@x.com",
                         perfil="Consultor", codigo_sicla="AN1", ativo=1))
        s.add(db.Usuario(nome="Beto", login="beto@x.com", email="beto@x.com",
                         perfil="Consultor", codigo_sicla="BT1", ativo=1))
        s.commit()
    with client.session_transaction() as sess:        # com usuários, o login fica ativo
        sess["auth"] = True; sess["perfil"] = "ADM"; sess["perfil_nome"] = "ADM"
    hoje = _dt.date.today()
    seg = (hoje - _dt.timedelta(days=hoje.weekday())).isoformat()      # 2ª da semana
    monkeypatch.setattr(D, "configurado", lambda: True)
    monkeypatch.setattr(D, "filtra_por_tecnico", lambda cfg=None: True)
    monkeypatch.setattr(D, "ocupacao_por_slot",
                        lambda di, df, tecnicos=None, cfg=None: {("an1", seg, "manha"): True})  # casa por código
    D._CACHE.clear()                                     # a view usa o cache TTL — começa limpo
    base = "/projetos/%s/agenda?ref=%s" % (pid, seg)
    assert "Ocupado: Ana" in client.get(base).get_data(as_text=True)                       # conjunta
    assert "Ocupado:" not in client.get(base + "&modo=individual&tec=Beto").get_data(as_text=True)   # Beto livre
    assert "Ocupado: Ana" in client.get(base + "&modo=individual&tec=Ana").get_data(as_text=True)    # Ana ocupada
    with db.Session() as s:
        for em in ("ana@x.com", "beto@x.com"):
            u = s.query(db.Usuario).filter(db.Usuario.email == em).first()
            if u:
                s.delete(u)
        s.commit()
    client.post("/projetos/%s/excluir" % pid)
