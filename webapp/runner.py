# -*- coding: utf-8 -*-
"""Executor: roda os geradores, o importador, o conversor e o verificador.
Ciente de empacotamento (.exe): grava em pasta gravável via _common."""
import os
import re
import io
import sys
import contextlib
import importlib

HERE = os.path.dirname(os.path.abspath(__file__))
if not getattr(sys, "frozen", False):
    TOOLS = os.path.join(os.path.dirname(HERE), "tools")
    if TOOLS not in sys.path:
        sys.path.insert(0, TOOLS)

import _common as C   # noqa: E402

DATA = C.DATA_WRITE   # pasta gravável (= tools/data no modo normal)


def run_generator(modname, yaml_basename=None):
    """Roda gerar_*.main(); retorna (caminho_do_arquivo, log)."""
    mod = importlib.import_module(modname)
    buf = io.StringIO()
    with contextlib.redirect_stdout(buf):
        mod.main("data/" + yaml_basename) if yaml_basename else mod.main()
    out = buf.getvalue().strip()
    matches = re.findall(r"-> (.+)", out)
    return (matches[-1].strip() if matches else None), out


def save_upload_yaml(file_storage, slug_fn):
    """Salva um .yaml enviado em <gravável>/upload_<slug>.yaml e devolve o basename."""
    base = "upload_" + slug_fn(os.path.splitext(file_storage.filename)[0]) + ".yaml"
    file_storage.save(os.path.join(DATA, base))
    return base


def run_sequencia(docx_path):
    """Importa o Levantamento e gera EM SEQUÊNCIA: Projeto + Check List + Termo,
    aplicando tempo verbal + ortografia. Retorna um dict com os caminhos."""
    import importar_mapeamento as I
    import catalogo as CAT
    import yaml as _yaml
    data = I.extract(docx_path)
    ydict = I.to_yaml_dict(data, aplicar_verbal=True)      # verbal + ortografia
    nome = data.get("client_name") or "cliente"
    slug = C.slug(nome)
    modulos = data.get("modulos") or []

    # 1) Projeto de Implantação
    proj_base = "projeto_%s.yaml" % slug
    with open(os.path.join(DATA, proj_base), "w", encoding="utf-8") as f:
        f.write("# Importação automática (tempo verbal + ortografia). Revisar antes de usar.\n")
        _yaml.safe_dump(ydict, f, allow_unicode=True, sort_keys=False, width=100)
    proj_path, _ = run_generator("gerar_projeto_implantacao", proj_base)

    # 2) Check List do Consultor (a partir dos módulos do levantamento)
    lev_base = "lev_%s.yaml" % slug
    with open(os.path.join(DATA, lev_base), "w", encoding="utf-8") as f:
        _yaml.safe_dump({"cliente": nome, "modulos_contratados": modulos},
                        f, allow_unicode=True, sort_keys=False)
    chk_path, _ = run_generator("gerar_checklist_consultor", lev_base)

    # 3) Termo de Encerramento (quadro de módulos a partir do catálogo)
    mods, _ = CAT.resolve(modulos)
    resumo = [{"modulo": m.get("descricao") or m.get("abrev"), "adicional": "",
               "processo": "Implantado", "status_uso": "Sim", "obs": ""} for m in mods]
    termo_base = "termo_%s.yaml" % slug
    with open(os.path.join(DATA, termo_base), "w", encoding="utf-8") as f:
        _yaml.safe_dump({"cliente": nome, "resumo_modulos": resumo, "alteracoes": [],
                         "observacoes": [], "pendencias": [],
                         "cidade_data": "Novo Hamburgo, ____ de ____________ de 2026."},
                        f, allow_unicode=True, sort_keys=False)
    termo_path, _ = run_generator("gerar_termo_encerramento", termo_base)

    return {"cliente": nome, "modulos": len(modulos), "projeto": proj_path,
            "checklist": chk_path, "termo": termo_path,
            "yaml": os.path.join(DATA, proj_base)}


def converter_verbal(texto):
    """Presente→Futuro + ortografia (offline) e, havendo chave de API, reconferência
    pela IA (Claude). Retorna (texto, mudanças_verbais do offline)."""
    import conversor_verbal as V
    import ortografia as O
    novo, mudancas = V.converter_texto(texto)
    novo = O.corrigir(novo)
    try:
        import ia
        novo = ia.revisar(novo)          # reconferência IA (no-op sem chave)
    except Exception:
        pass
    return novo, mudancas


def converter_docx(file_storage):
    """Lê um .docx, aplica tempo verbal + ortografia em todo o texto (preservando o
    layout) e salva uma cópia '(corrigido)'. Retorna o caminho."""
    import os
    import conversor_verbal as V
    import ortografia as O
    from docx import Document

    slug = C.slug(os.path.splitext(file_storage.filename or "documento")[0])
    src = os.path.join(C.DATA_WRITE, "_in_" + slug + ".docx")
    file_storage.save(src)
    doc = Document(src)

    runs = []

    def coleta(par):
        for r in par.runs:
            if r.text:
                runs.append(r)

    for p in doc.paragraphs:
        coleta(p)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    coleta(p)

    textos = [O.corrigir(V.converter(r.text)) for r in runs]   # offline (verbal+ortografia)
    try:
        import ia
        if ia.disponivel():
            textos = ia.revisar_lote(textos)                   # reconferência IA em lote
    except Exception:
        pass
    for r, t in zip(runs, textos):
        r.text = t

    C.ensure_out()
    out = os.path.join(C.OUT, "%s_corrigido.docx" % slug)
    doc.save(out)
    return out


def run_saude():
    import verificar
    verificar.R.clear()
    buf = io.StringIO()
    with contextlib.redirect_stdout(buf):
        code = verificar.main()
    return code, buf.getvalue()


def catalogo_por_area():
    """[(area, [módulos])] de TODO o catálogo, para a tela de seleção."""
    import catalogo as CAT
    return CAT.por_area(CAT.load())


def _yaml_de_form(form, modulos):
    """Escreve um YAML (cliente + identificação + modulos_contratados) e devolve o basename."""
    import yaml as _yaml
    nome = (form.get("cliente") or "Cliente").strip()
    doc = {
        "cliente": nome,
        "data": form.get("data", ""),
        "responsaveis": form.get("responsaveis", ""),
        "identificacao": {
            "razao_social": nome,
            "ramo": form.get("ramo", ""),
            "produto": form.get("produto", ""),
            "fornecedor_atual": form.get("fornecedor_atual", ""),
            "localizacao": form.get("localizacao", ""),
            "observacoes_objetivos": form.get("observacoes", ""),
        },
        "modulos_contratados": modulos,
    }
    base = "lev_" + C.slug(nome) + ".yaml"
    with open(os.path.join(DATA, base), "w", encoding="utf-8") as f:
        _yaml.safe_dump(doc, f, allow_unicode=True, sort_keys=False, width=120)
    return base


def gerar_levantamento_form(form, modulos):
    """Gera SÓ o Levantamento (Word). Retorna (caminho, log)."""
    return run_generator("gerar_levantamento", _yaml_de_form(form, modulos))


def gerar_checklist_form(form, modulos):
    """Gera SÓ o Check List do Consultor (Excel) — após o Projeto, módulos finais."""
    return run_generator("gerar_checklist_consultor", _yaml_de_form(form, modulos))


_MESES = ["janeiro", "fevereiro", "março", "abril", "maio", "junho",
          "julho", "agosto", "setembro", "outubro", "novembro", "dezembro"]


def criar_templates(form):
    """Tela 'Criação dos Templates': monta os YAMLs do Mapeamento (Levantamento)
    e do Termo a partir do formulário tabbado e gera os documentos marcados."""
    import yaml as _yaml
    g = form.getlist
    cliente = (form.get("cliente") or "Cliente").strip()
    slug = C.slug(cliente)
    data = "%s/%s/%s" % (form.get("data_dia", ""), form.get("data_mes", ""), form.get("data_ano", ""))
    res = {"ok": True, "mapa": None, "termo": None}

    def rows(*names):
        out = []
        for vals in zip(*[g(n) for n in names]):
            if any((v or "").strip() for v in vals):
                out.append(vals)
        return out

    if form.get("gen_mapa"):
        modb = [{"modulo": m, "necessidade": nec, "obs": o} for m, nec, o in rows("mb_mod", "mb_nec", "mb_obs")]
        lev = {
            "cliente": cliente, "data": data, "responsaveis": form.get("responsaveis", ""),
            "total_usuarios": form.get("total_usuarios", ""),
            "identificacao": {
                "razao_social": cliente, "ramo": form.get("ramo", ""),
                "produto": form.get("produto", ""), "fornecedor_atual": form.get("fornecedor_atual", ""),
                "localizacao": form.get("localizacao", ""), "observacoes_objetivos": form.get("observacoes", ""),
            },
            "usuarios": [{"nome": n, "email": e, "atribuicoes": a} for n, e, a in rows("u_nome", "u_email", "u_atrib")],
            "modulos_contratados": [x for x in g("modulos") if x and x.strip()],
            "modulos_identificados": modb,
            "conversoes": {"horas": form.get("cv_total", ""), "estimativas": {
                "clientes_fornecedores": form.get("cv_cli", ""), "produtos": form.get("cv_prod", ""),
                "financeiro": form.get("cv_fin", ""), "notas_fiscais": form.get("cv_nf", ""),
                "folha": form.get("cv_folha", "")}},
            "horas": {"cobradas": form.get("h_cob", ""), "bonificadas": form.get("h_bon", ""), "total": form.get("h_tot", "")},
        }
        base = "lev_%s.yaml" % slug
        with open(os.path.join(DATA, base), "w", encoding="utf-8") as f:
            _yaml.safe_dump(lev, f, allow_unicode=True, sort_keys=False, width=200)
        res["mapa"], _ = run_generator("gerar_levantamento", base)

    if form.get("gen_termo"):
        mes = form.get("data_mes", "")
        mes_nome = _MESES[int(mes) - 1] if mes.isdigit() and 1 <= int(mes) <= 12 else ""
        termo = {
            "cliente": cliente, "numero_projeto": form.get("numero_projeto", ""),
            "data_validacao": form.get("data_validacao", ""),
            "alteracoes": [{"modulo": m, "rotina": r, "descricao": dsc}
                           for m, r, dsc in rows("alt_modulo", "alt_rotina", "alt_desc")],
            "resumo_modulos": [{"modulo": m, "adicional": a, "processo": p, "status_uso": s, "obs": o}
                               for m, a, p, s, o in rows("tr_modulo", "tr_adic", "tr_proc", "tr_status", "tr_obs")],
            "observacoes": [l.strip() for l in (form.get("termo_obs", "") or "").splitlines() if l.strip()],
            "pendencias": [{"pendencia": p, "tecnico": t, "detalhamento": d2}
                           for p, t, d2 in rows("pend_p", "pend_t", "pend_d")],
            "cidade_data": "Novo Hamburgo, %s de %s de %s." % (form.get("data_dia", ""), mes_nome, form.get("data_ano", "")),
        }
        base = "termo_%s.yaml" % slug
        with open(os.path.join(DATA, base), "w", encoding="utf-8") as f:
            _yaml.safe_dump(termo, f, allow_unicode=True, sort_keys=False, width=200)
        res["termo"], _ = run_generator("gerar_termo_encerramento", base)

    if not (res["mapa"] or res["termo"]):
        return {"ok": False, "erro": "Marque ao menos um documento a gerar."}
    return res
